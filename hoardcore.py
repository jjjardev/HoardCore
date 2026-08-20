#!/usr/bin/env python3
"""
HoardCore - Research toolkit for AI agents: retrieval & deep research.
Ingests HTML, PDF, DOCX, EPUB, and TXT into a persistent, searchable SQLite Vault.
Hybrid retrieval fuses FTS5 keyword search with vector search (RRF), a
web-discovery action feeds the crawler from a live search query, and an
entry-point plugin system lets third-party parsers/fetchers/providers/chunkers
drop in without editing the module.

Usage:
    python hoardcore.py <URL> --action scrape|crawl|search --query "text"
    python hoardcore.py _ --action ingest --urls "u1,u2,u3"
    python hoardcore.py _ --action discover --query "negros occidental renewable energy" --limit 5
    python hoardcore.py _ --action search --query "solar" --mode fast   # FTS-only
    python hoardcore.py _ --action search --query "solar" --mode hybrid # force vector+RRF
    python hoardcore.py _ --action verify --claim "..." --hint          # nearest-vault coaching
    python hoardcore.py _ --action check --migrate  # rebuild vault at 16 KB pages
    python hoardcore.py _ --action stats            # vault summary
"""
from __future__ import annotations

__version__ = "0.14.3"

import argparse
import asyncio
import hashlib
import importlib.metadata
import inspect
import io
import ipaddress
import json
import logging
import os
import queue
import re
import socket
import sqlite3
import sys
import threading
import time
import tomllib
import unicodedata
import zipfile
from array import array
from collections.abc import Iterator
from contextlib import contextmanager, suppress
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from itertools import combinations
from typing import Any
from urllib.parse import unquote, urlparse

# --- GUARANTEED DEPENDENCIES (Installed via Makefile) ---
import aiohttp
import trafilatura
from aiohttp import ClientTimeout, TCPConnector
from readability import Document
from yarl import URL

# Heavy binary parsers are lazy-imported so that HTML/text-only usage works
# even without PDF/DOCX/EPUB libraries installed. Modules are fetched on first
# use in DocumentParser (see _import_binary_parsers) and cached.
FITZ_AVAILABLE = False
DOCX_AVAILABLE = False
EPUB_AVAILABLE = False
RAPIDOCR_AVAILABLE = False
_BINARY_IMPORTED = False

# numpy is optional but strongly recommended: it makes the brute-force cosine
# scan 50-100x faster than the pure-Python fallback (B1). fastembed pulls it in
# transitively, but sparse-mode-only installs can live without it.
try:
    import numpy as _np
    NP_AVAILABLE = True
except ImportError:
    _np = None
    NP_AVAILABLE = False

# curl_cffi is optional but highly recommended (installed via Makefile)
try:
    from curl_cffi import requests as curl_requests
    CURL_AVAILABLE = True
except ImportError:
    curl_requests = None  # type: ignore[assignment]
    CURL_AVAILABLE = False
    print("Warning: curl_cffi not found. Advanced TLS impersonation disabled.", file=sys.stderr)

# --- Logging ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("hoardcore")

# =============================================================================
# 1. DATA MODELS
# =============================================================================

@dataclass
class Chunk:
    """A semantic chunk of text ready for LLM injection."""
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {"text": self.text, "metadata": self.metadata}

# =============================================================================
# 2. CONFIGURATION MANAGER
# =============================================================================

DEFAULT_CONFIG = f"""
# HoardCore v{__version__} Configuration

[general]
timeout_seconds = 30
max_retries = 2
user_agent = "HoardCore-Bot/5.0 (LLM Agent)"

[network]
default_strategy = "aggressive"   # fast, balanced, aggressive
enable_preflight = true
ssrf_protection = true            # refuse non-public (private/local/link-local)
                                  # and non-http(s) fetch targets, and re-check
                                  # every redirect hop. Disable only on trusted
                                  # isolated networks.

[auth]
cookie_string = ""

[solver]
enabled = false
url = "http://localhost:8191/v1"
solver_timeout = 60

[storage]
root_dir = "hoardcore_data"
artifacts_dir = "artifacts"          # Finished research deliverables (reports, syntheses, audits)
artifacts_by_day = true          # Organize deliverables into artifacts/YYYY-MM-DD/ subfolders
grounding_subdir = "grounding"   # Research EMITs a grounding context (a working
                                 # instrument, not a deliverable): it lands in
                                 # artifacts/YYYY-MM-DD/<grounding_subdir>/ so it
                                 # never pollutes the day folder of finished
                                 # syntheses/audits.
local_dir = "local_inputs"       # Read-only local folder for --action local: only
                                 # this directory (relative to the project root)
                                 # may be ingested from disk; paths that resolve
                                 # outside it are refused.
save_binary = true               # Save original PDF/DOCX/EPUB files
save_raw_html = false            # Save raw HTML for debugging
page_size = 16384                # SQLite page size (bytes). 16 KB keeps 384-dim
                                 # vectors inline (no overflow pages): ~1.7x faster
                                 # vector lookups than the 4 KB default. Applies
                                 # to new vaults; migrate existing ones with
                                 # `--action check --migrate`.

[parsers]
enable_pdf = true
enable_docx = true
enable_epub = true
extract_pdf_tables = true
enable_pdf_ocr = true            # auto-OCR scanned/image-only PDF pages (needs rapidocr_onnxruntime)

[crawler]
respect_robots = true
sitemap_limit = 500
parallel_workers = 5

[indexer]
enable_fts = true
search_limit = 20
parallel = true           # threaded ingest: engages the reader→embed→write
                          # pipeline for batches of 8+ chunks (--no-parallel
                          # forces sequential)
near_dedup = false        # simhash near-duplicate chunk filter (off: preserves
                          # cross-source corroborating text as evidence)
near_dedup_threshold = 3  # hamming-distance cutoff for a near-dup block (0-64)

[embeddings]
enabled = true
mode = "dense"           # dense = ONNX sentence-transformer (default); sparse = lightweight hash fallback
dense_model = "BAAI/bge-small-en-v1.5"
                         # English (384-d) by default. Multilingual-alternatives:
                         #   intfloat/multilingual-e5-small  (384-d — same dim as
                         #     the default, so vectors do NOT need re-embedding)
                         #   BAAI/bge-small-zh-v1.5     (512-d — needs backfill re-embed)
                         # The model name participates in the embedding cache
                         # fingerprint, so switching models can never serve
                         # stale cross-model vectors.
dim = 256                # used in sparse mode; dense uses the model's dimension
mrl_dims = 0             # Matryoshka truncation: store only the first N dims of
                         # dense vectors (0 = keep the full model dim). Shrinks
                         # the vector table ~4x at 384->96; best on MRL-trained
                         # models. Existing rows rebuild via backfill.
hybrid_search = true       # merge FTS + vector via RRF
top_k = 40                 # candidate pool from vector search
batch_size = 16            # chunks per model forward pass at ingest (bit-identical
                           # output to per-chunk; 0 = no batching, one chunk at a time)
quantize = "float32"       # "float32" (default) or "int8" (1 byte/dim, ~4x smaller, tiny recall cost)
fts_fast_path = true       # skip the vector scan when FTS5 alone fills the result set (all-term AND match)
recency_half_life_days = 0 # recency weighting in RRF: 0 = disabled; e.g. 30 halves an old hit's score per month
reranker_model = ""        # optional cross-encoder re-ranker applied to the final
                           # recalled set: e.g. "BAAI/bge-reranker-base" (MIT,
                           # English) or "jinaai/jina-reranker-v2-base-multilingual"
                           # (CC-BY-NC, ~1.1 GB). Empty = disabled (default).
                           # Loads lazily via fastembed on first search.

[research]
answer_first = true      # skip live DISCOVER when the existing vault already
                         # returns a high-confidence hit (Adaptive-RAG style:
                         # most recurring questions need no new retrieval)
filter_low = true        # at EMIT, drop duplicate confidence='low' chunks but
                         # keep one low chunk per distinct source (all-low: keep all)
max_per_source = 2       # cap recall chunks per source URL so one rich page
                         # can't crowd out every other source (0 = unlimited)

[discovery]
enabled = true
provider = "duckduckgo_html"   # free HTML endpoint; uses the existing fetch/FlareSolverr chain (Mojeek auto-fallback)
max_results = 10
top_rank = 6                   # ingest only the top-N ranked results
max_retries = 2                # per-provider transient-failure retries
backoff_seconds = 1.5          # exponential backoff base

[chunking]
max_tokens = 512
overlap_tokens = 50                  # sliding-window overlap between chunks: N tokens
                                     # of the previous chunk reopen the next one so
                                     # boundary-split sentences stay in context (CJK-aware)
strategy = "heading"                 # heading, paragraph, or plugin.<name> for a
                                     # plugin chunker (throws back to the build-in
                                     # pipeline on any plugin failure)

[plugins]
enabled = true                       # discover entry-point plugins (hoardcore.parsers,
                                     # .fetchers, .providers, .chunkers). Safe to leave
                                     # on with zero plugins installed.

[cache]
ttl_seconds = 86400              # 24 hours
"""

class ConfigManager:
    CONFIG_PATH = "hoardcore.toml"
    _instance = None

    def __new__(cls, config_path: str | None = None):
        # Singleton ONLY on the default path. A per-instance config_path must
        # build a fresh, independent instance so callers can point at a
        # non-default file without shared-state contamination (D2).
        if config_path is None:
            if cls._instance is None:
                cls._instance = super().__new__(cls)
            return cls._instance
        return super().__new__(cls)

    def __init__(self, config_path: str | None = None):
        if not hasattr(self, '_initialized'):
            self._config: dict[str, Any] = {}
            self.config_path = config_path or self.CONFIG_PATH
            self._load()
            self._initialized = True

    def _load(self) -> None:
        if not os.path.exists(self.config_path):
            logger.info(f"Creating default config: {self.config_path}")
            with open(self.config_path, 'w', encoding='utf-8') as f:
                f.write(DEFAULT_CONFIG.strip())

        try:
            with open(self.config_path, 'rb') as f:
                self._config = tomllib.load(f)
        except Exception as e:
            logger.warning(f"Config parse error: {e}. Using defaults.")
            self._config = self._defaults()

        defaults = self._defaults()
        for key, value in defaults.items():
            if key not in self._config:
                self._config[key] = value
            elif isinstance(value, dict):
                for sub_key, sub_value in value.items():
                    if sub_key not in self._config[key]:
                        self._config[key][sub_key] = sub_value

    def _defaults(self) -> dict[str, Any]:
        return {
            "general": {"timeout_seconds": 30, "max_retries": 2, "user_agent": "HoardCore/5.0"},
            "network": {"default_strategy": "aggressive", "enable_preflight": True, "ssrf_protection": True},
            "auth": {"cookie_string": ""},
            "solver": {"enabled": False, "url": "http://localhost:8191/v1", "solver_timeout": 60},
            "storage": {"root_dir": "hoardcore_data", "artifacts_dir": "artifacts", "artifacts_by_day": True, "grounding_subdir": "grounding", "local_dir": "local_inputs", "save_binary": True, "save_raw_html": False, "page_size": 16384},
            "parsers": {"enable_pdf": True, "enable_docx": True, "enable_epub": True, "extract_pdf_tables": True, "enable_pdf_ocr": True},
            "crawler": {"respect_robots": True, "sitemap_limit": 500, "parallel_workers": 5},
            "indexer": {"enable_fts": True, "search_limit": 20, "parallel": True,
                        "near_dedup": False, "near_dedup_threshold": 3},
            "embeddings": {"enabled": True, "mode": "dense", "dense_model": "BAAI/bge-small-en-v1.5", "dim": 256, "mrl_dims": 0, "hybrid_search": True, "top_k": 40, "conf_mode": "relative", "conf_high_abs": 0.025, "conf_low_abs": 0.020, "quantize": "float32", "fts_fast_path": True, "recency_half_life_days": 0, "reranker_model": "", "batch_size": 16},
            "research": {"answer_first": True, "filter_low": True, "max_per_source": 2},
            "discovery": {"enabled": True, "provider": "duckduckgo_html", "max_results": 10, "top_rank": 6, "max_retries": 2, "backoff_seconds": 1.5},
            "chunking": {"max_tokens": 512, "overlap_tokens": 50, "strategy": "heading"},
            "plugins": {"enabled": True},
            "cache": {"ttl_seconds": 86400}
        }

    def get(self, path: str, default: Any = None) -> Any:
        keys = path.split('.')
        value = self._config
        for key in keys:
            if isinstance(value, dict) and key in value:
                value = value[key]
            else:
                return default
        return value

# =============================================================================
# 2.5 EMBEDDING ENGINE (Vector basis for hybrid retrieval)
# =============================================================================

def _fnv1a(data: bytes) -> int:
    """FNV-1a 64-bit hash. Deterministic feature hashing for the sparse fallback."""
    h = 0xcbf29ce484222325
    for b in data:
        h = ((h ^ b) * 0x100000001b3) & 0xFFFFFFFFFFFFFFFF
    return h


def hamming64(a: int, b: int) -> int:
    """Popcount of the XOR of two 64-bit hashes (simhash near-dup distance)."""
    return bin(a ^ b).count("1")


# Near-duplicate index bucketing. To avoid loading the whole chunks_simhash
# table for every ingest (O(N) memory + candidates), each stored simhash also
# records a `bucket` = low 16 bits. Because a hamming distance of <= D changes
# at most D of those 16 bits, every simhash within distance D of a candidate
# has a bucket that is the candidate's bucket with <= D bits flipped — so
# probing only those (sum C(16,i)) bucket values finds all near-duplicates.
_SIMHASH_BUCKET_BITS = 16
_SIMHASH_BUCKET_MASK = (1 << _SIMHASH_BUCKET_BITS) - 1
_SIMHASH_PATTERN_CACHE: dict[int, list[int]] = {}
# SQLite's default max bound variables is ~999; keep a headroom margin so
# bucket probing never raises "too many SQL variables" on large thresholds.
_MAX_SQL_VARIABLES = 900


def is_ad_tracking_url(url: str) -> bool:
    """True if a URL is an ad-redirect or tracking beacon, not content.

    Search engines occasionally surface ad-redirect links (e.g. DuckDuckGo's
    /y.js tracker, Bing's /aclick redirector) as if they were organic results.
    Ingesting one stores an ad landing page as a "source", polluting recall.
    Recognized by host, path, or query markers and dropped before fetch/index.
    """
    if not url or not url.startswith("http"):
        return False
    parts = urlparse(url)
    host = (parts.netloc or "").lower()
    path = (parts.path or "").lower()
    query = (parts.query or "").lower()
    if any(s in host for s in (
            "doubleclick.net", "googlesyndication.com", "googleadservices.com",
            "amazon-adsystem.com", "adsrvr.org", "adservice.google")):
        return True
    # DuckDuckGo ad tracker / Bing ad-click redirector.
    if host == "duckduckgo.com" and path.startswith("/y.js"):
        return True
    if host.endswith("bing.com") and path.startswith("/aclick"):
        return True
    # Explicit ad metadata carried in the query string.
    return any(k in query for k in ("ad_domain=", "ad_provider=", "ad_type=",
                                    "ad_url=", "ad_clickid=", "bct=ad"))


def normalize_claim(text: str) -> str:
    """Normalize a claim for verbatim matching: fold *typographic* noise only.
    It deliberately does NOT add/remove tokens: "400K" stays distinct from
    "400K+" and reordered words never match. Typography-blind, semantics-strict.
    """
    if not text:
        return ""
    # Strip parser-emitted markdown emphasis/code markers (**bold**, *italic*,
    # `code`) before folding: they are render artifacts, not wording, so a
    # sentence stored as "increased by **17% in 2025**" verifies against
    # "increased by 17% in 2025" (same principle as folding typographic dashes).
    # Applied symmetrically to claim and stored text; never adds/removes tokens.
    t = re.sub(r"[\*`]+", "", text)
    t = t.translate({
        0x2010: "-", 0x2011: "-", 0x2012: "-", 0x2013: "-",
        0x2014: "-", 0x2015: "-", 0x2212: "-",
        0x2018: "'", 0x2019: "'", 0x201A: "'", 0x201B: "'",
        0x201C: '"', 0x201D: '"', 0x201E: '"', 0x201F: '"',
    })
    # NFKC folds full-width variants (FVPLUS "+"/FVHYPHEN) onto ASCII, and
    # unifies look-alike punctuation; do it on the translated string.
    norm = unicodedata.normalize("NFKC", t)
    return re.sub(r"\s+", " ", norm).strip().lower()


def _tidy_markdown_text(text: str) -> str:
    """Strip parser-emitted markdown emphasis/code markers from stored chunk
    text so the vault's canonical text is clean prose (mirrors the verifier's
    marker-strip in `normalize_claim`, applied at ingest so the stored form
    already matches the compared form). Preserves newline structure and leaves
    code fences (```` ``` ````) and their content intact. Collapses only the
    horizontal whitespace a removed marker orphans, so "a **b** c" -> "a b c"
    but line structure is untouched."""
    if not text:
        return text
    in_fence = False
    out: list[str] = []
    for line in text.split("\n"):
        if line.lstrip().startswith("```"):
            in_fence = not in_fence
            out.append(line)          # keep the fence marker itself
            continue
        if in_fence:
            out.append(line)          # code content left untouched
            continue
        t = re.sub(r"[\*`]+", "", line)   # strip emphasis + inline code markers
        out.append(re.sub(r"[ \t]{2,}", " ", t))
    return "\n".join(out)


def _nearest_phrase_probe(text: str, needle: str) -> tuple[float, int, int]:
    """Best fuzzy overlap of `needle` in `text`; returns (ratio, start, size)."""
    hay = normalize_claim(text)
    if not hay:
        return (0.0, 0, 0)
    m = SequenceMatcher(None, needle, hay).find_longest_match(
        0, len(needle), 0, len(hay))
    ratio = (2.0 * m.size) / (len(needle) + len(hay)) if (len(needle) + len(hay)) else 0.0
    return (ratio, m.b, m.size)


def _fts_token(token: str) -> str:
    """Normalize a single query token for FTS MATCH, aligned with the index.

    The vault uses the `porter unicode61` tokenizer, which treats currency
    symbols as separators: `$13` is indexed as the token `13`, not `$13`. To
    make a keyword/OR-fallback MATCH agree with the stored index, a `$` directly
    followed by digits (e.g. `$13`, `$21.3`, `$1`) is reduced to its digit-only
    form for the FTS phrase — while `verify`'s raw-text `LIKE` still confirms
    the verbatim `$13` for `[V]` (see `normalize_claim`).
    """
    if "$" in token:
        stripped = token.replace("$", "")
        if stripped and stripped[0].isdigit():
            return stripped
    return token


def _token_matches(cursor: sqlite3.Cursor, token: str) -> bool:
    """True if a single quoted FTS token matches at least one stored chunk.

    Used by the hybrid OR-fallback guard: it counts how many distinct query
    tokens are genuinely present in the corpus, so a long research question is
    rescued by the OR fallback only when it is actually topical (>=2 matching
    tokens), never by a single coincidental word. `$`+number tokens are aligned
    to the index via `_fts_token`.
    """
    q = f'"{_fts_token(token)}"'
    try:
        cursor.execute("SELECT 1 FROM chunks_fts WHERE chunks_fts MATCH ? LIMIT 1", (q,))
        return cursor.fetchone() is not None
    except sqlite3.OperationalError:
        return False


def _simhash_bucket_patterns(k: int) -> list[int]:
    """All 16-bit masks whose popcount is <= min(k, 16), for hamming probing.

    Probe buckets for a candidate are `candidate_bucket ^ p` over these
    patterns; this set provably contains the bucket of any stored simhash
    within hamming distance k. For thresholds high enough that the pattern set
    would blow past SQLite's bound-variable limit, callers must fall back to a
    full-table scan (see _near_duplicate_candidates).
    """
    k = min(int(k), _SIMHASH_BUCKET_BITS)
    if k in _SIMHASH_PATTERN_CACHE:
        return _SIMHASH_PATTERN_CACHE[k]
    pats: list[int] = []
    for i in range(k + 1):
        for bits in combinations(range(_SIMHASH_BUCKET_BITS), i):
            m = 0
            for b in bits:
                m |= 1 << b
            pats.append(m)
    _SIMHASH_PATTERN_CACHE[k] = pats
    return pats


class EmbeddingsEngine:
    """Turns chunk text into fixed-dimension vectors for hybrid retrieval.

    Default mode is dense: an ONNX-quantized sentence-transformer embedding
    (via fastembed on onnxruntime, no PyTorch) that captures semantic
    similarity. A lightweight sparse fallback (FNV-1a feature hashing of word
    + char n-gram features into a unit vector) remains for environments where
    the dense model is unavailable. Both fuse with FTS5 keyword search via
    Reciprocal Rank Fusion, keeping HoardCore fast and single-file.
    """

    def __init__(self, config: ConfigManager):
        self.config = config
        self.enabled = config.get('embeddings.enabled', True)
        # Mode: "dense" (default, ONNX-quantized sentence-transformer via
        # fastembed) or "sparse" (dependency-light FNV-1a lexical hash). Dense
        # is lazily loaded and falls back to sparse if fastembed is missing.
        self.mode = str(config.get('embeddings.mode', 'dense')).lower()
        self.dim = int(config.get('embeddings.dim', 256))
        # Matryoshka-style dimension truncation (MRL): embed at full model
        # dim, then STORE only the first `mrl_dims` dimensions (models trained
        # for MRL like nomic/mxbai retain most quality when truncated; e.g.
        # 93% retention at 12x compression on OpenAI's text-embedding-3-large).
        # Queries are truncated the same way so cosine stays consistent; the
        # existing backfill dimension-migration rebuilds any stale rows.
        self.mrl_dims = int(config.get('embeddings.mrl_dims', 0) or 0)
        self.base_dim = self.dim
        # Vector storage format: "float32" (4 bytes/dim) or "int8" (1 byte/dim,
        # ~4x smaller vault, tiny recall cost). int8 is applied to dense output
        # only — the sparse hash already produces compact normalized vectors.
        self.quantize = str(config.get('embeddings.quantize', 'float32')).lower()
        if self.mode == 'dense' and self.quantize not in ('float32', 'int8'):
            logger.warning(
                f"Unrecognized quantize={self.quantize!r}; falling back to float32."
            )
            self.quantize = 'float32'
        self._dense = None  # lazy fastembed backend (model + dim)
        if self.mode == 'dense':
            try:
                self.dim = self._load_dense()
                self.base_dim = self.dim
            except Exception as e:  # fastembed missing or model download failed
                logger.warning(
                    f"Dense mode requested but unavailable ({e}); "
                    f"falling back to sparse lexical hashing."
                )
                self.mode = 'sparse'
        if self.mode == 'dense' and self.mrl_dims > 0 and self.mrl_dims < self.base_dim:
            self.dim = self.mrl_dims
            logger.info(
                f"MRL dim truncation active: storing first {self.dim} of "
                f"{self.base_dim} embedding dims."
            )

    @property
    def bytes_per_dim(self) -> int:
        """On-disk bytes per embedding dimension (4 for float32, 1 for int8)."""
        return 1 if (self.mode == 'dense' and self.quantize == 'int8') else 4

    def fingerprint(self) -> str:
        """Cache-safe fingerprint of the embedding configuration.

        The content-addressable vector cache (chunk_vectors_ca) is keyed by
        content hash; without this fingerprint a chunk embedded under an old
        model / dimension / mrl / quantize could be served to the current
        configuration. The fingerprint captures every knob that changes the
        stored bytes (mode, stored dim, mrl truncation, quantize, and the
        dense model name), so a cache hit is only ever served when it was
        built with today's embedding settings (A9).
        """
        model = str(self.config.get('embeddings.dense_model', 'BAAI/bge-small-en-v1.5'))
        parts = "|".join((self.mode, str(self.dim), self.quantize,
                          str(self.mrl_dims), model))
        return hashlib.blake2b(parts.encode("utf-8"), digest_size=8).hexdigest()

    def _load_dense(self) -> int:
        """Lazily import fastembed and load the ONNX-quantized model.

        Returns the model's embedding dimension. Raises if fastembed is not
        installed or the model cannot be loaded, so the caller can fall back
        to sparse.
        """
        from fastembed import TextEmbedding

        model_name = str(self.config.get('embeddings.dense_model',
                                         'BAAI/bge-small-en-v1.5'))
        model = TextEmbedding(model_name)
        # Probe a short real token to discover the embedding dimension (an empty
        # string can produce a degenerate vector on some tokenizers).
        probe = next(iter(model.embed(['probe'])), None)
        dim = len(probe) if probe is not None else 384
        if dim <= 0:
            # Fall back to the documented MiniLM dimension if probe is empty.
            dim = 384
        self._dense = (model, dim)
        logger.info(f"Dense embeddings loaded: {model_name} (dim={dim})")
        return dim

    @staticmethod
    def _tokens(text: str) -> list[str]:
        """Word unigrams + 3-gram shingles for sparse lexical features."""
        words = re.findall(r"[a-zà-ÿ0-9']+", text.lower())
        shingles: list[str] = []
        for w in words:
            if len(w) <= 3:
                shingles.append(w)
                continue
            for i in range(len(w) - 2):
                shingles.append(w[i:i + 3])
        return words + shingles

    def _hash_vector(self, text: str) -> bytes:
        """Feature-hash tokens into an L2-normalized float32 vector."""
        vec = [0.0] * self.dim
        counts: dict[str, int] = {}
        for t in self._tokens(text):
            counts[t] = counts.get(t, 0) + 1
        for token, n in counts.items():
            digest = _fnv1a(token.encode('utf-8'))
            idx = digest % self.dim
            sign = 1.0 if (digest >> 62) & 1 else -1.0
            vec[idx] += sign * (1.0 + n ** 0.5)
        norm = (sum(v * v for v in vec)) ** 0.5
        if norm > 0:
            vec = [v / norm for v in vec]
        arr = array('f', vec)
        return arr.tobytes()

    def vectorize(self, text: str) -> bytes:
        was_dense = (self.mode == 'dense' and self._dense is not None)
        raw = self._vectorize_dense(text) if was_dense else self._hash_vector(text)
        return self._finalize(raw, was_dense)

    def vectorize_batch(self, texts: list[str]) -> list[bytes]:
        """Embed many texts in a single model call; returns a list aligned to
        `texts`.

        Bit-identical to calling `vectorize` per text (shared normalization and
        post-processing), so batching is a pure throughput win with no recall
        cost. Falls back to per-item embedding if the model returns a mismatched
        count or raises.
        """
        if not texts:
            return []
        if self.mode == 'dense' and self._dense is not None:
            model, _dim = self._dense
            vectors: list = []
            try:
                vectors = list(model.embed(list(texts)))
            except Exception as e:
                logger.warning(f"Batch embed failed ({e}); per-item fallback.")
            if len(vectors) == len(texts):
                return [
                    self._finalize(self._dense_vec_to_bytes(v) if v is not None else b"", True)
                    for v in vectors
                ]
            logger.warning("Batch embed returned a different count; per-item fallback.")
        return [self.vectorize(t) for t in texts]

    def _dense_vec_to_bytes(self, vec) -> bytes:
        """Normalize a raw dense vector to a unit float32 blob (shared by the
        single- and batch-embedding paths so tokenization/output is identical)."""
        from array import array

        if _np is not None:
            arr = _np.asarray(vec, dtype=_np.float32).reshape(-1)
            norm = float(_np.linalg.norm(arr))
            if norm > 0:
                arr = arr / norm
            return array('f', arr).tobytes()
        arr = list(vec)
        norm = float(sum(v * v for v in arr)) ** 0.5
        if norm > 0:
            arr = [v / norm for v in arr]
        return array('f', arr).tobytes()

    def _finalize(self, vec: bytes, was_dense: bool) -> bytes:
        """Apply mode-specific post-processing (MRL truncation, int8) to a
        serialized vector. Shared by single and batched embedding."""
        if was_dense and self.mrl_dims > 0:
            vec = self._truncate_f32(vec, self.dim)
        if was_dense and self.quantize == 'int8' and vec:
            return self._quantize_int8(vec)
        return vec

    @staticmethod
    def _truncate_f32(vec: bytes, dims: int) -> bytes:
        """Slice a float32 vector blob to its first `dims` dimensions.

        Matryoshka-style truncation for models trained for progressive dim
        reduction (MRL). Query and stored vectors are truncated identically so
        the cosine dimension check stays consistent. Empty/invalid payloads
        pass through unchanged.
        """
        if not vec:
            return vec
        arr = array('f')
        arr.frombytes(vec)
        if len(arr) <= dims:
            return vec
        return array('f', arr[:dims]).tobytes()

    @staticmethod
    def _quantize_int8(vec: bytes) -> bytes:
        """Convert a float32 vector to signed int8 (scale to [-127, 127]).

        Dequantization for cosine is just dividing back by 127, so vectors stay
        comparable without a per-vector scale factor.
        """
        arr = array('f')
        arr.frombytes(vec)
        q = array('b', (max(-127, min(127, round(v * 127))) for v in arr))
        return q.tobytes()

    def _vectorize_dense(self, text: str) -> bytes:
        """Encode one text with the loaded dense model, normalized to a unit
        float32 blob. Post-processing (MRL truncation, int8) is applied by
        `_finalize` so the single and batched paths stay byte-identical."""
        dense = self._dense
        if dense is None:
            return b""
        model, _dim = dense
        vec = next(iter(model.embed([text])), None)
        if vec is None:
            return b""
        return self._dense_vec_to_bytes(vec)

    @staticmethod
    def cosine(a: bytes, b: bytes, dim: int) -> float:
        """Cosine similarity between two serialized vectors.

        Vectors are stored L2-normalized at build time, so a plain dot product
        equals cosine. int8-quantized vectors are exactly `dim` bytes; float32
        are `dim * 4`. Any other payload length signals corruption or a stale
        embedding format, and is surfaced instead of silently truncated (A7),
        so mismatches are never hidden by a plausible-but-wrong score.
        """
        int8_desc = (dim, dim)
        f32_desc = (dim * 4, dim * 4)
        if (len(a), len(b)) != int8_desc and (len(a), len(b)) != f32_desc:
            logger.warning(
                "vector dim mismatch: query=%d bytes, stored=%d bytes, "
                "expected int8=%d or float32=%d; scoring as 0.0",
                len(a), len(b), dim, dim * 4,
            )
            return 0.0
        if len(a) == dim:  # int8 path
            if _np is not None:
                # Must upcast to int32, NOT int16/int8: the worst-case element
                # product is 127*127 = 16129, and int16 overflows at 32767 in
                # just 2-3 dimensions; int8 would silently wrap. int8's value is
                # 4x smaller on-disk storage, not scan speed — the float32 path
                # below is the fast one (no copy/upcast).
                va = _np.frombuffer(a, dtype=_np.int8).astype(_np.int32)
                vb = _np.frombuffer(b, dtype=_np.int8).astype(_np.int32)
                return float(_np.dot(va, vb)) / (127.0 ** 2)
            va = array('b')
            va.frombytes(a)
            vb = array('b')
            vb.frombytes(b)
            # Both confirmed to be exactly `dim` bytes above, so strict is safe.
            return float(sum(x * y for x, y in zip(va, vb, strict=True))) / (127.0 ** 2)
        # float32 path (all dimensions are identical by the check above)
        if _np is not None:
            va = _np.frombuffer(a, dtype=_np.float32)
            vb = _np.frombuffer(b, dtype=_np.float32)
            return float(_np.dot(va, vb))
        va = array('f')
        va.frombytes(a)
        vb = array('f')
        vb.frombytes(b)
        # Both confirmed to be exactly `dim * 4` bytes above, so strict is safe.
        return float(sum(x * y for x, y in zip(va, vb, strict=True)))

# =============================================================================
# 3. PERSISTENT STORAGE & VAULT (SQLite + Filesystem)
# =============================================================================

# Ingest pipeline tuning (kept modest to stay lightweight)
DB_TIMEOUT           = 30.0
CONNECTION_POOL_SIZE = int(os.environ.get("HC_POOL_SIZE", "8"))
CHUNK_BATCH_SIZE     = 50
PIPELINE_QUEUE_SIZE  = 20
WORKER_THREADS       = int(os.environ.get("HC_WORKERS", "4"))

# Schema versioning (G3): PRAGMA user_version is the source of truth.
# 0 = legacy unversioned vault (pre-0.9.0 inline migrations still run), which
# gets stamped to _SCHEMA_VERSION on open. Future schema changes bump the
# constant and append a (description, DDL) entry to _SCHEMA_MIGRATIONS — never
# edit a shipped migration, append the next one.
_SCHEMA_VERSION = 1
_SCHEMA_MIGRATIONS: dict[int, tuple[str, str]] = {}

class ConnectionPool:
    """Reusable bounded pool of SQLite connections.

    Each connection is configured for WAL mode, a memory-mapped I/O window, an
    in-memory temp store, and a page cache for better concurrent read/write
    throughput than the previous open-a-new-connection-per-query behaviour.
    """

    def __init__(self, db_path: str, pool_size: int = CONNECTION_POOL_SIZE,
                 page_size: int = 16384):
        self.db_path = db_path
        self.pool_size = pool_size
        self.page_size = page_size
        self._pool: queue.Queue = queue.Queue(maxsize=pool_size)
        for _ in range(pool_size):
            self._pool.put(self._create_connection())

    def _create_connection(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.db_path, timeout=DB_TIMEOUT,
                               check_same_thread=False)
        conn.execute(f"PRAGMA page_size = {self.page_size}")
        conn.execute("PRAGMA journal_mode = WAL")
        conn.execute("PRAGMA mmap_size = 536870912")    # 512 MB mmap window
        conn.execute("PRAGMA temp_store = MEMORY")
        conn.execute("PRAGMA synchronous = NORMAL")
        conn.execute("PRAGMA cache_size = -64000")      # 64 MB page cache
        conn.execute("PRAGMA busy_timeout = 5000")
        return conn

    def get(self, timeout: float = 30.0) -> sqlite3.Connection:
        """Acquire a connection, creating an overflow one if the pool is busy."""
        try:
            conn = self._pool.get(timeout=timeout)
            try:
                conn.execute("SELECT 1")
                return conn
            except sqlite3.Error:
                # Discard the dead connection so its fd/WAL handles aren't
                # leaked, and hand out a fresh one in its place (S5).
                with suppress(Exception):
                    conn.close()
                return self._create_connection()
        except queue.Empty:
            return self._create_connection()

    def put(self, conn: sqlite3.Connection) -> None:
        """Return a connection to the pool, or close it if the pool is full."""
        try:
            self._pool.put_nowait(conn)
        except queue.Full:
            with suppress(Exception):
                conn.close()

    def close_all(self) -> None:
        """Close every connection currently sitting in the pool."""
        while not self._pool.empty():
            with suppress(Exception):
                self._pool.get_nowait().close()

class VaultManager:
    """Handles SQLite FTS indexing and filesystem storage."""

    def __init__(self, config: ConfigManager, vault_name: str | None = None,
                 event_bus: EventBus | None = None):
        self.config = config
        self.vault_name = vault_name
        self.bus = event_bus
        root_dir = config.get('storage.root_dir', 'hoardcore_data')
        if vault_name:
            # Guard against path traversal: a vault name must be a single
            # path-safe token (no separators, '..', or leading dots).
            safe = re.sub(r'[^A-Za-z0-9._-]+', '-', vault_name).strip('.-')
            if not safe or safe != vault_name:
                logger.warning(
                    f"Sanitized vault name {vault_name!r} -> {safe!r}"
                )
                vault_name = safe
            self.vault_name = vault_name
            root_dir = os.path.join(root_dir, vault_name)
        self.root_dir = root_dir
        os.makedirs(self.root_dir, exist_ok=True)
        self.artifacts_dir = config.get('storage.artifacts_dir', 'artifacts')
        os.makedirs(self.artifacts_dir, exist_ok=True)
        self.db_path = os.path.join(self.root_dir, 'vault.db')
        self.embeddings = EmbeddingsEngine(config)
        self._vector_dim = self.embeddings.dim
        self.page_size = int(config.get('storage.page_size', 16384))
        self._pool = ConnectionPool(self.db_path, CONNECTION_POOL_SIZE, self.page_size)
        self._init_db()
        self.backfill_vectors()

        # Lazy-loaded cross-encoder reranker (embeddings.reranker_model).
        self._reranker = None
        # Brute-force vector-scan cache (numpy matmul): keyed on the vault's
        # vector count so a rebuild happens exactly when a new row lands.
        self._vec_mat_cache: dict[str, Any] = {"count": None}

    @contextmanager
    def _db(self) -> Iterator[tuple[sqlite3.Connection, sqlite3.Cursor]]:
        """Yield a committed-on-success, surfaced-on-exception DB cursor.

        Uses the connection pool (WAL + mmap + page cache). Guarantees the
        connection is always returned to the pool and transactions are never
        left dangling, even if a query raises mid-method.
        """
        conn = self._pool.get()
        try:
            cursor = conn.cursor()
            yield conn, cursor
            conn.commit()
        except Exception:
            conn.rollback()
            raise
        finally:
            self._pool.put(conn)

    def _init_db(self) -> None:
        """Initialize SQLite with FTS5 virtual table."""
        with self._db() as (_conn, cursor):
            # Enable FTS5
            cursor.execute("PRAGMA journal_mode=WAL;")
            cursor.execute("PRAGMA synchronous=NORMAL;")

            # --- v0.6.0 schema migration ----------------------------------
            # Older vaults (pre-0.6.0) had documents without `version` /
            # `content_hash` and a UNIQUE(url) constraint. Rebuild the table in
            # place, preserving existing rows (each becomes version 1).
            cursor.execute(
                "SELECT sql FROM sqlite_master WHERE type='table' AND name='documents'"
            )
            row = cursor.fetchone()
            if row and "version" not in (row[0] or ""):
                logger.info("Migrating documents table to v0.6.0 schema (WORM versions).")
                cursor.execute("ALTER TABLE documents RENAME TO documents_old")
                cursor.execute("""
                    CREATE TABLE documents (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        url TEXT,
                        domain TEXT,
                        file_name TEXT,
                        content_type TEXT,
                        fetched_at REAL,
                        parser_used TEXT,
                        quality_score REAL,
                        total_chunks INTEGER,
                        metadata_json TEXT,
                        version INTEGER NOT NULL DEFAULT 1,
                        content_hash TEXT,
                        UNIQUE(url, version)
                    )
                """)
                cursor.execute("""
                    INSERT INTO documents (
                        id, url, domain, file_name, content_type, fetched_at,
                        parser_used, quality_score, total_chunks, metadata_json,
                        version
                    )
                    SELECT id, url, domain, file_name, content_type, fetched_at,
                           parser_used, quality_score, total_chunks, metadata_json,
                           1
                    FROM documents_old
                """)
                cursor.execute("DROP TABLE documents_old")

            # Main table for metadata
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    url TEXT,
                    domain TEXT,
                    file_name TEXT,
                    content_type TEXT,
                    fetched_at REAL,
                    parser_used TEXT,
                    quality_score REAL,
                    total_chunks INTEGER,
                    metadata_json TEXT,
                    version INTEGER NOT NULL DEFAULT 1,
                    content_hash TEXT,
                    UNIQUE(url, version)
                )
            """)

            # Content-addressable chunk index for cross-document deduplication.
            # chunk_hash is the BLAKE2b-256 of the raw chunk text. The vector
            # table is keyed by this hash so identical chunks share one vector.
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS chunks_ca (
                    chunk_hash TEXT PRIMARY KEY,
                    text TEXT NOT NULL,
                    url TEXT,
                    header_path TEXT,
                    metadata_json TEXT,
                    first_seen REAL
                )
            """)

            # FTS5 virtual table for full-text search
            cursor.execute("""
                CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(
                    url,
                    header_path,
                    text,
                    metadata_json,
                    tokenize = 'porter unicode61'
                )
            """)

            # No automatic cascade on document delete: chunks_fts rows carry no
            # version column, so a URL-scoped DELETE would nuke every version's
            # chunks on a single-row surgery (S6). Nothing in the normal flow
            # deletes documents (WORM append-only), so the old URL-scoped
            # trigger is dropped to remove that data-loss trap; users pruning
            # rows by hand own the chunks accordingly.
            cursor.execute("DROP TRIGGER IF EXISTS documents_after_delete")

            # Vector index for hybrid retrieval. chunk_rowid mirrors the implicit
            # rowid of chunks_fts rows so vector and FTS results can be fused.
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS chunk_vectors (
                    chunk_rowid INTEGER PRIMARY KEY,
                    url TEXT,
                    vector BLOB
                )
            """)
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_chunk_vec_url ON chunk_vectors(url)")

            # Recency weighting (P1.2) looks up documents by (url, fetched_at);
            # this index turns those searches from O(N) scans into index seeks.
            cursor.execute(
                "CREATE INDEX IF NOT EXISTS idx_documents_url_fetched "
                "ON documents(url, fetched_at DESC)"
            )

            # Content-addressable vector cache: identical chunks share one
            # embedding, so cross-document duplicate text is embedded once.
            # embed_fp stores the embedding-config fingerprint that built the
            # vector; a cache entry is only served when its fingerprint matches
            # the current embedding settings (model/dim/mrl/quantize), so a
            # stale cross-model vector can never be reused (A9).
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS chunk_vectors_ca (
                    chunk_hash TEXT PRIMARY KEY,
                    vector BLOB,
                    embed_fp TEXT
                )
            """)
            _ca_cols = [c[1] for c in cursor.execute(
                "PRAGMA table_info(chunk_vectors_ca)").fetchall()]
            if 'embed_fp' not in _ca_cols:
                cursor.execute("ALTER TABLE chunk_vectors_ca ADD COLUMN embed_fp TEXT")

            # Near-duplicate index (optional, indexer.near_dedup): a 64-bit
            # simhash per stored chunk. Exact content handling stays in
            # chunks_ca; this catches near-identical re-crawled text whose
            # boilerplate differs. First-write wins, like chunks_ca.
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS chunks_simhash (
                    simhash INTEGER PRIMARY KEY,
                    chunk_hash TEXT,
                    url TEXT,
                    text TEXT,
                    first_seen REAL,
                    bucket INTEGER
                )
            """)

            # Vaults created before the bucket-based near-dup probe index lack
            # the column; add and backfill it before any bucket queries run.
            cols = [c[1] for c in cursor.execute(
                "PRAGMA table_info(chunks_simhash)").fetchall()]
            if 'bucket' not in cols:
                cursor.execute("ALTER TABLE chunks_simhash ADD COLUMN bucket INTEGER")
                cursor.execute(
                    "UPDATE chunks_simhash SET bucket = simhash & ?",
                    (_SIMHASH_BUCKET_MASK,),
                )
            cursor.execute(
                "CREATE INDEX IF NOT EXISTS idx_chunks_simhash_bucket "
                "ON chunks_simhash(bucket)"
            )

            # --- Schema versioning (G3) ---
            # PRAGMA user_version is the single source of truth. 0 marks a
            # legacy (pre-0.9.0) vault: the inline migrations above already
            # brought the schema current, so stamp it to baseline. Newer
            # versions run the numbered migrations in order — each migration
            # is appended to _SCHEMA_MIGRATIONS, never edited, and the version
            # bump commits atomically with the DDL (SQLite DDL is
            # transactional).
            current_ver = cursor.execute("PRAGMA user_version").fetchone()[0]
            if current_ver == 0:
                cursor.execute(f"PRAGMA user_version = {_SCHEMA_VERSION}")
                logger.info(f"Vault stamped baseline schema version {_SCHEMA_VERSION}.")
            elif current_ver < _SCHEMA_VERSION:
                for target in range(current_ver + 1, _SCHEMA_VERSION + 1):
                    migration = _SCHEMA_MIGRATIONS.get(target)
                    if migration is None:
                        raise sqlite3.OperationalError(
                            f"Missing schema migration for version {target}"
                        )
                    description, statements = migration
                    logger.info(f"Applying schema migration {target}: {description}")
                    for statement in (s.strip() for s in statements.split(";") if s.strip()):
                        cursor.execute(statement)
                    cursor.execute(f"PRAGMA user_version = {target}")
            elif current_ver > _SCHEMA_VERSION:
                logger.error(
                    f"Vault schema version {current_ver} is newer than this build "
                    f"supports ({_SCHEMA_VERSION}); refusing to open it."
                )
                raise sqlite3.OperationalError(
                    f"vault schema v{current_ver} > supported v{_SCHEMA_VERSION}"
                )

    def _get_domain_folder(self, url: str) -> str:
        parsed = urlparse(url)
        domain = parsed.netloc or "unknown_domain"
        domain = re.sub(r'[^\w\-\.]', '_', domain)
        folder = os.path.join(self.root_dir, domain)
        os.makedirs(folder, exist_ok=True)
        os.makedirs(os.path.join(folder, 'binaries'), exist_ok=True)
        os.makedirs(os.path.join(folder, 'extracted'), exist_ok=True)
        return folder

    def _generate_filename(self, url: str, content_type: str) -> str:
        """Generate a safe filename from URL."""
        parsed = urlparse(url)
        path = parsed.path or '/'
        name = os.path.basename(path) or 'index'
        # Remove query strings
        name = re.sub(r'\?.*$', '', name)
        # Sanitize characters that are illegal in filenames on some platforms
        # (Windows `:`/`*`, path separators, etc.) instead of hoping the URL
        # author cooperated (C6).
        name = re.sub(r'[^A-Za-z0-9._-]+', '_', name).strip('.-') or 'index'
        if not name:
            name = hashlib.blake2b(url.encode(), digest_size=8).hexdigest()[:16]

        # Determine extension
        ext_map = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/epub+zip': '.epub',
            'text/html': '.html',
            'text/plain': '.txt'
        }
        ext = ext_map.get(content_type, '.bin')
        # Ensure unique by hashing. BLAKE2b-64 (16 hex chars) has a negligible
        # collision risk, vs the old 24-bit MD5 suffix which hit a 50% collision
        # probability around 4k URLs and silently overwrote files (C2).
        hash_suffix = hashlib.blake2b(url.encode(), digest_size=8).hexdigest()[:16]
        return f"{name}_{hash_suffix}{ext}"

    def save_binary(self, url: str, content_type: str, data: bytes) -> str:
        """Save binary file to domain/binaries/ folder."""
        folder = self._get_domain_folder(url)
        bin_dir = os.path.join(folder, 'binaries')
        filename = self._generate_filename(url, content_type)
        bin_path = os.path.join(bin_dir, filename)

        if not os.path.exists(bin_path):
            with open(bin_path, 'wb') as f:
                f.write(data)
            logger.info(f"Saved binary: {bin_path}")
        return bin_path

    def save_extracted_text(self, url: str, markdown: str, chunks: list[Chunk], meta: dict[str, Any]) -> None:
        """Save extracted text and chunks to domain/extracted/."""
        folder = self._get_domain_folder(url)
        ext_dir = os.path.join(folder, 'extracted')

        # Create a safe filename
        filename = self._generate_filename(url, meta.get('content_type', 'text/plain'))
        base_name = os.path.splitext(filename)[0]

        # Save full markdown
        md_path = os.path.join(ext_dir, f"{base_name}.content.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown)

        # Save chunks JSON
        chunks_path = os.path.join(ext_dir, f"{base_name}.chunks.json")
        with open(chunks_path, 'w', encoding='utf-8') as f:
            json.dump([c.to_dict() for c in chunks], f, indent=2, ensure_ascii=False)

        logger.info(f"Saved extracted text to {ext_dir}")

    @staticmethod
    def _simhash(tokens: list[str]) -> int:
        """63-bit simhash over token features (FNV-1a feature hashing).

        Standard weighted simhash: each token's 64-bit hash votes up/down on
        every bit; the sign of each bit's total vote becomes the bit value.
        The top bit is forced clear so the value always fits SQLite's *signed*
        64-bit INTEGER — otherwise ~44% of documents fail ingest with
        "Python int too large to convert to SQLite INTEGER". All bits still
        vote; comparisons only ever see masked values, so hamming distances
        between them are unaffected.
        """
        v = [0] * 64
        for t in set(tokens):
            h = _fnv1a(t.encode("utf-8"))
            for i in range(64):
                v[i] += 1 if (h >> i) & 1 else -1
        out = 0
        for i in range(63):
            if v[i] > 0:
                out |= 1 << i
        return out

    def _filter_near_dupes(self, cursor: sqlite3.Cursor, url: str,
                           chunks: list[Chunk]) -> list[Chunk]:
        """Drop chunks near-duplicate (simhash hamming <= threshold) to already
        stored text, optionally, when config indexer.near_dedup is true.

        Exact-duplicate handling stays in chunks_ca; this catches near-identical
        re-crawled pages whose boilerplate differs. Kept chunks get their
        simhash recorded (first-write wins). Off by default because collapsing
        cross-source corroborating text would hide evidence — enable when
        crawling large sites and duplicate growth matters.
        """
        if not self.config.get('indexer.near_dedup', False):
            return chunks
        threshold = int(self.config.get('indexer.near_dedup_threshold', 3))
        kept: list[Chunk] = []
        local_seen: set[int] = set()
        for chunk in chunks:
            tokens = EmbeddingsEngine._tokens(chunk.text)
            if not tokens:
                kept.append(chunk)  # no lexical features -> cannot judge
                continue
            sh = self._simhash(tokens)
            if self._is_near_duplicate(cursor, sh, threshold, local_seen):
                logger.debug(f"near-dup block (hamming<= {threshold}) for {url}")
                continue
            local_seen.add(sh)
            c_hash = hashlib.blake2b(chunk.text.encode("utf-8"),
                                     digest_size=32).hexdigest()
            cursor.execute(
                "INSERT OR IGNORE INTO chunks_simhash "
                "(simhash, chunk_hash, url, text, first_seen, bucket) "
                "VALUES (?, ?, ?, ?, ?, ?)",
                (sh, c_hash, url, chunk.text, time.time(),
                 sh & _SIMHASH_BUCKET_MASK),
            )
            kept.append(chunk)
        if len(kept) != len(chunks):
            logger.info(f"near-dedup kept {len(kept)}/{len(chunks)} chunks for {url}")
        return kept

    def _near_duplicate_candidates(self, cursor: sqlite3.Cursor,
                                   sh: int, threshold: int) -> list[int]:
        """Stored simhashes within hamming `threshold` of `sh`.

        Probes only the buckets that could contain a near-duplicate (see
        _simhash_bucket_patterns) instead of scanning the whole table, so
        ingest cost stays bounded by the near-dup cluster size rather than the
        vault's total chunk count.
        """
        base = sh & _SIMHASH_BUCKET_MASK
        probes = [base ^ p for p in _simhash_bucket_patterns(threshold)]
        if len(probes) > _MAX_SQL_VARIABLES:
            # Pattern count (sum of C(16,i)) explodes past SQLite's bound
            # variable limit at threshold >= 4-5; degrade gracefully to a
            # full-table scan, which stays correct (only slower).
            rows = cursor.execute(
                "SELECT simhash FROM chunks_simhash"
            ).fetchall()
            return [r[0] for r in rows]
        qmarks = ",".join("?" * len(probes))
        rows = cursor.execute(
            f"SELECT simhash FROM chunks_simhash WHERE bucket IN ({qmarks})",  # nosec B608
            probes).fetchall()
        return [r[0] for r in rows]

    def _is_near_duplicate(self, cursor: sqlite3.Cursor, sh: int,
                           threshold: int, local_seen: set[int]) -> bool:
        """True if `sh` is within hamming threshold of a stored or just-kept
        simhash. `local_seen` avoids re-querying the DB for chunks kept earlier
        in the same document.
        """
        for e in local_seen:
            if hamming64(sh, e) <= threshold:
                return True
        for e in self._near_duplicate_candidates(cursor, sh, threshold):
            if e in local_seen:
                continue
            if hamming64(sh, e) <= threshold:
                return True
        return False

    def index_document(self, url: str, chunks: list[Chunk], meta: dict[str, Any]) -> None:
        """Insert/update document and chunks in SQLite FTS.

        WORM semantics: re-ingesting the same URL creates a new *version* row
        rather than overwriting the previous one, so the vault is append-only.
        Chunks are content-addressed (BLAKE2b-256); identical chunk text across
        documents shares a single canonical entry and is embedded only once.
        """
        if not self.config.get('indexer.enable_fts', True):
            return

        embed_ok = self.config.get('embeddings.enabled', True)
        domain = urlparse(url).netloc

        with self._db() as (_conn, cursor):
            # Optional near-duplicate filter (indexer.near_dedup) runs inside
            # the same transaction so skipped chunks never half-persist.
            chunks = self._filter_near_dupes(cursor, url, chunks)
            content_hash = hashlib.blake2b(
                "\n".join(c.text for c in chunks).encode("utf-8"),
                digest_size=32,
            ).hexdigest()

            # WORM: determine the next version for this URL (never overwrite).
            cursor.execute(
                "SELECT COALESCE(MAX(version), 0) FROM documents WHERE url = ?",
                (url,),
            )
            version = cursor.fetchone()[0] + 1

            # Insert document metadata (append-only; UNIQUE(url, version)).
            cursor.execute("""
                INSERT INTO documents (
                    url, domain, file_name, content_type, fetched_at,
                    parser_used, quality_score, total_chunks, metadata_json,
                    version, content_hash
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                url,
                domain,
                meta.get('file_name', ''),
                meta.get('content_type', ''),
                time.time(),
                meta.get('parser_used') or meta.get('parser', 'unknown'),
                meta.get('quality_score', 0.0),
                len(chunks),
                json.dumps(meta),
                version,
                content_hash,
            ))

            # Insert chunks into FTS + content-addressable dedup. The CA-vector
            # lookup runs on the SAME connection (no nested pool acquires).
            embed_jobs: list[tuple[int, str, str]] = []  # (rowid, c_hash, text)
            for chunk in chunks:
                text = chunk.text
                header = chunk.metadata.get('header_path', 'Root')
                c_hash = hashlib.blake2b(text.encode("utf-8"),
                                         digest_size=32).hexdigest()

                # Canonical dedup entry (INSERT OR IGNORE = first-write wins).
                cursor.execute("""
                    INSERT OR IGNORE INTO chunks_ca (
                        chunk_hash, text, url, header_path, metadata_json, first_seen
                    ) VALUES (?, ?, ?, ?, ?, ?)
                """, (c_hash, text, url, header, json.dumps(chunk.metadata), time.time()))

                cursor.execute("""
                    INSERT INTO chunks_fts (url, header_path, text, metadata_json)
                    VALUES (?, ?, ?, ?)
                """, (url, header, text, json.dumps(chunk.metadata)))
                rowid = cursor.lastrowid

                if embed_ok and rowid is not None:
                    embed_jobs.append((rowid, c_hash, text))

            # Batch-embed the whole document (one model call over many chunks),
            # then write vectors keyed by their FTS rowid. Dedup-aware and
            # bit-identical to per-chunk embedding.
            if embed_jobs:
                vec_by_hash = self._embed_chunk_batch(
                    cursor, [(c_hash, text) for _r, c_hash, text in embed_jobs])
                for rowid, c_hash, _text in embed_jobs:
                    vec = vec_by_hash.get(c_hash)
                    if vec is not None:
                        cursor.execute(
                            "INSERT OR REPLACE INTO chunk_vectors (chunk_rowid, url, vector) "
                            "VALUES (?, ?, ?)",
                            (rowid, url, vec),
                        )

        # chunk_vectors changed; drop the stale vector-scan matrix (S1).
        self._vec_mat_cache.clear()

        if self.bus is not None:
            self.bus.emit("document.ingested", url=url, version=version,
                          chunks=len(chunks))
        logger.info(f"Indexed {len(chunks)} chunks for {url} (v{version})")

    def _embed_chunk(self, cursor: sqlite3.Cursor, c_hash: str, text: str) -> bytes | None:
        """Return a vector for *text*, reusing a cached one for an identical
        chunk hash AND the current embedding fingerprint. Runs on the caller's
        connection (no nested pool acquire). Returns None if embedding fails
        (non-fatal)."""
        embed_fp = self.embeddings.fingerprint()
        try:
            cursor.execute(
                "SELECT vector FROM chunk_vectors_ca "
                "WHERE chunk_hash = ? AND embed_fp = ?",
                (c_hash, embed_fp),
            )
            row = cursor.fetchone()
            if row is not None and row[0] is not None:
                expected = self.embeddings.dim * self.embeddings.bytes_per_dim
                if len(row[0]) == expected:
                    return row[0]
                # Stale format (dim/model/quantize changed since the chunk was
                # last embedded): fall through and recompute so a cached vector
                # can never silently mismatch the current embedding config (A9).
        except Exception as e:
            # Best-effort cache probe: a broken query must not stall ingest.
            logger.debug(f"CA cache lookup failed for {c_hash[:16]}: {e}")
        try:
            vec = self.embeddings.vectorize(text)
        except Exception as e:  # embedding failures must not block indexing
            logger.warning(f"Embedding failed: {e}")
            return None
        if vec:
            with suppress(Exception):
                cursor.execute(
                    "INSERT OR REPLACE INTO chunk_vectors_ca "
                    "(chunk_hash, vector, embed_fp) VALUES (?, ?, ?)",
                    (c_hash, vec, embed_fp),
                )
                if self.bus is not None:
                    self.bus.emit("chunk.embedded", chunk_hash=c_hash,
                                  vector_dim=self.embeddings.dim)
        return vec

    def _embed_chunk_batch(self, cursor: sqlite3.Cursor,
                           items: list[tuple[str, str]]) -> dict[str, bytes]:
        """Embed many (c_hash, text) items in a single model call, reusing
        content-addressed vectors for any already embedded under the current
        fingerprint.

        Returns `{c_hash: vector_bytes}` for every item that produced a vector
        (cached or freshly embedded). Runs on the caller's connection. Batch
        lookup + batch embed so the model does one forward pass over many texts
        instead of one per chunk; bit-identical output to per-chunk embedding.
        Non-fatal: a failed item is simply absent from the result.
        """
        embed_fp = self.embeddings.fingerprint()
        expected = self.embeddings.dim * self.embeddings.bytes_per_dim
        result: dict[str, bytes] = {}

        # 1. Batch content-addressed cache lookup for every hash at once.
        cached: dict[str, bytes] = {}
        hashes = [h for h, _t in items]
        if hashes:
            ph = ",".join("?" * len(hashes))
            try:
                cursor.execute(
                    f"SELECT chunk_hash, vector FROM chunk_vectors_ca "  # nosec B608
                    f"WHERE chunk_hash IN ({ph}) AND embed_fp = ?",
                    (*hashes, embed_fp),
                )
                for ch, vec in cursor.fetchall():
                    if vec is not None and len(vec) == expected:
                        cached[ch] = vec
            except Exception as e:
                logger.debug(f"CA batch cache lookup failed: {e}")

        pending: list[tuple[str, str]] = []
        for c_hash, text in items:
            if c_hash in cached:
                result[c_hash] = cached[c_hash]
            else:
                pending.append((c_hash, text))

        if not pending:
            return result

        # 2. Batch-embed only the missing texts, in slices so huge documents
        # don't allocate one giant batch.
        batch_size = int(self.config.get('embeddings.batch_size', 16) or 0)
        if batch_size <= 0:
            batch_size = len(pending)
        for i in range(0, len(pending), batch_size):
            slice_ = pending[i:i + batch_size]
            try:
                vecs = self.embeddings.vectorize_batch([t for _h, t in slice_])
            except Exception as e:
                logger.warning(f"Batch embedding failed ({e}); item skipped.")
                vecs = []
            for (c_hash, _text), vec in zip(slice_, vecs, strict=False):
                if not vec:
                    continue
                result[c_hash] = vec
                with suppress(Exception):
                    cursor.execute(
                        "INSERT OR REPLACE INTO chunk_vectors_ca "
                        "(chunk_hash, vector, embed_fp) VALUES (?, ?, ?)",
                        (c_hash, vec, embed_fp),
                    )
                    if self.bus is not None:
                        self.bus.emit("chunk.embedded", chunk_hash=c_hash,
                                      vector_dim=self.embeddings.dim)
        return result

    def backfill_vectors(self) -> int:
        """Compute and store embeddings for chunks missing one (or with a stale
        dimension). Returns count backfilled.

        If the configured embedding mode/dimension differs from what is stored
        (e.g. switching from the 256-dim sparse hash to a 384-dim dense model),
        the mismatched rows are recomputed in place. Rebuilds are resumable: an
        interrupted run simply leaves some rows stale, which the next run picks
        up, so no destructive DELETE-all is ever needed.
        """
        if not self.config.get('embeddings.enabled', True):
            return 0
        expected_bytes = self._vector_dim * self.embeddings.bytes_per_dim
        stale_dim = False
        with self._db() as (_conn, cursor):
            # Detect dimension mismatch across ALL stored vectors, not just a
            # LIMIT 1 sample: an interrupted earlier backfill can leave some
            # rows fresh and some stale, and a one-row probe would miss it.
            cursor.execute("SELECT COUNT(*) FROM chunk_vectors")
            vec_count = cursor.fetchone()[0]
            if vec_count > 0:
                cursor.execute(
                    "SELECT COUNT(*) FROM chunk_vectors WHERE length(vector) != ?",
                    (expected_bytes,),
                )
                stale_count = cursor.fetchone()[0]
                stale_dim = stale_count > 0
                if stale_dim:
                    logger.info(
                        f"{stale_count} of {vec_count} vectors have a stale "
                        f"dimension (expected {expected_bytes} bytes); "
                        f"recomputing them in place."
                    )

        # Cheap count check: fully vectorized AND matching dim -> 0 work.
        # Skip the count shortcut when dims are stale (rows exist but are wrong).
        if not stale_dim:
            with self._db() as (_conn, cursor):
                cursor.execute("SELECT COUNT(*) FROM chunks_fts")
                fts_count = cursor.fetchone()[0]
            if fts_count == vec_count:
                return 0

        # Select chunks that are missing a vector OR carry a wrong-dimension one.
        # Streams in batches (fetchmany) instead of loading the whole table,
        # committing per batch so a long backfill is resumable (B5) and never
        # pins a giant write transaction.
        count = 0
        embed_fp = self.embeddings.fingerprint()
        with self._db() as (_conn, cursor):
            if stale_dim:
                # Any content-addressable vector built by a different embedding
                # setting can never be served again; drop it so the next embed
                # recomputes fresh instead of lazily expiring one row at a time.
                cursor.execute(
                    "DELETE FROM chunk_vectors_ca WHERE embed_fp IS NULL OR embed_fp != ?",
                    (embed_fp,),
                )
                if cursor.rowcount:
                    logger.info(
                        f"Cleared {cursor.rowcount} stale CA cache entries "
                        f"(embedding config changed)."
                    )
            sql = f"""
                SELECT c.rowid, c.url, c.text
                FROM chunks_fts c
                LEFT JOIN chunk_vectors v ON v.chunk_rowid = c.rowid
                WHERE v.chunk_rowid IS NULL
                   OR length(v.vector) != {expected_bytes}
            """  # nosec B608 (expected_bytes is an int, not user input)
            cursor.execute(sql)
            while True:
                rows = cursor.fetchmany(1000)
                if not rows:
                    break
                for rowid, url, text in rows:
                    try:
                        vec = self.embeddings.vectorize(text)
                    except Exception as e:
                        logger.warning(f"Backfill embedding failed {url}: {e}")
                        continue
                    cursor.execute(
                        "INSERT OR REPLACE INTO chunk_vectors (chunk_rowid, url, vector) VALUES (?, ?, ?)",
                        (rowid, url, vec)
                    )
                    count += 1
                _conn.commit()
        if count:
            logger.info(f"Backfilled {count} chunk embeddings.")
            self._vec_mat_cache.clear()
        return count

    def verify_vault(self) -> bool:
        """Run a three-phase integrity check over the vault.

        Phase 1 — verify every document's chunk count and content hash.
        Phase 2 — verify content-addressable chunks are internally consistent.
        Phase 3 — verify every stored vector's dimension matches the engine.

        Returns True if no errors were found.
        """
        errors = 0
        checks = 0

        # Phase 1: for each URL, the total FTS chunk count must equal the sum of
        # the declared chunk counts across all of that URL's document versions
        # (WORM means one URL may span several version rows). Two GROUP BY
        # queries joined as dicts — not a COUNT query per URL (B4 N+1).
        with self._db() as (_conn, cursor):
            declared_by_url = dict(cursor.execute(
                "SELECT url, SUM(total_chunks) FROM documents GROUP BY url"
            ).fetchall())
            actual_by_url = dict(cursor.execute(
                "SELECT url, COUNT(*) FROM chunks_fts GROUP BY url"
            ).fetchall())
            for url, declared_total in declared_by_url.items():
                fts_count = actual_by_url.get(url, 0)
                checks += 1
                if fts_count != declared_total:
                    logger.error(
                        "MISMATCH: %s declares %d chunks (all versions) but FTS has %d",
                        url, declared_total, fts_count,
                    )
                    errors += 1

        # Phase 2: content-addressable chunks must have non-empty text.
        with self._db() as (_conn, cursor):
            cursor.execute("SELECT chunk_hash, text FROM chunks_ca")
            for c_hash, text in cursor.fetchall():
                checks += 1
                recomputed = hashlib.blake2b(
                    (text or "").encode("utf-8"), digest_size=32
                ).hexdigest()
                if recomputed != c_hash:
                    logger.error(
                        "CORRUPTION: chunk %s… text hash does not match",
                        c_hash[:16],
                    )
                    errors += 1

        # Phase 3: every vector must match the configured dimension.
        expected_bytes = self._vector_dim * self.embeddings.bytes_per_dim
        with self._db() as (_conn, cursor):
            cursor.execute("SELECT chunk_rowid, length(vector) FROM chunk_vectors")
            for rid, vlen in cursor.fetchall():
                checks += 1
                if vlen != expected_bytes:
                    logger.error(
                        "BAD DIM: vector for chunk %d is %d bytes (expected %d)",
                        rid, vlen, expected_bytes,
                    )
                    errors += 1

        if errors == 0:
            logger.info(f"Vault integrity PASS ({checks} checks).")
        else:
            logger.error(f"Vault integrity FAIL ({errors} error(s) across {checks} checks).")
        return errors == 0

    def stats(self) -> dict[str, Any]:
        """High-level vault statistics for the `stats` action.

        Counts sources (distinct URLs), document versions, chunks, embedded
        vectors, schema version and page size — the numbers a promotion or
        maintenance pass needs to quantify the vault. Uses single aggregate
        queries (B4 N+1 discipline).
        """
        with self._db() as (_conn, cursor):
            sources = cursor.execute(
                "SELECT COUNT(DISTINCT url) FROM documents").fetchone()[0]
            doc_versions = cursor.execute(
                "SELECT COUNT(*) FROM documents").fetchone()[0]
            chunks = cursor.execute(
                "SELECT COUNT(*) FROM chunks_fts").fetchone()[0]
            vectors = cursor.execute(
                "SELECT COUNT(*) FROM chunk_vectors").fetchone()[0]
            schema_version = cursor.execute(
                "PRAGMA user_version").fetchone()[0]
            page_size = cursor.execute(
                "PRAGMA page_size").fetchone()[0]
        try:
            db_bytes = os.path.getsize(self.db_path)
        except OSError:
            db_bytes = 0
        dim = getattr(self.embeddings, "dim", None) or 0
        return {
            "vault": self.vault_name,
            "sources": sources,
            "doc_versions": doc_versions,
            "chunks": chunks,
            "vectors": vectors,
            "dim": dim,
            "mode": str(getattr(self.embeddings, "mode", "")),
            "conf_mode": str(self.config.get('embeddings.conf_mode', 'relative')),
            "schema_version": schema_version,
            "page_size": page_size,
            "db_bytes": db_bytes,
            "preferred_name": self.vault_name,
        }

    def confidence_distribution(self, probes: int = 4, recall: int = 6) -> dict[str, int]:
        """Run a handful of diagnostic probes through the real hybrid search and
        aggregate the resulting confidence bands.

        Confidence is computed at retrieval time (it is query-relative), so it
        cannot be read from stored rows. This samples how the vault's own
        content actually ranks, so the histogram reveals "all-medium" flatness
        (the sign of a mis-tuned band) versus a healthy high/medium/low spread.

        Probe queries are the most *distinctive* header segments — the deepest,
        longest portion of each header path (e.g. the "… > Yield by region"
        tail, preferring multi-word segments). Generic single-word headers
        ("Production", "Farmers") are avoided: they match too broadly to be
        keyword-backed, so probing them would report a misleading all-medium
        distribution even when real topical queries spread normally.
        """
        with self._db() as (_conn, cursor):
            rows = cursor.execute(
                "SELECT header_path FROM chunks_ca "
                "WHERE header_path IS NOT NULL AND trim(header_path) != '' "
                "LIMIT 200").fetchall()
        # Split each header path and keep its deepest segment; rank by length
        # (multi-word, specific) and drop generic single-word labels.
        candidates: list[str] = []
        seen: set[str] = set()
        for (hp,) in rows:
            segs = [s.strip() for s in hp.split(">") if s and s.strip()]
            seg = segs[-1] if segs else ""
            if not seg or seg in seen:
                continue
            seen.add(seg)
            if seg.lower() in ("production", "farmers", "history", "introduction",
                               "overview", "summary", "conclusion", "abstract"):
                continue
            candidates.append(seg)
        # Prefer the longest (most specific, keyword-dense) segments.
        candidates.sort(key=lambda s: len(s.split()), reverse=True)
        phrases = candidates[: max(1, probes * 4)] or sorted(seen) or ["crop"]
        if not phrases:
            return {"high": 0, "medium": 0, "low": 0}
        chosen = phrases[: max(1, probes)]
        counts: dict[str, int] = {"high": 0, "medium": 0, "low": 0}
        for q in chosen:
            for chunk in self.search_vault(q, limit=recall, hybrid=True):
                counts[chunk.metadata.get("confidence", "low")] = \
                    counts.get(chunk.metadata.get("confidence", "low"), 0) + 1
        return counts

    def migrate_page_size(self, target: int | None = None) -> bool:
        """Rewrite the vault DB at a different SQLite page size via `VACUUM INTO`.

        SQLite only honors `PRAGMA page_size` while a database file is still
        empty, so vaults created before the 16 KB default keep their old page
        size (typically 4096). This rebuilds the file at `target` bytes per
        page without touching live connections, preserving all data.

        Returns True if the vault was rewritten, False if it was already at the
        target size (or the rewrite failed).
        """
        target = int(target or self.page_size)
        with self._db() as (_conn, cursor):
            current = cursor.execute("PRAGMA page_size").fetchone()[0]
        if current == target:
            return False

        tmp_path = f"{self.db_path}.ps{target}"
        # SQLite's VACUUM INTO refuses to overwrite an existing file, so clear
        # any stale temp from a previous failed attempt before retrying.
        with suppress(FileNotFoundError):
            os.remove(tmp_path)
        ok = False
        try:
            # A dedicated connection outside the pool so WAL is quiet.
            conn = sqlite3.connect(self.db_path, timeout=DB_TIMEOUT)
            try:
                conn.execute(f"PRAGMA page_size = {target}")
                # Cannot parameterize VACUUM INTO's filename; escape any single
                # quote in the path so a root_dir with one can't break the SQL.
                escaped = tmp_path.replace("'", "''")
                conn.execute(f"VACUUM INTO '{escaped}'")
            finally:
                conn.close()
            with sqlite3.connect(tmp_path) as v:
                new_size = v.execute("PRAGMA page_size").fetchone()[0]
            if new_size != target:
                logger.error(
                    f"migrate_page_size: VACUUM INTO produced {new_size}, "
                    f"expected {target}."
                )
                with suppress(FileNotFoundError):
                    os.remove(tmp_path)
                return False
            ok = True
        except Exception as e:
            logger.error(f"migrate_page_size failed: {e}")
            with suppress(FileNotFoundError):
                os.remove(tmp_path)
            return False

        # Swap the rebuilt file into place and refresh the pool.
        self._pool.close_all()
        os.replace(tmp_path, self.db_path)
        for suffix in ("-wal", "-shm"):
            with suppress(FileNotFoundError):
                os.remove(self.db_path + suffix)
        self._pool = ConnectionPool(self.db_path, CONNECTION_POOL_SIZE, self.page_size)
        logger.info(f"Vault page size migrated {current} -> {target} bytes.")
        return ok

    def ingest_chunks_parallel(self, url: str, chunks: list[Chunk],
                               meta: dict[str, Any]) -> None:
        """Ingest chunks through a parallel reader→embed→write pipeline.

        Kept optional (guarded by config 'indexer.parallel'): for typical
        research vaults the sequential path is fast enough, and this avoids
        thread overhead on small batches. When enabled, embedding work is
        spread across WORKER_THREADS threads while the DB writer stays single.
        """
        if not self.config.get('indexer.enable_fts', True):
            return
        if not self.config.get('indexer.parallel', False) or len(chunks) < 8:
            return self.index_document(url, chunks, meta)

        # Filter near-duplicates BEFORE the embed pipeline. Vectors are
        # keyed by position in *this* list, so dedup must not shrink the list
        # between embedding and writing (B1: it used to, misaligning every
        # vector after the first dropped chunk).
        with self._db() as (_conn, cursor):
            chunks = self._filter_near_dupes(cursor, url, chunks)
        if not chunks:
            return

        embed_ok = self.config.get('embeddings.enabled', True)
        batch_size = max(1, int(self.config.get('embeddings.batch_size', 16) or 0) or 16)
        # Slice chunks into embedding mini-batches keyed by their start index in
        # `chunks`, so the model does one forward pass over several texts per
        # worker instead of one pass per chunk. Bit-identical output.
        work_batches: list[tuple[int, list[str]]] = [
            (i, [c.text for c in chunks[i:i + batch_size]])
            for i in range(0, len(chunks), batch_size)
        ]
        work_q: queue.Queue = queue.Queue(maxsize=PIPELINE_QUEUE_SIZE)
        result_q: queue.Queue = queue.Queue(maxsize=PIPELINE_QUEUE_SIZE)
        error_holder: list[Exception] = []
        results: list[tuple[int, str, bytes | None]] = [None] * len(chunks)  # type: ignore[list-item]
        # A sentinel is pushed to the WORK queue (not result_q) and consumed by
        # the workers themselves, so shutdown can never race with the consumer
        # (A11). The reader then collects exactly len(work_batches) results.
        sentinel = (-1, None)

        def _embed_worker() -> None:
            while True:
                start, texts = work_q.get()
                try:
                    if start == -1:
                        return
                    vecs = (self.embeddings.vectorize_batch(texts) if embed_ok
                            else [None] * len(texts))
                except Exception as e:
                    error_holder.append(e)
                    vecs = [None] * len(texts)
                finally:
                    work_q.task_done()
                result_q.put((start, vecs))

        # Start the workers, then feed work and sentinels.
        threads = [threading.Thread(target=_embed_worker, daemon=True)
                   for _ in range(WORKER_THREADS)]
        for t in threads:
            t.start()

        # Drain result_q CONCURRENTLY with feeding work_q via a reader thread
        # started BEFORE any feeding. result_q is bounded, so it must be read
        # while the main thread feeds work; otherwise the workers block on
        # result_q.put() once it fills and stop consuming work_q, deadlocking
        # the main thread's work_q.put() — a real hang for any batch larger
        # than PIPELINE_QUEUE_SIZE once embeddings are non-trivial.
        collected: list[tuple[int, object]] = []

        def _collector() -> None:
            # Exactly len(work_batches) real results are produced (sentinels are
            # consumed by workers and produce no result), so this cannot hang.
            for _ in range(len(work_batches)):
                start, vecs = result_q.get()
                collected.append((start, vecs))

        reader = threading.Thread(target=_collector, daemon=True)
        reader.start()

        for start, texts in work_batches:
            work_q.put((start, texts))
        # Wake EXACTLY the number of workers we started so none hang waiting.
        for _ in threads:
            work_q.put(sentinel)
        # reader has now collected exactly len(work_batches) results; every
        # worker has seen its sentinel and exited, so join cannot hang.
        reader.join()
        for start, vecs in collected:
            for j, vec in enumerate(vecs):  # type: ignore[arg-type]
                results[start + j] = vec  # type: ignore[assignment]
        for t in threads:
            t.join()

        # Single writer thread commits the whole batch (dedup + vectors).
        with self._db() as (_conn, cursor):
            domain = urlparse(url).netloc
            content_hash = hashlib.blake2b(
                "\n".join(c.text for c in chunks).encode("utf-8"),
                digest_size=32,
            ).hexdigest()
            cursor.execute(
                "SELECT COALESCE(MAX(version), 0) FROM documents WHERE url = ?",
                (url,),
            )
            version = cursor.fetchone()[0] + 1
            cursor.execute("""
                INSERT INTO documents (
                    url, domain, file_name, content_type, fetched_at,
                    parser_used, quality_score, total_chunks, metadata_json,
                    version, content_hash
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                url, domain, meta.get('file_name', ''),
                meta.get('content_type', ''), time.time(),
                meta.get('parser_used') or meta.get('parser', 'unknown'),
                meta.get('quality_score', 0.0), len(chunks),
                json.dumps(meta), version, content_hash,
            ))
            for idx, chunk in enumerate(chunks):
                text = chunk.text
                c_hash = hashlib.blake2b(text.encode("utf-8"),
                                         digest_size=32).hexdigest()
                cursor.execute("""
                    INSERT OR IGNORE INTO chunks_ca (
                        chunk_hash, text, url, header_path, metadata_json, first_seen
                    ) VALUES (?, ?, ?, ?, ?, ?)
                """, (c_hash, text, url, chunk.metadata.get('header_path', 'Root'),
                      json.dumps(chunk.metadata), time.time()))
                cursor.execute("""
                    INSERT INTO chunks_fts (url, header_path, text, metadata_json)
                    VALUES (?, ?, ?, ?)
                """, (url, chunk.metadata.get('header_path', 'Root'), text,
                      json.dumps(chunk.metadata)))
                rowid = cursor.lastrowid
                vec = results[idx]
                if vec:
                    cursor.execute(
                        "INSERT OR REPLACE INTO chunk_vectors (chunk_rowid, url, vector) "
                        "VALUES (?, ?, ?)",
                        (rowid, url, vec),
                    )
        # chunk_vectors changed; drop the stale vector-scan matrix (S1).
        self._vec_mat_cache.clear()
        if error_holder:
            logger.warning(f"{len(error_holder)} embedding errors during parallel ingest.")
        if self.bus is not None:
            self.bus.emit("document.ingested", url=url, version=version,
                          chunks=len(chunks), parallel=True)
        logger.info(f"Indexed {len(chunks)} chunks for {url} (v{version}, parallel)")

    def latest_content_hash(self, url: str) -> str | None:
        """Return the newest stored `content_hash` for a URL, or None.

        Local ingestion uses this to skip re-indexing a file whose extracted
        content is unchanged since the last ingest (content-based freshness
        instead of the HTTP cache TTL): identical content -> identical hash ->
        no new WORM version is written. `--force` bypasses the check.
        """
        with self._db() as (_conn, cursor):
            cursor.execute(
                "SELECT content_hash FROM documents WHERE url = ? "
                "ORDER BY version DESC LIMIT 1",
                (url,),
            )
            row = cursor.fetchone()
        return row[0] if row else None

    @staticmethod
    def _fts_query(query: str, op: str = "AND") -> str | None:
        """Build a safe FTS5 MATCH expression from a raw user query.

        Wraps each whitespace token as a quoted phrase so FTS operators
        (quotes, parentheses, *, ^, :, -) in user input cannot alter the
        query semantics or raise syntax errors. Returns None if the query
        contains no usable tokens (e.g. empty, punctuation-only).

        `op` joins the quoted tokens: "AND" requires every token (strict, used
        for the fts_fast strong-signal gate and exact recall), "OR" matches any
        token (lenient fallback so a long research question with a couple of
        absent terms still surfaces keyword-backed candidates instead of
        collapsing to a pure-vector search).
        """
        tokens = []
        for token in re.findall(r'\S+', query):
            cleaned = re.sub(r'["()*^:\-]', ' ', token)
            cleaned = ' '.join(cleaned.split())
            if cleaned:
                tokens.append(f'"{_fts_token(cleaned)}"')
        if not tokens:
            return None
        return f' {op} '.join(tokens)

    def search_vault(self, query: str, limit: int = 20, domain: str | None = None,
                     hybrid: bool | None = None, max_per_source: int = 0) -> list[Chunk]:
        """Perform FTS5 search (or hybrid FTS+vector when enabled).

        Args:
            query: free-text query.
            limit: max results.
            domain: restrict to a netloc substring, if given.
            hybrid: None -> use config; True/False -> override.
            max_per_source: cap on chunks from any one source URL in the result
                (0 = unlimited). Research recall uses this to keep the set
                source-diverse so a single rich page can't crowd out every other
                source.
        """
        if not self.config.get('indexer.enable_fts', True):
            return []

        if not query or not query.strip():
            return []  # empty/whitespace query -> no results, not a crash

        use_hybrid = self.config.get('embeddings.hybrid_search', True) if hybrid is None else hybrid
        if use_hybrid and self.embeddings.enabled:
            return self._search_hybrid(query, limit, domain, max_per_source=max_per_source)

        results = []
        with self._db() as (_conn, cursor):
            fts_match = self._fts_query(query)
            if not fts_match:
                return []  # punctuation-only query -> nothing to match
            where = "chunks_fts MATCH ?"
            params: list[Any] = [fts_match]

            if domain:
                where += " AND url LIKE ?"
                params.append(f'%{domain}%')

            # FTS5 query with ranking
            sql = f"""
                SELECT url, header_path, text, metadata_json, rank, rowid
                FROM chunks_fts
                WHERE {where}
                ORDER BY rank
                LIMIT ?
            """  # nosec B608 (where is a fixed clause; values are bound)
            cursor.execute(sql, (*params, limit))

            for row in cursor.fetchall():
                url, header_path, text, meta_json, rank, rowid = row
                meta = json.loads(meta_json)
                meta['search_rank'] = rank
                meta['source_url'] = url
                meta['chunk_id'] = rowid
                results.append(Chunk(text=text, metadata=meta))

        return results

    def _vector_scan(self, cursor: sqlite3.Cursor, qvec: bytes,
                     top_k: int, domain: str | None) -> list[tuple[float, int, str]]:
        """Brute-force cosine over the whole vector table in one numpy matmul.

        All stored vectors are loaded into a single contiguous float32 buffer
        and dotted with the query vector — BLAS/LAPACK speed instead of a
        per-row Python cosine loop, which is ~15x slower on a 100k-row vault
        (H2.11, B2: one fetchall, one buffer, no per-row Python). The matrix
        is cached keyed on (row count, byte width) so repeated searches on an
        unchanged vault skip the reload. Falls back to the per-row cosine path
        when numpy is absent or a payload is malformed.
        """
        vec_where = ""
        vec_params: list[Any] = []
        if domain:
            vec_where = " WHERE url LIKE ?"
            vec_params.append(f'%{domain}%')
        cursor.execute(
            "SELECT chunk_rowid, url, vector FROM chunk_vectors" + vec_where,  # nosec B608
            vec_params,
        )
        rows = cursor.fetchall()
        if not rows or not qvec:
            return []

        expected_bytes = self._vector_dim * self.embeddings.bytes_per_dim
        mat: Any = None
        rids: list[int] = []
        urls: list[str] = []
        # Only the unfiltered full-table scan can reuse the cached matrix.
        if domain is None and self._vec_mat_cache.get("count") == len(rows) \
                and self._vec_mat_cache.get("expected") == expected_bytes:
            mat = self._vec_mat_cache["mat"]
            rids = self._vec_mat_cache["rids"]
            urls = self._vec_mat_cache["urls"]
        elif _np is not None:
            rids = [r[0] for r in rows]
            urls = [r[1] for r in rows]
            blobs = [r[2] for r in rows]
            well_formed = all(isinstance(b, (bytes, bytearray))
                              and len(b) == expected_bytes for b in blobs)
            if well_formed:
                try:
                    if self.embeddings.bytes_per_dim == 1:
                        arr = _np.frombuffer(b"".join(blobs), dtype=_np.int8)
                        arr = arr.astype(_np.float32) / 127.0
                    else:
                        arr = _np.frombuffer(b"".join(blobs), dtype=_np.float32)
                    mat = arr.reshape(-1, self._vector_dim)
                    if domain is None:
                        self._vec_mat_cache.update(
                            count=len(rows), expected=expected_bytes,
                            mat=mat, rids=rids, urls=urls,
                        )
                except ValueError:
                    mat = None

        if mat is not None and _np is not None:
            try:
                if self.embeddings.bytes_per_dim == 1:
                    q = _np.frombuffer(qvec, dtype=_np.int8).astype(_np.float32) / 127.0
                else:
                    q = _np.frombuffer(qvec, dtype=_np.float32)
                sims = mat @ q
                if len(sims) == len(rids):
                    k = min(top_k, len(sims))
                    if k <= 0:
                        return []
                    # argpartition is O(N) for the top-k, then a tiny sort.
                    idx = _np.argpartition(-sims, k - 1)[:k]
                    idx = idx[_np.argsort(-sims[idx])]
                    return [(float(sims[i]), rids[i], urls[i]) for i in idx]
            except (ValueError, TypeError):
                pass  # malformed query payload -> fall through to per-row

        # Fallback: per-row cosine (numpy absent, or malformed payloads).
        scored: list[tuple[float, int, str]] = []
        for rid, u, blob in rows:
            s = EmbeddingsEngine.cosine(qvec, blob, self._vector_dim)
            scored.append((s, rid, u))
        scored.sort(key=lambda t: t[0], reverse=True)
        return scored[:top_k]

    def _rerank(self, query: str, chunks: list[Chunk]) -> list[Chunk]:
        """Optionally re-order a recalled set with a cross-encoder reranker.

        Gated on `embeddings.reranker_model` (empty = disabled, the default).
        Loaded lazily via fastembed's TextCrossEncoder and cached on the
        instance. Any load/run failure degrades to the input order (a reranker
        is an enhancement, never a recall gate).
        """
        model_name = str(self.config.get('embeddings.reranker_model', '') or '')
        if not model_name or len(chunks) < 2:
            return chunks
        try:
            if self._reranker is None:
                from fastembed.rerank.cross_encoder import TextCrossEncoder
                self._reranker = TextCrossEncoder(model_name)
            docs = [c.text for c in chunks]
            ranked = list(self._reranker.rerank(query, docs))
            score_by_idx: dict[int, float] = {}
            for item in ranked:
                if isinstance(item, tuple):
                    idx, score = item[0], item[1]
                else:
                    idx, score = item.index, item.score  # pyright: ignore[reportAttributeAccessIssue]
                score_by_idx[idx] = float(score)
            ordered = sorted(range(len(chunks)),
                             key=lambda i: score_by_idx.get(i, 0.0),
                             reverse=True)
            reranked = [chunks[i] for i in ordered]
            for pos, orig_idx in enumerate(ordered):
                reranked[pos].metadata['rerank_score'] = score_by_idx.get(orig_idx, 0.0)
            return reranked
        except Exception as e:
            logger.warning(f"Reranker unavailable ({e}); keeping original order.")
            return chunks

    @staticmethod
    def _diverse_order(fused: list[tuple[int, float]],
                       limit: int, max_per_source: int,
                       url_by_rid: dict[int, str]) -> list[tuple[int, float]]:
        """Rebalance a rank-ordered RRF list so no single source dominates the
        top-N, keeping the set source-diverse while preserving relevance.

        Walks the fused (already relevance-ranked) list in order and admits each
        hit unless its source URL has already contributed `max_per_source` hits,
        continuing to fill until `limit` is reached. The top-ranked hit is always
        admitted (its source can't be saturated yet), so the most relevant result
        is never demoted. Returns the full list when `limit <= 0` (no cap) or
        `max_per_source <= 0` (diversity off)."""
        if max_per_source <= 0 or limit <= 0:
            return fused[:limit] if limit > 0 else fused
        per_src: dict[str, int] = {}
        selected: list[tuple[int, float]] = []
        for rid, score in fused:
            if len(selected) >= limit:
                break
            src = url_by_rid.get(rid, '?')
            if per_src.get(src, 0) >= max_per_source:
                continue
            per_src[src] = per_src.get(src, 0) + 1
            selected.append((rid, score))
        return selected

    def _search_hybrid(self, query: str, limit: int, domain: str | None,
                       max_per_source: int = 0) -> list[Chunk]:
        """Fuse FTS5 keyword ranks and vector-similarity ranks via Reciprocal
        Rank Fusion (RRF). Returns Chunks best matching the query.

        max_per_source > 0 caps how many chunks any single source URL may
        contribute to the returned set (see `_diverse_order`); the top-ranked
        hit is always kept regardless of source, so relevance stays first.
        Reads the FTS candidate pool in a SINGLE query (fetching the full rows
        once and stashing them by rowid) instead of a second rowid-IN round
        trip for the fast path and the fused result set (B3)."""
        if not query or not query.strip():
            return []
        k = 60  # RRF constant
        fts_pool = int(self.config.get('indexer.search_limit', 20) * 3)
        vec_pool = self.config.get('embeddings.top_k', 40)

        with self._db() as (_conn, cursor):
            # --- FTS candidate list: one query, rows stashed by rowid ---
            fts_match = self._fts_query(query)
            fts_where = "chunks_fts MATCH ?"
            fts_params: list[Any] = [fts_match]
            if domain:
                fts_where += " AND url LIKE ?"
                fts_params.append(f'%{domain}%')
            # A query with no FTS tokens (e.g. "*", operators only) still has
            # semantic content the embedding model can match: run the vector
            # scan alone instead of returning [] (A6).
            fts_rows: list[tuple[int, str]] = []
            row_by_id: dict[int, tuple[str, str, str, str]] = {}
            if fts_match:
                sql = f"""
                    SELECT rowid, url, header_path, text, metadata_json
                    FROM chunks_fts
                    WHERE {fts_where}
                    ORDER BY rank
                    LIMIT ?
                """  # nosec B608 (fts_where is a fixed clause; values are bound)
                cursor.execute(sql, (*fts_params, fts_pool))
                full_rows = cursor.fetchall()
                fts_rows = [(rid, url) for rid, url, _hp, _text, _mj in full_rows]
                row_by_id = {rid: (url, hp, text, mj)
                             for rid, url, hp, text, mj in full_rows}

            # --- FTS5 strong-signal fast path (P1.1) ---
            # When the FTS5 AND-match alone fills the requested result set, the
            # query is a strong keyword signal: skip the (more expensive) vector
            # scan entirely and return the FTS ranking directly. Tagged
            # retrieval='fts_fast' so downstream provenance is explicit.
            if (self.config.get('embeddings.fts_fast_path', True)
                    and len(fts_rows) >= limit):
                # Recency weighting must apply uniformly: without it, a 2-year-old
                # keyword match would rank identically to a fresh page even when
                # recency_half_life_days is configured (A3). We reorder the FTS
                # candidates by recency decay before slicing to the limit.
                half_life = float(self.config.get('embeddings.recency_half_life_days', 0) or 0)
                selected: list[tuple[int, str]] = fts_rows[:limit]
                if half_life > 0 and selected:
                    now = time.time()
                    urls = list({u for _rid, u in selected if u})
                    fetched_by_url: dict[str, float] = {}
                    if urls:
                        placeholders = ",".join("?" * len(urls))
                        cursor.execute(
                            f"SELECT url, MAX(fetched_at) FROM documents "  # nosec B608
                            f"WHERE url IN ({placeholders}) GROUP BY url",
                            urls,
                        )
                        fetched_by_url = dict(cursor.fetchall())

                    def _decay(u: str) -> float:
                        fetched = fetched_by_url.get(u)
                        if not fetched:
                            return 1.0
                        age = max(0.0, (now - fetched) / 86400.0)
                        return 0.5 ** (age / half_life)

                    selected.sort(key=lambda t: (_decay(t[1]), t[0]), reverse=True)
                results: list[Chunk] = []
                for rid, url in selected:
                    url, _hp, text, meta_json = row_by_id[rid]
                    meta = json.loads(meta_json)
                    meta['source_url'] = url
                    meta['chunk_id'] = rid
                    meta['retrieval'] = 'fts_fast'
                    meta['hybrid_score'] = None
                    # Confidence band is deliberately 'medium', not 'high':
                    # the vector scan was skipped, so semantic closeness is
                    # unverified. 'high' is reserved for hybrid hits that
                    # matched BOTH the keyword AND vector lists (see the
                    # confidence-band derivation below).
                    meta['confidence'] = 'medium'
                    results.append(Chunk(text=text, metadata=meta))
                return results

            # --- OR-fallback for keyword-backed candidates (A-OR) ---
            # A long research question whose strict AND-match is empty (any one
            # absent term zeroes the whole AND) would otherwise fall through to
            # a pure-vector search, tagging every hit 'medium' because no FTS
            # keyword match exists to back a 'high'. Retry with an any-token OR
            # so the RRF set is keyword-backed again and confidence can spread.
            # The fallback only rescues genuinely topical queries: it requires
            # at least two distinct query tokens to match the corpus, so a
            # single coincidental token (e.g. one stray "marketplace" in an
            # otherwise off-topic query) does NOT fake a keyword-backed set.
            if not fts_rows and fts_match and ' AND ' in fts_match:
                distinct = sum(
                    1 for _t in re.findall(r'\S+', query)
                    if (cand := re.sub(r'["()*^:\-]', ' ', _t).strip())
                    and _token_matches(cursor, cand)
                )
                if distinct >= 2:
                    or_match = self._fts_query(query, op="OR")
                    if or_match:
                        cursor.execute("""
                            SELECT rowid, url, header_path, text, metadata_json
                            FROM chunks_fts
                            WHERE chunks_fts MATCH ?
                            ORDER BY rank
                            LIMIT ?
                        """, (or_match, fts_pool))
                        or_rows = cursor.fetchall()
                        if or_rows:
                            fts_rows = [(rid, url) for rid, url, _hp, _t, _m in or_rows]
                            row_by_id = {rid: (url, hp, text, mj)
                                         for rid, url, hp, text, mj in or_rows}

            # --- vector candidate list (one numpy matmul over the whole table) ---
            scored: list[tuple[float, int, str]] = []
            if vec_pool > 0:
                qvec = self.embeddings.vectorize(query)
                scored = self._vector_scan(cursor, qvec, vec_pool, domain)

            # --- RRF fuse ---
            rrf: dict[int, float] = {}
            fts_rids: set[int] = set()
            for rank, (rid, _u) in enumerate(fts_rows):
                rrf[rid] = rrf.get(rid, 0.0) + 1.0 / (k + rank + 1)
                fts_rids.add(rid)
            vec_rids: set[int] = set()
            for rank, (_score, rid, _u) in enumerate(scored):
                rrf[rid] = rrf.get(rid, 0.0) + 1.0 / (k + rank + 1)
                vec_rids.add(rid)

            # --- Recency weighting (P1.2) ---
            # Optionally dampen stale hits: rrf *= 0.5 ** (age_days / half_life).
            # Half-life is 0 (disabled) by default.
            half_life = float(self.config.get('embeddings.recency_half_life_days', 0) or 0)
            if half_life > 0 and rrf:
                now = time.time()
                url_by_rid = dict(fts_rows)
                url_by_rid.update({rid: u for _s, rid, u in scored})
                urls = list({u for u in url_by_rid.values() if u})
                if urls:
                    placeholders = ",".join("?" * len(urls))
                    cursor.execute(
                        f"SELECT url, MAX(fetched_at) FROM documents "  # nosec B608
                        f"WHERE url IN ({placeholders}) GROUP BY url",
                        urls,
                    )
                    fetched_by_url = dict(cursor.fetchall())
                    for rid, url in url_by_rid.items():
                        fetched = fetched_by_url.get(url)
                        if fetched:
                            age_days = max(0.0, (now - fetched) / 86400.0)
                            rrf[rid] *= 0.5 ** (age_days / half_life)

            if not rrf:
                return []

            # URL per result, for source-diversity rebalancing (built outside
            # the recency block so `_diverse_order` always has it).
            url_by_rid: dict[int, str] = dict(fts_rows)
            url_by_rid.update({rid: u for _s, rid, u in scored})

            fused = sorted(rrf.items(), key=lambda kv: kv[1], reverse=True)
            order = self._diverse_order(fused, limit, max_per_source, url_by_rid)
            ids = [rid for rid, _ in order]

            # --- Confidence bands ---
            # Confidence is derived from how strong the fused evidence is, not
            # just ratio-to-top (which stays ~0.9 even for weak queries because
            # RRF scores cluster). RRF scores are rank-bucketed, so on a
            # homogeneous vault the absolute fused scores cluster into one band
            # and leave every hit "medium" — the failure the set-relative mode
            # fixes. Two modes:
            #   "relative" (default): confidence is ordinal WITHIN the returned
            #      set. The top hit(s) that clearly clear the set's own tail are
            #      "high"; hits hugging the coincidence floor are "low"; the
            #      middle is "medium". A flat homogeneous recall set therefore
            #      produces a spread, not all-"medium". Only a keyword-backed
            #      set (a genuine FTS match near the top) can crown "high" —
            #      a pure-vector/off-topic set never does, matching verify's
            #      corpus-scaled coincidence-floor logic.
            #   "absolute" (legacy): old thresholds conf_high_abs/conf_low_abs
            #      on the raw fused score (matched both lists, or score above
            #      an absolute ceiling).
            conf_mode = str(self.config.get('embeddings.conf_mode', 'relative')).lower()
            conf_by_rid: dict[int, str] = {}
            if conf_mode == 'absolute':
                conf_high_abs = float(self.config.get('embeddings.conf_high_abs', 0.025))
                conf_low_abs = float(self.config.get('embeddings.conf_low_abs', 0.020))
                for rid, score in order:
                    matched_both = rid in fts_rids and rid in vec_rids
                    if matched_both or score >= conf_high_abs:
                        conf_by_rid[rid] = "high"
                    elif score >= conf_low_abs:
                        conf_by_rid[rid] = "medium"
                    else:
                        conf_by_rid[rid] = "low"
            elif order:
                n = len(order)
                top_score = order[0][1]
                tail_score = order[-1][1]
                spread = max(top_score - tail_score, 1e-9)
                # Keyword-backed: a real FTS match sits near the top. A set
                # whose top hits are vector-only (off-topic / keyword-free) can
                # never be "high".
                top_half = order[:max(1, (n + 1) // 2)]
                keyword_backed = any(rid in fts_rids for rid, _score in top_half)
                # Only the top ~20% of the set may be "high"; only the bottom
                # ~half (hugging the coincidence floor) may be "low".
                n_high = max(1, (n * 2) // 10 + (1 if n % 10 >= 5 else 0))
                n_low = max(1, n // 2)
                for idx, (rid, score) in enumerate(order):
                    matched_both = rid in fts_rids and rid in vec_rids
                    rel = (score - tail_score) / spread
                    if (matched_both and idx == 0) or (
                            keyword_backed and rel >= 0.66 and idx < n_high):
                        conf_by_rid[rid] = "high"
                    elif rel <= 0.10 and idx >= n_low:
                        conf_by_rid[rid] = "low"
                    else:
                        conf_by_rid[rid] = "medium"

            results: list[Chunk] = []
            if ids:
                # Fetch only the rows the single FTS query did not already
                # return (vector-only hits), then assemble in RRF order (B3).
                missing = [rid for rid in ids if rid not in row_by_id]
                fetched_by_id: dict[int, tuple[str, str, str, str]] = {}
                if missing:
                    placeholders = ",".join("?" * len(missing))
                    sql = f"""
                        SELECT rowid, url, header_path, text, metadata_json
                        FROM chunks_fts WHERE rowid IN ({placeholders})
                    """  # nosec B608 (placeholders are positional ? marks)
                    cursor.execute(sql, missing)
                    fetched_by_id = {rid: (url, hp, text, mj)
                                     for rid, url, hp, text, mj in cursor.fetchall()}
                for rid in ids:
                    row = row_by_id.get(rid) or fetched_by_id.get(rid)
                    if row is None:
                        continue
                    url, _hp, text, meta_json = row
                    meta = json.loads(meta_json)
                    meta['hybrid_score'] = rrf.get(rid, 0.0)
                    meta['confidence'] = conf_by_rid.get(rid, "low")
                    meta['source_url'] = url
                    meta['chunk_id'] = rid
                    meta['retrieval'] = 'hybrid'
                    results.append(Chunk(text=text, metadata=meta))
            # Optional cross-encoder re-ranking of the final recalled set.
            if results and self.config.get('embeddings.reranker_model', ''):
                results = self._rerank(query, results)
            return results

    def document_exists(self, url: str, ttl_seconds: int) -> bool:
        """Check if a document is in the vault and not expired."""
        with self._db() as (_conn, cursor):
            cursor.execute(
                "SELECT fetched_at FROM documents WHERE url = ?",
                (url,)
            )
            row = cursor.fetchone()
        if not row:
            return False
        fetched_at = row[0]
        # ttl_seconds <= 0 means "never expire" (documented contract).
        if ttl_seconds <= 0:
            return True
        return (time.time() - fetched_at) < ttl_seconds

    def get_chunks_for_url(self, url: str) -> list[Chunk]:
        """Return every stored chunk for a URL (in insertion order)."""
        results: list[Chunk] = []
        with self._db() as (_conn, cursor):
            cursor.execute(
                "SELECT header_path, text, metadata_json FROM chunks_fts WHERE url = ?",
                (url,)
            )
            for header_path, text, meta_json in cursor.fetchall():
                meta = json.loads(meta_json)
                if "source_url" not in meta:
                    meta["source_url"] = url
                if header_path and "header_path" not in meta:
                    meta["header_path"] = header_path
                results.append(Chunk(text=text, metadata=meta))
        return results

# =============================================================================
# 4. NETWORK RESILIENCE CORE
# =============================================================================

class NetworkFetcher:
    """Executes explicit strategy chain with pre-flight check."""

    # Hostnames that must never be fetched regardless of how they resolve:
    # cloud metadata endpoints are the classic SSRF prize. `.internal` /
    # `.local` are private-DNS TLDs for resolver search domains / mDNS.
    _SSRF_BLOCKED_HOSTS = frozenset({
        "metadata.google.internal", "metadata.google", "localhost",
    })

    def __init__(self, config: ConfigManager):
        self.config = config
        self._cookie_string = config.get('auth.cookie_string', '')
        self._solver_enabled = config.get('solver.enabled', False)
        self._solver_url = config.get('solver.url', 'http://localhost:8191/v1')
        self._solver_timeout = config.get('solver.solver_timeout', 60)
        self._user_agent = config.get('general.user_agent', 'HoardCore/5.0')
        self._timeout = config.get('general.timeout_seconds', 30)
        self._max_retries = config.get('general.max_retries', 2)
        self._enable_preflight = config.get('network.enable_preflight', True)
        # SSRF protection (network.ssrf_protection): refuse non-http(s) and
        # non-public fetch targets, and re-validate every redirect hop. Default
        # on; disable only on trusted isolated networks.
        self._ssrf_protected = bool(config.get('network.ssrf_protection', False))
        # Plugin fetchers (hoardcore.fetchers entry points), appended to the
        # strategy fallback chain (G1). Signature: async fn(url) -> tuple or
        # None, or a plain callable returning (text|None, binary|None, ctype,
        # status).
        self.plugin_fetchers: dict[str, Any] = {}

    @staticmethod
    def _is_public_ip(ip_text: str) -> bool:
        """True if *ip_text* is a routable public address.

        `is_global` already excludes private, loopback, link-local (incl. the
        169.254.x.x cloud-metadata ranges), CGNAT 100.64.0.0/10, multicast,
        reserved and unspecified space. IPv4-mapped IPv6 addresses (the
        `::ffff:169.254.169.254` encoding trick) are normalized to their IPv4
        form first, because `is_global` on the mapped form can consult the
        wrong table.
        """
        try:
            ip = ipaddress.ip_address(ip_text)
        except ValueError:
            return False
        if isinstance(ip, ipaddress.IPv6Address) and ip.ipv4_mapped is not None:
            ip = ipaddress.ip_address(ip.ipv4_mapped)
        return ip.is_global

    @staticmethod
    def validate_url_target(url: str) -> bool:
        """SSRF gate for a fetch target: scheme + host allowlist, resolve-then-
        validate, with a blocklist backstop for metadata/DNS-special names.

        Every resolved address must be public (a host that resolves even ONE
        internal address is refused). An unresolvable host is allowed through —
        it cannot reach anything, and the subsequent real fetch will fail on
        its own. DNS rebinding (public IP at check time, internal at fetch
        time) is the residual risk, mitigated here by validating the final URL
        after any redirect chain; full connection-time pinning is a server-mode
        hardening (not needed for the single-user CLI threat model).
        """
        try:
            parsed = urlparse(url)
        except ValueError:
            return False
        if (parsed.scheme or "").lower() not in ("http", "https"):
            return False
        host = (parsed.hostname or "").lower()
        if not host:
            return False
        if host in NetworkFetcher._SSRF_BLOCKED_HOSTS:
            return False
        if host.endswith(".internal") or host.endswith(".local"):
            return False
        try:
            port = parsed.port or (443 if parsed.scheme == "https" else 80)
            infos = socket.getaddrinfo(host, port, proto=socket.IPPROTO_TCP)
        except (socket.gaierror, OSError):
            # Cannot resolve -> cannot reach anything -> allow (fetch fails).
            return True
        except (OverflowError, ValueError):
            return False
        return all(NetworkFetcher._is_public_ip(str(info[4][0])) for info in infos)

    def _parse_cookies(self) -> dict[str, str]:
        cookies = {}
        if not self._cookie_string:
            return cookies
        for part in self._cookie_string.split(';'):
            part = part.strip()
            if '=' in part:
                key, val = part.split('=', 1)
                cookies[key.strip()] = val.strip()
        return cookies

    async def preflight(self, url: str) -> bool:
        if not self._enable_preflight or not self._parse_cookies():
            return True
        if self._ssrf_protected and not self.validate_url_target(url):
            # A SSRF block is a security refusal, NOT a Cloudflare cookie
            # expiry — propagate it under its own marker so the caller reports
            # the right diagnostic (S4).
            logger.warning(f"Preflight: SSRF guard refused target {url}")
            raise RuntimeError("SSRF_BLOCKED")

        try:
            async with aiohttp.ClientSession() as session, session.head(
                url,
                cookies=self._parse_cookies(),
                headers={'User-Agent': self._user_agent},
                timeout=ClientTimeout(total=5),
                allow_redirects=False
            ) as resp:
                blocked = resp.status in (403, 429)
                captcha_redirect = (
                    resp.status in (302, 303)
                    and 'captcha' in resp.headers.get('Location', '').lower()
                )
                return not (blocked or captcha_redirect)
        except Exception as e:
            # Fail CLOSED: if the preflight itself errors (DNS, TLS, transient
            # 5xx), the right default is to abort the fetch, not to proceed.
            # Proceeding can hit Cloudflare with a request the probe could
            # already have rejected, and failing open makes a soft-left case
            # on any server deployment (C1).
            logger.warning(f"Preflight check failed for {url}: {e}")
            return False

    async def _fetch_aiohttp(self, url: str) -> tuple[str | None, bytes | None, str, int | None]:
        """Attempt 1: Standard aiohttp. Returns (text, binary, content_type, status).

        Redirects are followed MANUALLY (`allow_redirects=False` + resolving
        each `Location`) so every hop is re-validated against the SSRF gate
        before it is fetched — a redirect into a private/metadata address is
        refused instead of silently followed (C1).
        """
        cookies = self._parse_cookies()
        headers = {'User-Agent': self._user_agent}
        connector = TCPConnector(force_close=True, enable_cleanup_closed=True, ttl_dns_cache=300)
        timeout = ClientTimeout(total=self._timeout)

        try:
            async with aiohttp.ClientSession(connector=connector, headers=headers) as session:
                current = url
                for _ in range(6):  # initial request + up to 5 redirect hops
                    if self._ssrf_protected and not self.validate_url_target(current):
                        logger.warning(f"aiohttp: SSRF guard refused target {current}")
                        return None, None, 'text/html', 403
                    async with session.get(
                        current, cookies=cookies, timeout=timeout, allow_redirects=False
                    ) as resp:
                        if resp.status in (301, 302, 303, 307, 308):
                            location = resp.headers.get('Location')
                            if not location:
                                return None, None, '', resp.status
                            try:
                                current = str(resp.url.join(URL(location)))
                            except ValueError:
                                return None, None, '', resp.status
                            continue
                        content_type = resp.headers.get('Content-Type', 'text/plain').split(';')[0].strip()
                        if resp.status == 200:
                            if 'text' in content_type:
                                return await resp.text(), None, content_type, resp.status
                            else:
                                return None, await resp.read(), content_type, resp.status
                        elif resp.status == 403:
                            logger.warning("aiohttp: 403 Blocked.")
                            return None, None, content_type, resp.status
                        else:
                            logger.warning(f"aiohttp: Status {resp.status}")
                            return None, None, content_type, resp.status
                logger.warning("aiohttp: redirect limit exceeded")
                return None, None, '', None
        except Exception as e:
            logger.debug(f"aiohttp failed: {e}")
            return None, None, '', None

    async def _fetch_curl_cffi(self, url: str) -> tuple[str | None, bytes | None, str, int | None]:
        if not CURL_AVAILABLE:
            return None, None, '', None

        cookies = self._parse_cookies()
        try:
            def _sync_fetch():
                if curl_requests is None:
                    raise RuntimeError("curl_cffi unexpectedly unavailable")
                resp = curl_requests.get(
                    url,
                    cookies=cookies,
                    headers={'User-Agent': self._user_agent},
                    impersonate="chrome120",
                    timeout=self._timeout,
                )
                return resp
            resp = await asyncio.to_thread(_sync_fetch)

            # curl follows redirects internally; re-validate the FINAL URL so a
            # redirect chain into an internal address is never accepted (C1).
            final_url = str(getattr(resp, 'url', url) or url)
            if self._ssrf_protected and not self.validate_url_target(final_url):
                logger.warning(f"curl_cffi: SSRF guard refused final URL {final_url}")
                return None, None, 'text/html', 403

            content_type = resp.headers.get('Content-Type', 'text/plain').split(';')[0].strip()
            if resp.status_code == 200:
                if 'text' in content_type:
                    return resp.text, None, content_type, resp.status_code
                else:
                    return None, resp.content, content_type, resp.status_code
            elif resp.status_code == 403:
                logger.warning("curl_cffi: 403 Blocked.")
                return None, None, content_type, resp.status_code
            else:
                logger.warning(f"curl_cffi: Status {resp.status_code}")
                return None, None, content_type, resp.status_code
        except Exception as e:
            logger.debug(f"curl_cffi failed: {e}")
            return None, None, '', None

    async def _fetch_flaresolverr(self, url: str) -> tuple[str | None, bytes | None, str, int | None]:
        if not self._solver_enabled:
            return None, None, '', None

        logger.info("FlareSolverr: Solving challenge...")
        payload = {
            "cmd": "request.get",
            "url": url,
            "maxTimeout": self._solver_timeout * 1000,
            "userAgent": self._user_agent,
        }
        if self._cookie_string:
            payload["cookies"] = self._parse_cookies()

        try:
            timeout = ClientTimeout(total=self._solver_timeout + 10)
            async with aiohttp.ClientSession() as session, session.post(
                self._solver_url, json=payload, timeout=timeout
            ) as resp:
                if resp.status != 200:
                    return None, None, '', None
                data = await resp.json()
                if data.get("status") != "ok":
                    return None, None, '', None
                solution = data.get("solution", {})
                if solution.get("status") == 200:
                    # FlareSolverr reports the post-redirect final URL; refuse
                    # a chain that landed on an internal address (C1).
                    final_url = str(solution.get('url') or url)
                    if self._ssrf_protected and not self.validate_url_target(final_url):
                        logger.warning(f"FlareSolverr: SSRF guard refused final URL {final_url}")
                        return None, None, 'text/html', 403
                    content_type = solution.get('headers', {}).get('Content-Type', 'text/html').split(';')[0]
                    response = solution.get('response', '')
                    if 'text' in content_type:
                        return response, None, content_type, int(solution.get('status', 200))
                    else:
                        # FlareSolverr usually returns binary as b64, but we handle text mostly
                        return None, response.encode('utf-8'), content_type, int(solution.get('status', 200))
                return None, None, '', None
        except Exception as e:
            logger.error(f"FlareSolverr failed: {e}")
            return None, None, '', None

    @staticmethod
    def _normalize_fetch(result: tuple[Any, ...]) -> tuple[Any, ...]:
        """Pad a fetch-strategy result to (text, binary, content_type, status).

        Strategy methods return the full 4-tuple; test doubles/older callers
        may return 3-tuples, which are padded with status=None.
        """
        r = tuple(result)
        if len(r) < 4:
            return r + (None,) * (4 - len(r))
        return r[:4]

    async def fetch(self, url: str, strategy: str) -> tuple[str | None, bytes | None, str, int | None]:
        """
        Execute the explicit strategy chain.
        Returns: (text, binary_data, content_type, status) where status is the
        final HTTP status (None if unknown) so the caller can refuse soft-404
        and other error bodies instead of indexing them as content.
        """
        logger.info(f"Fetching {url} with strategy: {strategy}")

        # SSRF entry gate: refuse non-http(s) / non-public targets before any
        # strategy or preflight touches them (C1).
        if self._ssrf_protected and not self.validate_url_target(url):
            logger.warning(f"SSRF guard refused fetch target: {url}")
            raise RuntimeError("SSRF_BLOCKED")

        # Preflight validation
        if (self._enable_preflight and self._parse_cookies()
                and not await self.preflight(url)):
            raise RuntimeError("CF_COOKIE_EXPIRED")

        # Strategy dispatch
        text, binary, ctype, status = None, None, '', None

        if strategy == "fast":
            text, binary, ctype, status = self._normalize_fetch(
                await self._fetch_aiohttp(url))
            if text is not None or binary is not None:
                return text, binary, ctype, status
            raise RuntimeError("FETCH_FAILED")

        elif strategy == "balanced":
            # aiohttp and curl_cffi run CONCURRENTLY: when aiohttp is
            # anti-bot-blocked, we don't wait a full serialized round-trip for
            # curl_cffi to fail too — the first leg that returns content wins.
            # Both legs SSRF-validate independently, so concurrency changes no
            # security semantics (C1). FlareSolverr stays a serialized terminal.
            aio, curl = await asyncio.gather(
                self._fetch_aiohttp(url), self._fetch_curl_cffi(url))
            text, binary, ctype, status = self._pick_fetch(aio, curl)
            if text is not None or binary is not None:
                return text, binary, ctype, status
            raise RuntimeError("FETCH_FAILED")

        elif strategy == "aggressive":
            aio, curl = await asyncio.gather(
                self._fetch_aiohttp(url), self._fetch_curl_cffi(url))
            text, binary, ctype, status = self._pick_fetch(aio, curl)
            if text is not None or binary is not None:
                return text, binary, ctype, status
            text, binary, ctype, status = self._normalize_fetch(
                await self._fetch_flaresolverr(url))
            if text is not None or binary is not None:
                return text, binary, ctype, status
            text, binary, ctype, status = self._normalize_fetch(
                await self._try_plugin_fetchers(url))
            if text is not None or binary is not None:
                return text, binary, ctype, status
            raise RuntimeError("FETCH_FAILED")

        raise RuntimeError("FETCH_FAILED")

    @staticmethod
    def _pick_fetch(aio, curl):
        """Pick the first leg that produced content, preferring aiohttp on a tie
        (aiohttp's returned status is authoritative for soft-404 detection).
        Tolerates legacy 3-tuples from test doubles by padding status=None."""
        aio = NetworkFetcher._pad(aio)
        curl = NetworkFetcher._pad(curl)
        a_text, a_bin, a_ctype, a_status = aio
        c_text, c_bin, c_ctype, c_status = curl
        if (a_text is not None or a_bin is not None) and (
                c_text is not None or c_bin is not None):
            # Both succeeded: prefer the one with a 200; a 404/soft-block from
            # one leg while the other returns 200 means the 200 is the truth.
            if a_status == 200:
                return aio
            if c_status == 200:
                return curl
            return aio  # tie: aiohttp is authoritative
        if a_text is not None or a_bin is not None:
            return aio
        return curl

    @staticmethod
    def _pad(result):
        """Normalize a (text, binary, ctype) 3-tuple to a 4-tuple with
        status=None (legacy test doubles and older legs)."""
        if len(result) == 3:
            return result[0], result[1], result[2], None
        return result

    async def _try_plugin_fetchers(self, url: str) -> tuple[str | None, bytes | None, str, int | None]:
        """Try registered plugin fetchers in order (G1).

        Returns an all-None tuple when every plugin either raised, returned
        None, or returned an all-None tuple — so the strategy chain can move
        on without confusing a None for a successful fetch.
        """
        for name, fn in self.plugin_fetchers.items():
            try:
                out = fn(url)
                if asyncio.iscoroutine(out):
                    out = await out
                if out is None:
                    continue
                text, binary, ctype, status = self._normalize_fetch(out)
                if text is not None or binary is not None:
                    logger.info(f"Plugin fetcher '{name}' succeeded for {url}")
                    return text, binary, ctype, status
            except Exception as e:
                logger.warning(f"Plugin fetcher '{name}' failed for {url}: {e}")
        return None, None, 'text/html', None

# =============================================================================
# 5. DOCUMENT PARSERS (Universal)
# =============================================================================

class DocumentParser:
    """Parses HTML, PDF, DOCX, EPUB into markdown text."""

    # Lazy/optional binary parsers. Imported on first use so that HTML-only
    # scraping works without the heavy PDF/DOCX/EPUB libraries installed.
    _fitz: Any = None
    _docx: Any = None
    _epub: Any = None
    _ocr_engine: Any = None
    _ocr_engine_ready = False
    # Plugin parsers keyed by content type (hoardcore.parsers entry points).
    _plugin_parsers: dict[str, Any] = {}

    @classmethod
    def register_parser(cls, content_type: str, fn: Any) -> None:
        """Register a plugin parser for an (otherwise unknown) content type."""
        cls._plugin_parsers[content_type] = fn

    @classmethod
    def _import_binary_parsers(cls) -> None:
        global FITZ_AVAILABLE, DOCX_AVAILABLE, EPUB_AVAILABLE, RAPIDOCR_AVAILABLE, _BINARY_IMPORTED
        if _BINARY_IMPORTED:
            return
        _BINARY_IMPORTED = True
        try:
            import fitz  # type: ignore[import-not-found]  # PyMuPDF (optional)
            cls._fitz = fitz
            FITZ_AVAILABLE = True
        except ImportError:
            FITZ_AVAILABLE = False
            print("Warning: PyMuPDF (fitz) not installed. PDF parsing disabled.", file=sys.stderr)
        try:
            import docx  # type: ignore[import-not-found]  # python-docx (optional)
            cls._docx = docx
            DOCX_AVAILABLE = True
        except ImportError:
            DOCX_AVAILABLE = False
            print("Warning: python-docx not installed. DOCX parsing disabled.", file=sys.stderr)
        try:
            from ebooklib import epub  # type: ignore[import-not-found]  # ebooklib (optional)
            cls._epub = epub
            EPUB_AVAILABLE = True
        except ImportError:
            EPUB_AVAILABLE = False
            print("Warning: ebooklib not installed. EPUB parsing disabled.", file=sys.stderr)
        try:
            # optional PDF OCR fallback
            from rapidocr_onnxruntime import RapidOCR  # type: ignore[import-not-found]
            cls._RapidOCR = RapidOCR
            RAPIDOCR_AVAILABLE = True
        except ImportError:
            RAPIDOCR_AVAILABLE = False
            print("Warning: rapidocr_onnxruntime not installed. PDF OCR fallback disabled.", file=sys.stderr)

    @classmethod
    def _get_ocr_engine(cls):
        """Return the shared RapidOCR engine (one instance, lazy) or None."""
        if not RAPIDOCR_AVAILABLE:
            return None
        if not cls._ocr_engine_ready:
            try:
                cls._ocr_engine = cls._RapidOCR()
            except Exception as e:
                logger.warning(f"Failed to initialise RapidOCR engine: {e}")
                cls._ocr_engine = None
            finally:
                cls._ocr_engine_ready = True
        return cls._ocr_engine

    @staticmethod
    def _ocr_page(page, dpi: int = 200) -> str:
        """OCR a single rendered page; returns extracted lines or '' if unavailable."""
        engine = DocumentParser._get_ocr_engine()
        if engine is None:
            return ""
        try:
            pix = page.get_pixmap(dpi=dpi)
            result = engine(pix.tobytes("png"))
            items = result[0] if isinstance(result, (list, tuple)) else result
            if not items:
                return ""
            lines = []
            for item in items:
                try:
                    box, text = item[0], item[1]
                    top = min(pt[1] for pt in box)
                    left = min(pt[0] for pt in box)
                except (TypeError, ValueError, IndexError):
                    top, left, text = 0, 0, ""
                text = str(text).strip()
                if text:
                    lines.append((top, left, text))
            lines.sort(key=lambda t: (int(t[0]) // 4, t[1]))
            return "\n".join(t[2] for t in lines)
        except Exception as e:
            logger.warning(f"OCR failed on a page: {e}")
            return ""

    @staticmethod
    async def parse_pdf(binary: bytes) -> tuple[str, dict[str, Any]]:
        """Extract text from PDF using PyMuPDF."""
        DocumentParser._import_binary_parsers()
        if not FITZ_AVAILABLE:
            return "", {"parser": "failed", "error": "PyMuPDF not installed"}
        try:
            doc = DocumentParser._fitz.open(stream=binary, filetype="pdf")
            try:
                text_parts = []
                ocr_pages = 0
                meta = {"page_count": doc.page_count, "parser": "pymupdf"}

                for page_num in range(doc.page_count):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(f"## Page {page_num + 1}\n\n{text.strip()}")
                        continue
                    # Scanned / image-only page: fall back to OCR when available.
                    ocr_text = DocumentParser._ocr_page(page)
                    if ocr_text:
                        text_parts.append(f"## Page {page_num + 1} (ocr)\n\n{ocr_text}")
                        ocr_pages += 1
                    else:
                        text_parts.append(f"## Page {page_num + 1} (ocr: no text extracted)\n\n")

                if ocr_pages:
                    meta["parser"] = "pymupdf+ocr"
                    meta["ocr_pages"] = ocr_pages
                full_text = "\n\n".join(text_parts)
                return full_text, meta
            finally:
                doc.close()
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            return "", {"parser": "failed", "error": str(e)}

    @staticmethod
    async def parse_docx(binary: bytes) -> tuple[str, dict[str, Any]]:
        """Extract text from DOCX."""
        DocumentParser._import_binary_parsers()
        if not DOCX_AVAILABLE:
            return "", {"parser": "failed", "error": "python-docx not installed"}
        try:
            doc = DocumentParser._docx.Document(io.BytesIO(binary))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            return full_text, {"parser": "python-docx", "paragraph_count": len(paragraphs)}
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            return "", {"parser": "failed", "error": str(e)}

    @staticmethod
    async def parse_epub(binary: bytes) -> tuple[str, dict[str, Any]]:
        """Extract text from EPUB."""
        DocumentParser._import_binary_parsers()
        if EPUB_AVAILABLE:
            try:
                # Use ebooklib to parse
                book = DocumentParser._epub.read_epub(io.BytesIO(binary))
                text_parts = []
                for item in book.get_items():
                    if item.get_type() == 9:  # ITEM_DOCUMENT
                        # Extract text from XHTML using regex
                        content = item.get_content().decode('utf-8', errors='ignore')
                        # Simple tag stripping
                        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL)
                        content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL)
                        content = re.sub(r'<[^>]+>', ' ', content)
                        content = re.sub(r'\s+', ' ', content).strip()
                        if content:
                            text_parts.append(content)
                full_text = "\n\n".join(text_parts)
                return full_text, {"parser": "ebooklib", "section_count": len(text_parts)}
            except Exception as e:
                logger.error(f"EPUB parsing failed: {e}")
                # Fall through to zipfile fallback
        else:
            logger.error("EPUB parsing skipped: ebooklib not installed.")
        # Fallback to zipfile (works without ebooklib)
        try:
            with zipfile.ZipFile(io.BytesIO(binary)) as zf:
                text_parts = []
                for name in zf.namelist():
                    if name.endswith('.xhtml') or name.endswith('.html'):
                        content = zf.read(name).decode('utf-8', errors='ignore')
                        content = re.sub(r'<[^>]+>', ' ', content)
                        content = re.sub(r'\s+', ' ', content).strip()
                        if content:
                            text_parts.append(content)
                full_text = "\n\n".join(text_parts)
                return full_text, {"parser": "zipfile_fallback", "section_count": len(text_parts)}
        except Exception as e2:
            logger.error(f"EPUB fallback failed: {e2}")
            return "", {"parser": "failed", "error": str(e2)}

    @staticmethod
    async def clean_html(html: str, url: str) -> tuple[str, dict[str, Any]]:
        """Clean HTML, trafilatura-first with a readability fallback (B8).

        Trafilatura alone produces quality extraction ~90% of the time and is
        cheaper, so readability is only spawned when trafilatura's yield is
        too small (< 100 chars) instead of always running both in parallel.
        """
        def _extract_trafilatura():
            return trafilatura.extract(
                html,
                url=url,
                include_comments=False,
                include_tables=True,
                output_format="markdown"
            ) or ""

        def _extract_readability():
            doc = Document(html)
            return doc.summary()

        try:
            traf_md = await asyncio.to_thread(_extract_trafilatura)
        except Exception as e:
            logger.warning(f"trafilatura extraction failed: {e}")
            traf_md = ""

        # Trafilatura-first: accept its output whenever it cleared the bar.
        if len(traf_md) > 100:
            return traf_md, {"parser": "trafilatura"}

        # Only a weak trafilatura yield pays for the readability pass.
        read_md = ""
        try:
            read_html = await asyncio.to_thread(_extract_readability)
            if read_html:
                # Rough conversion to markdown-ish text
                clean = re.sub(r'<script[^>]*>.*?</script>', '', read_html, flags=re.DOTALL | re.IGNORECASE)
                clean = re.sub(r'<style[^>]*>.*?</style>', '', clean, flags=re.DOTALL | re.IGNORECASE)
                for tag in ['div', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'tr']:
                    clean = re.sub(f'</{tag}>', '\n', clean, flags=re.IGNORECASE)
                clean = re.sub(r'<br\s*/?>', '\n', clean, flags=re.IGNORECASE)
                clean = re.sub(r'<[^>]+>', '', clean)
                import html as html_parser
                clean = html_parser.unescape(clean)
                lines = [line.strip() for line in clean.split('\n') if line.strip()]
                read_md = '\n\n'.join(lines)
        except Exception as e:
            logger.warning(f"readability extraction failed: {e}")

        if len(read_md) > 100:
            return read_md, {"parser": "readability"}

        # Fallback: strip everything
        body_match = re.search(r'<body[^>]*>(.*?)</body>', html, re.DOTALL | re.IGNORECASE)
        body = body_match.group(1) if body_match else html
        body = re.sub(r'<script[^>]*>.*?</script>', '', body, flags=re.DOTALL | re.IGNORECASE)
        body = re.sub(r'<style[^>]*>.*?</style>', '', body, flags=re.DOTALL | re.IGNORECASE)
        body = re.sub(r'<[^>]+>', ' ', body)
        body = re.sub(r'\s+', ' ', body).strip()
        return body, {"parser": "fallback"}

    @staticmethod
    async def parse_text(content: str, url: str) -> tuple[str, dict[str, Any]]:
        """Parse a plain-text / markdown local file as-is (no HTML cleaning).

        Returns the raw text for the chunker (which applies `_tidy_markdown_text`
        itself), and tags the parser as "text" so downstream provenance can tell
        a raw-text source from an HTML/binary one.
        """
        return (content or "").strip(), {"parser": "text"}

    @staticmethod
    async def parse_binary(content_type: str, binary: bytes) -> tuple[str, dict[str, Any]]:
        """Route binary to appropriate parser based on content type."""
        parsers = {
            'application/pdf': DocumentParser.parse_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentParser.parse_docx,
            'application/epub+zip': DocumentParser.parse_epub,
        }

        parser = parsers.get(content_type)
        if parser:
            return await parser(binary)
        else:
            # Unknown binary: consult plugin parsers keyed by content type,
            # then try to decode as text.
            plugin = DocumentParser._plugin_parsers.get(content_type)
            if plugin is not None:
                try:
                    text, meta = plugin(binary)
                    return text, dict(meta or {})
                except Exception as e:
                    logger.warning(f"Plugin parser {content_type!r} failed: {e}")
            try:
                text = binary.decode('utf-8', errors='ignore')
                return text, {"parser": "binary_as_text"}
            except Exception as e:
                # Surface the real error instead of a silent generic one (D1).
                return "", {"parser": "unknown_binary", "error": f"Cannot parse: {e}"}

# =============================================================================
# 6. SEMANTIC CHUNKER
# =============================================================================

# CJK scripts (Chinese/Japanese/Korean incl. full-width/ideographic punctuation)
# tokenize much denser than Latin: ~1-1.5 chars per token vs ~4 for English.
# Used by the chunker so CJK-heavy docs don't get chunks 4x over budget (A14).
_CJK_RE = re.compile(
    r"[\u3400-\u4dbf\u4e00-\u9fff\u3040-\u30ff\u31f0-\u31ff"
    r"\uac00-\ud7af\u3000-\u303f\uff00-\uffef]"
)


class SemanticChunker:
    """Splits text into semantic chunks respecting headers."""

    def __init__(self, config: ConfigManager):
        self.config = config
        self.max_tokens = config.get('chunking.max_tokens', 512)
        self.overlap_tokens = int(config.get('chunking.overlap_tokens', 0) or 0)
        self.strategy = config.get('chunking.strategy', 'heading')
        # Plugin chunkers (hoardcore.chunkers entry points), selected by
        # setting chunking.strategy = "plugin.<name>".
        self._plugins: dict[str, Any] = {}

    def register_plugin(self, name: str, fn: Any) -> None:
        """Register an entry-point chunker under *name* (G1)."""
        self._plugins[name] = fn

    @staticmethod
    def _estimate_tokens(text: str) -> int:
        """CJK-aware token estimate (A14).

        English ~= 4 chars/token; CJK scripts average ~1-1.5 chars per token,
        so a naive `len // 4` badly underestimates Chinese/Japanese/Korean
        chunks. CJK chars count a full token each — conservative, i.e. it
        over-estimates tokens and yields marginally *smaller* chunks, the safe
        failure direction for a chunker. Non-CJK chars keep the ~4 chars/token
        rule (which includes whitespace, as before).
        """
        if not text:
            return 0
        cjk = len(_CJK_RE.findall(text))
        return cjk + (len(text) - cjk) // 4

    @staticmethod
    def _overlap_tail(text: str, overlap_tokens: int) -> str:
        """Trailing slice of *text* worth ~overlap_tokens using the CJK model.

        Walks the text backwards accumulating tokens (1 per CJK char, 0.25 per
        other char) until the budget is spent, then returns that tail. The
        previous chunk's tail becomes the next chunk's opening so sentences
        split across the boundary stay laterally searchable/contextual (A13).
        """
        if overlap_tokens <= 0 or not text:
            return ""
        budget = float(overlap_tokens)
        i = len(text)
        while i > 0 and budget > 0:
            i -= 1
            budget -= 1.0 if _CJK_RE.match(text[i]) else 0.25
        tail = text[i:].lstrip("\n")
        # Don't return a tail that is nothing but text itself re-pasted whole
        # when overlap budget exceeds the chunk length.
        return tail if len(tail) < len(text) else ""

    async def chunk(self, markdown: str, url: str, parser_meta: dict[str, Any]) -> list[Chunk]:
        if not markdown:
            return [Chunk(text="[Empty content]", metadata={"source": url, "empty": True})]

        # Tidy parser-emitted markdown markers (**bold**, *italic*, `code`) out
        # of the canonical text up front, so stored chunks read as clean prose
        # and match the verifier's normalized form (G1).
        markdown = _tidy_markdown_text(markdown)

        # Plugin chunker: chunking.strategy = "plugin.<name>" (G1). On any
        # failure we fall through to the built-in pipeline rather than abort.
        if self.strategy.startswith("plugin."):
            name = self.strategy[len("plugin."):]
            fn = self._plugins.get(name)
            if fn is not None:
                try:
                    out = fn(markdown, url, parser_meta)
                    if asyncio.iscoroutine(out):
                        out = await out
                    if isinstance(out, (list, tuple)) and out and all(
                        isinstance(c, Chunk) for c in out
                    ):
                        return list(out)
                    logger.warning(
                        f"Plugin chunker '{name}' returned non-Chunk output; "
                        "falling back to built-in."
                    )
                except Exception as e:
                    logger.warning(f"Plugin chunker '{name}' failed ({e}); "
                                   "falling back to built-in.")

        # If source is a binary (PDF/DOCX), use paragraph strategy
        if parser_meta.get('parser') in ['pymupdf', 'pymupdf+ocr', 'python-docx', 'ebooklib']:
            strategy = "paragraph"
        else:
            strategy = self.strategy

        lines = markdown.split('\n')
        chunks = []
        current_chunk_lines = []
        header_stack = []  # (depth, title)

        def get_header_path() -> str:
            return " > ".join([h[1] for h in header_stack])

        def flush_chunk(carry_overlap: bool = False):
            nonlocal current_chunk_lines
            if not current_chunk_lines:
                return
            text = '\n'.join(current_chunk_lines).strip()
            if text:
                chunks.append(Chunk(
                    text=text,
                    metadata={
                        "source": url,
                        "header_path": get_header_path(),
                        "parser": parser_meta.get('parser', 'unknown')
                    }
                ))
            if carry_overlap and self.overlap_tokens > 0:
                # Sliding-window overlap (A13): the token-capped flush hands
                # its trailing ~overlap_tokens ahead as the next chunk's head,
                # so a sentence broken at the boundary still begins in-context.
                tail = SemanticChunker._overlap_tail(text, self.overlap_tokens)
                current_chunk_lines = [line for line in tail.split('\n') if line] if tail else []
            else:
                current_chunk_lines = []

        for line in lines:
            stripped = line.strip()
            header_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)

            if header_match and strategy == "heading":
                depth = len(header_match.group(1))
                title = header_match.group(2).strip()

                flush_chunk()

                while header_stack and header_stack[-1][0] >= depth:
                    header_stack.pop()
                header_stack.append((depth, title))
                current_chunk_lines.append(line)
                continue

            current_chunk_lines.append(line)

            if self._estimate_tokens('\n'.join(current_chunk_lines)) > self.max_tokens:
                flush_chunk(carry_overlap=True)

        flush_chunk()

        if not chunks:
            chunks.append(Chunk(
                text=markdown.strip(),
                metadata={"source": url, "header_path": "Root", "parser": parser_meta.get('parser', 'unknown')}
            ))

        return chunks

# =============================================================================
# 7. CRAWLER (Sitemap & Robots)
# =============================================================================

class CrawlerPlanner:
    """Parses robots.txt and sitemaps to discover URLs."""

    def __init__(self, config: ConfigManager):
        self.config = config
        self.respect_robots = config.get('crawler.respect_robots', True)
        self.sitemap_limit = config.get('crawler.sitemap_limit', 500)

    async def get_robots_urls(self, domain: str) -> list[str]:
        """Fetch robots.txt and extract sitemap URLs."""
        if not self.respect_robots:
            return []

        base_url = f"{domain}/robots.txt"
        try:
            async with aiohttp.ClientSession() as session, session.get(
                base_url, timeout=ClientTimeout(total=10)
            ) as resp:
                if resp.status == 200:
                    text = await resp.text()
                    sitemap_urls = re.findall(r'^Sitemap:\s*(.+)$', text, re.MULTILINE | re.IGNORECASE)
                    parsed = [url.strip() for url in sitemap_urls]
                    if parsed:
                        return parsed
                    # robots.txt exists but declares no Sitemap directives;
                    # probe the conventional location before giving up.
                    logger.info(f"{base_url} had no Sitemap: directive; "
                                "falling back to /sitemap.xml.")
        except Exception as e:
            logger.warning(f"Failed to fetch robots.txt: {e}")

        # Fallback to default sitemap location
        return [f"{domain}/sitemap.xml"]

    @staticmethod
    def _extract_locs(xml: str) -> list[str]:
        """Extract <loc> URLs from sitemap XML.

        Namespace-aware via lxml (a guaranteed dependency); falls back to a
        regex scan if the payload cannot be parsed as XML.
        """
        try:
            import lxml.etree as _etree
            root = _etree.fromstring(xml.encode("utf-8"))
            locs: list[str] = []
            for el in root.iter():
                if el.text and el.text.strip() and _etree.QName(el).localname == "loc":
                    locs.append(el.text.strip())
        except Exception:
            locs = [
                loc.strip()
                for loc in re.findall(
                    r"<loc[^>]*>\s*(.*?)\s*</loc>", xml, re.IGNORECASE | re.DOTALL
                )
            ]
        return locs

    async def parse_sitemap(self, sitemap_url: str) -> list[str]:
        """Parse sitemap XML and extract URLs."""
        try:
            async with aiohttp.ClientSession() as session, session.get(
                sitemap_url, timeout=ClientTimeout(total=30)
            ) as resp:
                if resp.status != 200:
                    return []
                xml = await resp.text()
            return list(dict.fromkeys(self._extract_locs(xml)))[:self.sitemap_limit]
        except Exception as e:
            logger.warning(f"Failed to parse sitemap {sitemap_url}: {e}")
            return []

    async def discover_urls(self, url: str) -> list[str]:
        """Discover URLs for a given domain."""
        parsed = urlparse(url)
        domain = f"{parsed.scheme}://{parsed.netloc}"

        sitemap_urls = await self.get_robots_urls(domain)
        all_urls = []
        for sitemap_url in sitemap_urls:
            urls = await self.parse_sitemap(sitemap_url)
            all_urls.extend(urls)

        # Deduplicate
        return list(dict.fromkeys(all_urls))

# =============================================================================
# 8. WEB DISCOVERY (feed the crawler from a live search query)
# =============================================================================

@dataclass
class SearchResult:
    title: str
    url: str
    snippet: str = ""


class WebSearchProvider:
    """Discovers URLs from a live web query.

    Tries providers in order (DuckDuckGo HTML -> Mojeek HTML), each driven
    through the SAME resilient fetch chain as the crawler (aiohttp ->
    curl_cffi -> FlareSolverr), with bounded retry + exponential backoff on
    transient failures. Returns candidate SearchResults for the orchestrator
    to feed into _ingest_many.
    """

    def __init__(self, config: ConfigManager, fetcher: NetworkFetcher):
        self.config = config
        self.fetcher = fetcher
        self.max_retries = int(config.get('discovery.max_retries', 2))
        self.backoff_base = float(config.get('discovery.backoff_seconds', 1.5))
        self.plugin_providers: dict[str, Any] = {}

    @staticmethod
    def _clean_title(raw: str) -> str:
        return re.sub(r"<[^>]+>", "", raw).strip()

    async def _fetch_with_backoff(self, url: str, strategy: str) -> str | None:
        """Fetch a search page, retrying transient failures with backoff."""
        last_exc: Exception | None = None
        for attempt in range(self.max_retries + 1):
            try:
                data = await self.fetcher.fetch(url, strategy)
                text = data[0] if data else None
                if text:
                    return text
                raise RuntimeError("empty search response")
            except Exception as e:  # transient: rate-limit, 5xx, timeout
                last_exc = e
                if attempt < self.max_retries:
                    delay = self.backoff_base * (2 ** attempt)
                    logger.warning(
                        f"Discovery attempt {attempt + 1}/{self.max_retries + 1} "
                        f"failed ({e}); retrying in {delay:.1f}s"
                    )
                    await asyncio.sleep(delay)
        logger.error(f"Discovery search failed after retries: {last_exc}")
        return None

    @staticmethod
    def _parse_duckduckgo(text: str, max_results: int) -> list[SearchResult]:
        results: list[SearchResult] = []
        for m in re.finditer(
            r'class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>',
            text, re.IGNORECASE | re.DOTALL
        ):
            href, title_raw = m.group(1), m.group(2)
            target = href
            if "uddg=" in href:
                target = unquote(href.split("uddg=", 1)[1].split("&", 1)[0])
            if not target.startswith("http"):
                continue
            if is_ad_tracking_url(target):
                continue  # ad-redirect/tracker URLs are not search content
            results.append(SearchResult(
                title=WebSearchProvider._clean_title(title_raw),
                url=target
            ))
            if len(results) >= max_results:
                break
        return results

    @staticmethod
    def _parse_mojeek(text: str, max_results: int) -> list[SearchResult]:
        results: list[SearchResult] = []
        for m in re.finditer(
            r'<a class="ob"[^>]*href="([^"]+)"[^>]*>(.*?)</a>',
            text, re.IGNORECASE | re.DOTALL
        ):
            href, title_raw = m.group(1), m.group(2)
            if not href.startswith("http"):
                continue
            if is_ad_tracking_url(href):
                continue  # ad-redirect/tracker URLs are not search content
            results.append(SearchResult(
                title=WebSearchProvider._clean_title(title_raw),
                url=href
            ))
            if len(results) >= max_results:
                break
        return results

    async def _try_provider(self, url: str, strategy: str, max_results: int,
                            parser) -> list[SearchResult]:
        text = await self._fetch_with_backoff(url, strategy)
        if not text:
            return []
        return parser(text, max_results)

    async def search(self, query: str, max_results: int = 10,
                     strategy: str = "aggressive") -> list[SearchResult]:
        q = re.sub(r"\s+", "+", query.strip())
        # (label, url, parser) ordered by preference; later entries are fallbacks.
        providers = [
            ("duckduckgo", f"https://html.duckduckgo.com/html/?q={q}",
             self._parse_duckduckgo),
            ("mojeek", f"https://www.mojeek.com/search?q={q}",
             self._parse_mojeek),
        ]

        last_results: list[SearchResult] = []
        for label, url, parser in providers:
            results = await self._try_provider(url, strategy, max_results, parser)
            if results:
                logger.info(f"Discovery provider '{label}' returned {len(results)} results.")
                return results
            logger.warning(f"Discovery provider '{label}' returned nothing; trying fallback.")
            last_results = results

        # Plugin-provided discovery backends form the tail of the fallback
        # chain (built-ins first). A plugin provider is a callable
        # `search(query, max_results) -> list[SearchResult]`, optionally async.
        for label, backend in self.plugin_providers.items():
            try:
                res = backend(query, max_results)
                if inspect.isawaitable(res):
                    res = await res
                results = list(res or [])
                if results:
                    logger.info(
                        f"Discovery plugin provider '{label}' returned "
                        f"{len(results)} results."
                    )
                    return results
                logger.warning(
                    f"Discovery plugin provider '{label}' returned nothing; "
                    "trying fallback."
                )
            except Exception as e:
                logger.warning(
                    f"Discovery plugin provider '{label}' failed ({e}); "
                    "trying fallback."
                )

        return last_results


# =============================================================================
# 8.5 EVENT BUS & PLUGIN SYSTEM
# =============================================================================


class EventBus:
    """Tiny publish/subscribe bus for lifecycle hooks (G2).

    Handlers are best-effort: an exception in one never propagates into the
    ingest/search path that triggered it. Emitted events today:
      'document.ingested'   (url, version, chunks[, parallel])
      'chunk.embedded'      (chunk_hash, vector_dim)
      'discovery.completed' (query, urls)
      'search.completed'    (query, n_results[, domain])
    """

    def __init__(self) -> None:
        self._handlers: dict[str, list[Any]] = {}

    def on(self, event: str, handler: Any | None = None):
        """Register *handler* for *event*; also usable as a decorator:
        ``@bus.on('document.ingested')``."""
        if handler is None:
            def _decorate(fn: Any) -> Any:
                self._handlers.setdefault(event, []).append(fn)
                return fn
            return _decorate
        self._handlers.setdefault(event, []).append(handler)
        return handler

    def emit(self, event: str, **kwargs: Any) -> None:
        for fn in self._handlers.get(event, []):
            try:
                fn(**kwargs)
            except Exception as e:
                logger.warning(f"Event handler '{event}' raised: {e}")

    def handlers(self, event: str) -> list[Any]:
        return list(self._handlers.get(event, []))


class PluginManager:
    """Discovers third-party plugins via importlib.metadata entry points (G1).

    One group per extension point, following packaging's "Creating and
    discovering plugins" model (the same architecture virtualenv uses):

      hoardcore.parsers    -> binary parsers, keyed by content type
      hoardcore.fetchers   -> extra fetch strategies appended to the chain
      hoardcore.providers  -> extra discovery backends (fallback chain)
      hoardcore.chunkers   -> custom chunkers, selected by chunking.strategy

    An installed distribution advertises, e.g.::

        [project.entry-points."hoardcore.parsers"]
        myparser = "mylib:parser_entry"

    A broken plugin can never kill startup: every load()/registration is
    wrapped and logged. Entry points only resolve for *installed*
    distributions, so built-ins remain the fallback for repo-source runs.
    """

    GROUPS = ("hoardcore.parsers", "hoardcore.fetchers",
              "hoardcore.providers", "hoardcore.chunkers")

    def __init__(self, config: ConfigManager):
        self.config = config
        self.enabled = bool(config.get('plugins.enabled', True))
        self.parsers: dict[str, Any] = {}
        self.fetchers: dict[str, Any] = {}
        self.providers: dict[str, Any] = {}
        self.chunkers: dict[str, Any] = {}
        self._discovered = False

    def discover(self) -> PluginManager:
        if self._discovered:
            return self
        self._discovered = True
        if not self.enabled:
            return self
        targets = {
            "hoardcore.parsers": self.parsers,
            "hoardcore.fetchers": self.fetchers,
            "hoardcore.providers": self.providers,
            "hoardcore.chunkers": self.chunkers,
        }
        try:
            eps = importlib.metadata.entry_points()
            for group in self.GROUPS:
                try:
                    selected = eps.select(group=group)
                except (AttributeError, TypeError):
                    # Python < 3.10 returned a dict; keep the project 3.10+.
                    selected = eps.get(group, [])  # pyright: ignore[reportAttributeAccessIssue]
                for ep in selected:
                    try:
                        targets[group][ep.name] = ep.load()
                    except Exception as e:
                        logger.warning(
                            f"Plugin {ep.name} ({group}) failed to load: {e}"
                        )
        except Exception as e:
            logger.warning(f"Plugin discovery failed: {e}")
        total = sum(len(t) for t in targets.values())
        if total:
            counts = {g.split('.')[-1]: len(t)
                      for g, t in targets.items() if t}
            logger.info(f"Loaded plugin(s): {counts}")
        return self


# =============================================================================
# 9. MAIN ORCHESTRATOR
# =============================================================================

# Local ingestion (--action local): the only file types read from
# storage.local_dir, and only from that directory (containment enforced).
_LOCAL_EXTENSIONS = frozenset({".pdf", ".docx", ".epub", ".html", ".htm",
                               ".txt", ".md"})

class HoardCore:
    """Main entry point for scraping, crawling, and searching."""

    def __init__(self, vault_name: str | list[str] | tuple[str, ...] | None = None):
        self.config = ConfigManager()
        self.bus = EventBus()
        self.plugins = PluginManager(self.config).discover()

        # Cross-vault read: `vault_name` may be a comma/space list
        # ("career,negros_ai_jobs"). The FIRST vault is the primary
        # (self.vault): the only one that receives new ingest/discover/
        # research-write traffic. The rest are read-only companions used by
        # search/verify/audit (recall reads across all of them). A single
        # name or None behaves exactly as before.
        if isinstance(vault_name, str):
            names = [n for n in re.split(r"[,\s]+", vault_name.strip()) if n]
        elif isinstance(vault_name, (list, tuple)):
            names = [str(n).strip() for n in vault_name if str(n).strip()]
        else:
            names = []
        self.vault = VaultManager(self.config, names[0] if names else None,
                                  event_bus=self.bus)
        self.vaults = [self.vault] + [
            VaultManager(self.config, n, event_bus=self.bus) for n in names[1:]
        ]
        self.vault_name = self.vault.vault_name
        self.fetcher = NetworkFetcher(self.config)
        self.parser = DocumentParser()
        self.chunker = SemanticChunker(self.config)
        self.crawler = CrawlerPlanner(self.config)
        self.discovery = WebSearchProvider(self.config, self.fetcher)
        self._register_plugins()
        self.save_binary = self.config.get('storage.save_binary', True)
        self.save_raw_html = self.config.get('storage.save_raw_html', False)

    def _register_plugins(self) -> None:
        """Wire discovered entry-point plugins into the pipeline (G1).

        - parsers   -> DocumentParser keyed by content type
        - chunkers  -> SemanticChunker, selectable via chunking.strategy
        - providers -> WebSearchProvider fallback chain
        - fetchers  -> NetworkFetcher fallback chain
        All religiously failure-tolerant: a bad plugin degrades nothing.
        """
        # Parsers: accept either dict {content_type: fn} or a callable with a
        # .content_type attribute.
        for name, obj in self.plugins.parsers.items():
            try:
                if isinstance(obj, dict):
                    for ct, fn in obj.items():
                        DocumentParser.register_parser(str(ct), fn)
                elif callable(obj):
                    ct = getattr(obj, "content_type", None)
                    if ct:
                        DocumentParser.register_parser(str(ct), obj)
                else:
                    logger.warning(
                        f"Parser plugin {name}: expected a dict or a callable "
                        "with .content_type"
                    )
            except Exception as e:
                logger.warning(f"Parser plugin {name} failed to register: {e}")

        for name, fn in self.plugins.chunkers.items():
            try:
                self.chunker.register_plugin(name, fn)
            except Exception as e:
                logger.warning(f"Chunker plugin {name} failed to register: {e}")

        self.discovery.plugin_providers.update(self.plugins.providers)
        self.fetcher.plugin_fetchers.update(self.plugins.fetchers)

    # --- Local directory ingestion (--action local) -------------------------
    # Reads ONLY from storage.local_dir (default local_inputs/); any path that
    # resolves outside that directory is refused (mirrors write_artifact's
    # traversal guard). No network, no SSRF, no HTTP cache-TTL: freshness is
    # content-based (a file whose extracted content is unchanged since the last
    # ingest is skipped unless --force).

    def _local_dir_root(self) -> str:
        """Absolute root of the local ingestion directory."""
        return os.path.abspath(str(self.config.get('storage.local_dir', 'local_inputs')))

    def _local_url(self, relpath: str) -> str:
        """Synthetic provenance URL for a local file. The host is fixed to
        `local` so `documents.domain` is always "local" and the chunk carries a
        recognizable, human-readable address (`local://local/papers/x.pdf`)."""
        return f"local://local/{relpath}"

    def _local_guard(self, relpath: str) -> str:
        """Resolve `relpath` inside storage.local_dir, refusing escapes."""
        root = self._local_dir_root()
        abs_path = os.path.abspath(os.path.join(root, relpath))
        if abs_path != root and os.path.commonpath([root, abs_path]) != root:
            raise RuntimeError(
                f"LOCAL_PATH_REFUSED: {relpath!r} escapes "
                f"{self.config.get('storage.local_dir', 'local_inputs')}"
            )
        return abs_path

    def _local_files(self, relpath: str = ".") -> list[str]:
        """Every supported file under `relpath` of local_dir (sorted, recursive)."""
        root = self._local_dir_root()
        base = self._local_guard(relpath)
        if os.path.isfile(base):
            return [os.path.relpath(base, root)]
        if not os.path.isdir(base):
            raise RuntimeError(f"LOCAL_PATH_NOT_FOUND: {relpath!r}")
        files: list[str] = []
        for dirpath, _dirnames, filenames in os.walk(base):
            for name in sorted(filenames):
                if os.path.splitext(name)[1].lower() in _LOCAL_EXTENSIONS:
                    files.append(os.path.relpath(os.path.join(dirpath, name), root))
        return files

    async def _process_local(self, relpath: str, force_refresh: bool = False) -> tuple[list[Chunk], dict[str, Any]]:
        """Index one local file into the primary vault, bypassing the network
        pipeline entirely (no fetch chain / SSRF / HTTP cache). Freshness is
        content-based: a file whose extracted content is unchanged since the
        last ingest is skipped unless `force_refresh`. Returns (chunks, meta);
        meta may carry `cached`, `junk`, or `error`."""
        try:
            abs_path = self._local_guard(relpath)
        except RuntimeError as e:
            return [], {"error": str(e)}
        if not os.path.isfile(abs_path):
            return [], {"error": f"LOCAL_FILE_NOT_FOUND: {relpath!r}"}
        ext = os.path.splitext(abs_path)[1].lower()
        if ext not in _LOCAL_EXTENSIONS:
            return [], {"error": f"LOCAL_UNSUPPORTED_EXT: {relpath!r}"}

        url = self._local_url(relpath)
        try:
            with open(abs_path, "rb") as fh:
                data = fh.read()
        except OSError as e:
            return [], {"error": f"LOCAL_READ_FAILED: {relpath!r}: {e}"}

        text = data.decode("utf-8", errors="ignore")
        parser_meta: dict[str, Any] = {
            "source_type": "local", "local_path": abs_path, "url": url,
        }
        try:
            if ext in (".html", ".htm"):
                markdown, html_meta = await DocumentParser.clean_html(text, url)
                parser_meta.update(html_meta)
                parser_meta["content_type"] = "text/html"
            elif ext == ".pdf":
                markdown, pdf_meta = await DocumentParser.parse_pdf(data)
                parser_meta.update(pdf_meta)
                parser_meta["content_type"] = "application/pdf"
            elif ext == ".docx":
                markdown, doc_meta = await DocumentParser.parse_docx(data)
                parser_meta.update(doc_meta)
                parser_meta["content_type"] = (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                )
            elif ext == ".epub":
                markdown, epub_meta = await DocumentParser.parse_epub(data)
                parser_meta.update(epub_meta)
                parser_meta["content_type"] = "application/epub+zip"
            else:  # .txt / .md
                markdown, _t = await DocumentParser.parse_text(text, url)
                parser_meta["content_type"] = "text/plain"
        except Exception as e:
            return [], {"error": f"LOCAL_PARSE_FAILED: {relpath!r}: {e}"}

        if not (markdown or "").strip():
            markdown = "[No extractable content found]"

        # Boilerplate detection applies to real HTML; raw text/markdown is
        # ingested as-authored (a short note is still evidence).
        if ext in (".html", ".htm"):
            junk = self._detect_junk(markdown, text, parser_meta, 0.5)
            if junk:
                logger.info(f"local: skipping {relpath} (junk: {junk}).")
                return [], {"junk": True, "junk_reason": junk}

        chunks = await self.chunker.chunk(markdown, url, parser_meta)
        # Content-hash basis matches the rest of the pipeline: a blake2b over the
        # joined chunk texts. Identical extracted content -> identical hash ->
        # no new WORM version is written unless --force.
        content_hash = hashlib.blake2b(
            "\n".join(c.text for c in chunks).encode("utf-8"),
            digest_size=32,
        ).hexdigest()
        if not force_refresh and content_hash == self.vault.latest_content_hash(url):
            logger.info(f"local: {relpath} unchanged (content-hash); skipping.")
            return [], {"cached": True, "url": url, "content_hash": content_hash}
        self.vault.ingest_chunks_parallel(url, chunks, parser_meta)
        logger.info(f"local: indexed {relpath} ({len(chunks)} chunks, ext={ext}).")
        return chunks, {**parser_meta, "content_hash": content_hash}

    async def local_ingest(self, relpath: str = ".", force_refresh: bool = False) -> list[dict[str, Any]]:
        """Index every supported local file under `relpath` of local_dir."""
        try:
            files = self._local_files(relpath)
        except RuntimeError as e:
            return [{"text": str(e), "metadata": {"error": str(e)}}]
        results: list[dict[str, Any]] = []
        for rel in files:
            chunks, meta = await self._process_local(rel, force_refresh)
            if meta.get("cached") or meta.get("junk"):
                continue
            if meta.get("error"):
                results.append({"text": str(meta["error"]), "metadata": meta})
            else:
                results.extend(c.to_dict() for c in chunks)
        return results

    def scan_local(self, relpath: str = ".") -> list[dict[str, Any]]:
        """Read-only listing of supported files under `relpath` of local_dir."""
        root = self._local_dir_root()
        try:
            files = self._local_files(relpath)
        except RuntimeError as e:
            return [{"error": str(e)}]
        entries: list[dict[str, Any]] = []
        for rel in files:
            try:
                st = os.stat(os.path.join(root, rel))
            except OSError:
                continue
            entries.append({
                "path": rel,
                "bytes": st.st_size,
                "modified": st.st_mtime,
            })
        return entries

    @property
    def artifacts_dir(self) -> str:
        """Directory for finished research deliverables (reports, syntheses, audits)."""
        return self.vault.artifacts_dir

    def _artifact_day_subdir(self) -> str:
        """Day-scoped artifacts subdirectory, e.g. artifacts/2026-08-10/.

        Used when storage.artifacts_by_day is enabled so finished deliverables
        are grouped by the day they were written instead of piling up flat in
        the artifacts root.
        """
        day = time.strftime("%Y-%m-%d")
        path = os.path.join(self.artifacts_dir, day)
        os.makedirs(path, exist_ok=True)
        return path

    def resolve_artifact_out(self, out_path: str | None) -> str:
        """Map an artifact target to its on-disk path honoring artifacts_by_day.

        When storage.artifacts_by_day is true, any path that targets the
        artifacts directory is re-scoped into the current-day subfolder; paths
        elsewhere (e.g. /tmp/scratch.md) are left untouched so callers keep
        full control of out-of-vault writes.
        """
        if not self.config.get('storage.artifacts_by_day', True):
            if out_path is None:
                base = os.path.join(self.artifacts_dir, "grounding_context.md")
                if not os.path.exists(base):
                    return base
                i = 2
                while os.path.exists(os.path.join(self.artifacts_dir,
                                                  f"grounding_context_{i}.md")):
                    i += 1
                return os.path.join(self.artifacts_dir, f"grounding_context_{i}.md")
            return out_path
        if out_path is None:
            # Default deliverables share one name; if today's file already
            # exists, suffix it so a second research run never clobbers the
            # first session's grounding context. Grounding contexts are working
            # instruments (not finished deliverables), so they live in a
            # subfolder of the day directory and never pollute the day folder
            # of syntheses/audits.
            sub = self.config.get('storage.grounding_subdir', 'grounding')
            gdir = os.path.join(self._artifact_day_subdir(), sub)
            base = os.path.join(gdir, "grounding_context.md")
            if not os.path.exists(base):
                return base
            i = 2
            while os.path.exists(os.path.join(gdir, f"grounding_context_{i}.md")):
                i += 1
            return os.path.join(gdir, f"grounding_context_{i}.md")
        artifacts_root = os.path.abspath(self.artifacts_dir) + os.sep
        if os.path.abspath(out_path).startswith(artifacts_root):
            return os.path.join(self._artifact_day_subdir(), os.path.basename(out_path))
        return out_path

    def write_artifact(self, filename: str, content: str) -> str:
        """Write a research deliverable into the artifacts directory.

        Returns the absolute path written. Raises if the filename would escape
        the artifacts directory. Honors storage.artifacts_by_day: files land in
        an artifacts/YYYY-MM-DD/ subfolder.
        """
        if os.path.basename(filename) != filename:
            raise ValueError(f"artifact filename must be a bare name, got {filename!r}")
        target = self.resolve_artifact_out(os.path.join(self.artifacts_dir, filename))
        os.makedirs(os.path.dirname(os.path.abspath(target)), exist_ok=True)
        with open(target, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"Artifact written -> {target}")
        return target

    def organize_artifacts_by_day(self) -> list[str]:
        """Move flat artifacts/ files into per-day subfolders by mtime.

        One-time housekeeping for the artifacts-by-day feature: any deliverable
        living directly in the artifacts root is re-homed to
        artifacts/YYYY-MM-DD/ based on its file mtime. Only files (not
        directories) at the artifacts root are touched, and any day-subfolder
        names already present are skipped. Returns the list of new paths.

        Run when storage.artifacts_by_day is true; no-op otherwise.
        """
        if not self.config.get('storage.artifacts_by_day', True):
            return []
        root = self.artifacts_dir
        migrated: list[str] = []
        for name in sorted(os.listdir(root)):
            src = os.path.join(root, name)
            if not os.path.isfile(src):
                continue
            day = time.strftime("%Y-%m-%d", time.localtime(os.path.getmtime(src)))
            dest_dir = os.path.join(root, day)
            dest = os.path.join(dest_dir, name)
            if os.path.abspath(dest) == os.path.abspath(src):
                continue
            os.makedirs(dest_dir, exist_ok=True)
            if os.path.exists(dest):
                logger.warning(f"Migration: {dest} exists; skipping {src}.")
                continue
            os.replace(src, dest)
            migrated.append(dest)
            logger.info(f"Artifact organized -> {dest}")
        return migrated

    @staticmethod
    def citation_list(sources: list[str] | dict[str, str]) -> str:
        """Render a **Source Links / Citations** block for an artifact.

        Accepts either a list of source URLs or a ``{label: url}`` mapping. The
        block closes every artifact whose provenance tags use the ``[V#N]``
        convention (each number N resolves to the Nth entry here), so
        ``[V#3]`` -> ``[#3] <label> — <url>``.

        Returns a ready-to-append markdown string (leading + trailing newline
        included). Labels default to the bare URL when only a list is given.
        """
        if isinstance(sources, dict):
            items: list[tuple[str, str]] = list(sources.items())
        else:
            items = [(u, u) for u in sources]
        lines = ["\n## Source Links / Citations", ""]
        for i, (label, url) in enumerate(items, 1):
            lines.append(f"[#{i}] {label} — {url}")
        lines.append("")
        return "\n".join(lines)

    @staticmethod
    def _drop_low_confidence(chunks: list[Chunk]) -> list[Chunk]:
        """EMIT hygiene: drop confidence='low' hits unless they are all we have
        (a lone low hit is still better than nothing). Keeps at least one chunk
        per distinct source, so a low-banded but authoritative primary source
        never vanishes entirely from the deliverable (stress-test regression:
        the IEA primary source was filtered out while blogs survived)."""
        strong = [c for c in chunks
                  if c.metadata.get('confidence') != 'low']
        if not strong:
            return chunks
        kept_sources = {
            c.metadata.get('source_url') or c.metadata.get('source') or '?'
            for c in strong
        }
        result = list(strong)
        for c in chunks:
            if c.metadata.get('confidence') != 'low':
                continue
            src = c.metadata.get('source_url') or c.metadata.get('source') or '?'
            if src not in kept_sources:
                result.append(c)
                kept_sources.add(src)
        return result

    def _search_across_vaults(self, query: str, limit: int = 20,
                             domain: str | None = None,
                             hybrid: bool | None = None,
                             max_per_source: int = 0) -> list[Chunk]:
        """Cross-vault hybrid recall over every named vault.

        Each vault's own `search_vault` already returns an RRF-fused ranking;
        those rankings are fused again via Reciprocal Rank Fusion so every vault
        contributes `1 / (k + rank + 1)`. Exact-duplicate text appearing in more
        than one vault is deduped (first vault wins) while its scores accumulate,
        confidence bands are recomputed set-relative over the fused set, and
        `max_per_source` is enforced globally (the top-ranked hit is always
        admitted). Every chunk is stamped with the vault that sourced it so
        provenance can be traced per `[V#N]`. A single-vault run short-circuits
        to the original `search_vault` path unchanged."""
        if len(self.vaults) <= 1:
            return self.vault.search_vault(query, limit, domain=domain,
                                           hybrid=hybrid, max_per_source=max_per_source)
        pool = max(limit, int(self.config.get('indexer.search_limit', 20)))
        k = 60  # RRF constant (matches _search_hybrid)
        score: dict[str, float] = {}
        results: dict[str, Chunk] = {}
        for vault in self.vaults:
            per = vault.search_vault(query, limit=pool, domain=domain,
                                     hybrid=hybrid, max_per_source=0)
            for rank, c in enumerate(per):
                text_norm = normalize_claim(c.text)
                if not text_norm:
                    continue
                key = hashlib.blake2b(text_norm.encode("utf-8"),
                                      digest_size=16).hexdigest()
                score[key] = score.get(key, 0.0) + 1.0 / (k + rank + 1)
                if key not in results:
                    c.metadata.setdefault("vault", vault.vault_name or "(default)")
                    results[key] = c

        keys = sorted(score, key=lambda k: score[k], reverse=True)
        if max_per_source > 0:
            selected: list[str] = []
            per_src: dict[str, int] = {}
            for key in keys:
                if len(selected) >= limit:
                    break
                src = results[key].metadata.get('source_url') or '?'
                if per_src.get(src, 0) >= max_per_source:
                    continue
                per_src[src] = per_src.get(src, 0) + 1
                selected.append(key)
            keys = selected[:limit]
        else:
            keys = keys[:limit]

        chunks = []
        for key in keys:
            c = results[key]
            c.metadata['hybrid_score'] = score[key]
            c.metadata['retrieval'] = 'cross_vault'
            chunks.append(c)

        # Set-relative confidence bands over the fused set (mirrors the
        # relative mode of _search_hybrid: top ~20% can be high, bottom ~half
        # hugging the coincidence floor is low, the middle stays medium).
        n = len(chunks)
        if n:
            top, tail = score[keys[0]], score[keys[-1]]
            spread = max(top - tail, 1e-9)
            n_high = max(1, (n * 2) // 10 + (1 if n % 10 >= 5 else 0))
            n_low = max(1, n // 2)
            for idx, key in enumerate(keys):
                rel = (score[key] - tail) / spread
                conf = "medium"
                if rel >= 0.66 and idx < n_high:
                    conf = "high"
                elif rel <= 0.10 and idx >= n_low:
                    conf = "low"
                results[key].metadata['confidence'] = conf
        return chunks

    async def research(self, question: str, out_path: str | None = None,
                       discover: int = 5, recall: int = 6,
                       strategy: str | None = None,
                       answer_first: bool | None = None,
                       keep_low: bool = False) -> str | None:
        """Agentic research workflow: DISCOVER -> INGEST -> RECALL -> EMIT.

        Live web-searches the question (via the configured discovery provider),
        ingests the top-ranked sources into the vault, hybrid-retrieves the best
        chunks, and writes a grounding-context file. Returns the path written,
        or None if nothing was retrieved.

        strategy: "fast", "balanced", or "aggressive"; defaults to
            network.default_strategy from config. Controls the fetch chain used
            for both discovery and ingestion (e.g. "aggressive" enables the
            FlareSolverr path for anti-bot-protected sources).

        answer_first: when True (config research.answer_first, default), the
            existing vault is queried BEFORE any discovery; if a high-confidence
            memory hit answers the question, live discovery is skipped entirely
            (Adaptive-RAG-style routing — most recurring questions need no new
            retrieval). Pass False to always run live discovery.
        """
        if strategy is None:
            strategy = str(self.config.get("network.default_strategy", "aggressive"))

        if answer_first is None:
            answer_first = bool(self.config.get('research.answer_first', True))

        # [0/ANSWER-FIRST] memory check before touching the web.
        # max_per_source keeps recall source-diverse (DeepResearch wants breadth,
        # not one rich page crowding out the rest); 0 disables it.
        max_per_source = int(self.config.get('research.max_per_source', 2) or 0)
        raw_memory: list[Chunk] = self._search_across_vaults(
            question, limit=recall, domain=None, hybrid=True,
            max_per_source=max_per_source)
        memory_chunks = (raw_memory if keep_low
                         else self._drop_low_confidence(raw_memory))
        answered = (answer_first and memory_chunks and any(
            c.metadata.get('confidence') == 'high' for c in memory_chunks))

        if answered:
            print("\n[0/ANSWER-FIRST] memory answers the question; skipping DISCOVER", flush=True)
            chunks: list[Chunk] = memory_chunks
            dropped_low: list[str] = ([] if keep_low else [
                c.metadata.get('source_url', '?')
                for c in raw_memory
                if c.metadata.get('confidence') == 'low' and c not in memory_chunks
            ])
        else:
            # --discover N controls the hunt breadth; --discover 0 is an
            # explicit "recall-only" run (never touch the web). Omitted
            # (None) falls back to the config default live discovery.
            if discover is not None and discover <= 0:
                logger.info("Skipping DISCOVER (--discover 0 = recall-only).")
            else:
                print(f"\n[1/DISCOVER] searching web for: {question!r} (top {discover})", flush=True)
                await self._discover_and_ingest(question, discover, strategy, force_refresh=False)

            print(f"\n[2/RECALL] hybrid-retrieving top {recall} chunks", flush=True)
            raw_chunks = self._search_across_vaults(
                question, limit=recall, domain=None, hybrid=True,
                max_per_source=max_per_source)
            chunks = raw_chunks if keep_low else self._drop_low_confidence(raw_chunks)
            # Track how many low-confidence hits filter_low actually removed, so
            # the grounding file explains the count honestly (chunks retained
            # for source preservation are NOT counted as dropped). No-op when
            # keep_low is set.
            dropped_low = ([] if keep_low else [
                c.metadata.get('source_url', '?')
                for c in raw_chunks
                if c.metadata.get('confidence') == 'low' and c not in chunks
            ])
            if not chunks:
                print("  -> no chunks retrieved")
                return None

        out_path = self.resolve_artifact_out(out_path)
        os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)

        print("\n[3/EMIT] writing grounding context", flush=True)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(f"# Grounding Context\n## Question\n{question}\n\n")
            if answered:
                f.write("> Answer-first recall: live DISCOVER was skipped "
                        "(existing high-confidence memory hit).\n\n")
            f.write(f"## Retrieved sources ({len(chunks)})\n\n")
            if dropped_low:
                f.write(f"> Note: `filter_low` (research.filter_low) dropped "
                        f"{len(dropped_low)} low-confidence hit(s) from the raw "
                        f"recall before EMIT; they are excluded below.\n\n")
            seen: set = set()
            for i, c in enumerate(chunks, 1):
                src = c.metadata.get("source_url", "?")
                seen.add(src)
                score = c.metadata.get("hybrid_score")
                score_s = f"{score:.4f}" if score is not None else "n/a"
                cid = c.metadata.get("chunk_id")
                cid_s = f" | chunk {cid}" if cid is not None else ""
                vault = c.metadata.get("vault", "")
                vault_s = f" | vault {vault}" if vault else ""
                f.write(f"### [{i}] {src}  (score {score_s} | {c.metadata.get('confidence', 'n/a')}{cid_s}{vault_s})\n")
                f.write(f"{c.text}\n\n")
            f.write(f"## Run budget\n- `--discover` {discover if discover is not None else '(config)'} | `--recall` {recall}\n\n")
            f.write(f"## Distinct sources in recall: {len(seen)}\n")
            for s in sorted(seen):
                f.write(f" - {s}\n")
            if dropped_low:
                f.write("## Low-confidence hits filtered at EMIT\n")
                for s in sorted(set(dropped_low)):
                    f.write(f" - {s}\n")
            f.write(self.citation_list(sorted(seen)))

        abs_path = os.path.abspath(out_path)
        print(f"\n=== DONE. {len(chunks)} chunks, {len(seen)} sources -> {abs_path}")
        return abs_path

    def verify_claim(self, claim: str) -> str:
        """Programmatic adversarial-audit: confirm a claim against the vault(s).

        Checks the raw stored chunk text for the claim and reports whether it is
        supported. Returns one of:
          "verified"   - the claim (or a distinctive normalized substring of it)
                         appears verbatim in a stored chunk.
          "partial"    - the vault has strong FTS5 keyword support for the
                         claim, but no verbatim match.
          "unverified" - no vault support for the claim.

        With multiple vaults selected (`--vault a,b`), the verdict folds across
        them: VERIFIED if ANY vault verifies, else PARTIAL if any is partial,
        else UNVERIFIED — so a claim grounded in a companion vault audits clean
        (the RF-Tech cross-vault pain point).

        This makes the [V] honor-system tag machine-checkable: the caller can
        refuse to tag [V] unless this returns "verified".

        Comparison is strict but typography-blind (see `normalize_claim`):
        en/em dashes, smart quotes, NBSP and full-width Unicode are folded so a
        typesetter's dash choice never flips the verdict — while token identity
        ("400K" vs "400K+") and word order are still enforced.
        """
        if not claim or not claim.strip():
            return "unverified"
        best = "unverified"
        for vault in self.vaults:
            result = self._verify_against_vault(vault, claim)
            if result == "verified":
                return "verified"
            if result == "partial":
                best = "partial"
        return best

    def _verify_against_vault(self, vault: VaultManager, claim: str) -> str:
        """Verify a claim against a single VaultManager (the shared per-vault
        logic behind `verify_claim`'s cross-vault fold)."""
        needle = normalize_claim(claim)
        candidates: list[str] = []
        with vault._db() as (_conn, cursor):
            # Slide a 60-char window across the needle so a claim whose
            # *distinctive* portion is not its first 60 chars still matches
            # verbatim (A1). Every windowed fragment is tested; the first hit
            # confirms the claim.
            step = 60
            window = 60
            # A short needle is one fragment; a long one is broken into
            # overlapping windows so a distinctive tail gets a chance.
            fragments = [needle[i:i + window]
                         for i in range(0, max(1, len(needle) - window + 1), step)]
            for fragment in fragments:
                if not fragment.strip():
                    continue
                # Escape LIKE wildcards in the fragment itself before widening.
                like_fragment = fragment.replace("\\", "\\\\").replace("%", r"\%").replace("_", r"\_")
                # Widen whitespace runs AND the typographic equivalences to %:
                # stored text may use an en-dash where the claim has a hyphen,
                # curly quotes where the claim has straight ones, etc. The
                # Python confirm below is authoritative, so over-widening here
                # only costs a few extra candidate rows, never a wrong verdict.
                like_fragment = re.sub(r"[\s\-'`\*\u2010\u2011\u2012\u2013\u2014\u2015\u2212\u2018\u2019\u201a\u201b\u201c\u201d\u201e\u201f]+", "%", like_fragment)
                cursor.execute(
                    "SELECT text FROM chunks_fts WHERE lower(text) LIKE ? ESCAPE '\\'",
                    (f"%{like_fragment}%",)
                )
                candidates = [row[0] for row in cursor.fetchall()]
                for raw in candidates:
                    if needle in normalize_claim(raw):
                        return "verified"

        # 2) FTS5 keyword overlap: build a proper AND-of-phrases MATCH for the
        #    claim and measure the strength of the top hit. "Partial" is only
        #    reported when the top all-terms hit measurably beats the vault's
        #    coincidence floor — the best rank any single claim term achieves
#     alone. Semantics are the same as an absolute-BM25 bar (co-
        #     occurrence of a few common stopwords is NOT evidence, A2) but
        #     corpus-size independent: raw FTS5 ranks are ~1e-6 on a small
        #     vault and ~-20 on a large one, so a fixed absolute cutoff like
        #     -2.0 is unreachable at small scale and trivially passed at
        #     scale. The margin below is relative to the single-term floor,
        #     so a fixed ratio of separation is always required.
        fts = vault._fts_query(claim)
        if fts:
            with vault._db() as (_conn, cursor):
                cursor.execute(
                    "SELECT rank FROM chunks_fts WHERE chunks_fts MATCH ? "
                    "ORDER BY rank LIMIT 1",
                    (fts,)
                )
                row = cursor.fetchone()
                if row is None:
                    return "unverified"  # no chunk matches every claim term
                top_all = row[0]
                # Coincidence floor: the least-negative best single-term rank.
                floor: float | None = None
                for single in fts.split(" AND "):
                    cursor.execute(
                        "SELECT rank FROM chunks_fts WHERE chunks_fts MATCH ? "
                        "ORDER BY rank LIMIT 1",
                        (single,)
                    )
                    sr = cursor.fetchone()
                    if sr is None:  # a claim term matches nothing anywhere
                        floor = None
                        break
                    floor = sr[0] if floor is None else max(floor, sr[0])
                if floor is not None:
                    # Margin is relative to the coincidence floor so it is
                    # corpus-size free: FTS5 ranks are ~1e-6 on a small vault
                    # and ~-20 on a large one, so a fixed absolute margin is
                    # either unattainable at scale or trivially passed.
                    margin = 0.15 * abs(floor)
                    if top_all <= floor - margin:
                        return "partial"
        return "unverified"

    def verify_hint(self, claim: str, recall: int = 5) -> str | None:
        """Coaching message for a non-verified claim: the nearest vault phrase.

        `verify` is exact-phrasing by design: on PARTIAL/UNVERIFIED the caller
        is expected to re-express the claim in the source's own words. This
        returns a ready-to-print hint showing the closest stored phrase (best
        fuzzy overlap versus the normalized claim, across every selected vault)
        plus a reformulation nudge, so the denial reads as an instruction
        instead of a dead-end. None when no vault has anything remotely similar.
        """
        if not claim or not claim.strip():
            return None
        needle = normalize_claim(claim)
        if not needle:
            return None
        best: tuple[float, str] | None = None  # (ratio, phrase)
        for vault in self.vaults:
            candidate = self._verify_hint_against(vault, claim, needle, recall)
            if candidate is not None and (best is None or candidate[0] > best[0]):
                best = candidate
        if best is None:
            return None
        ratio, phrase = best
        # Trim the normalized phrase to a tight window around the fuzzy match.
        m = SequenceMatcher(None, needle, phrase).find_longest_match(
            0, len(needle), 0, len(phrase))
        lo_full = max(0, m.b - 40)
        hi_full = min(len(phrase), m.b + m.size + 40)
        window = phrase[lo_full:hi_full]
        shown = (("…" if lo_full > 0 else "") + window +
                 ("…" if hi_full < len(phrase) else ""))
        matched = window[m.b - lo_full:m.b - lo_full + m.size] if m.size else ""
        return (
            f"  nearest vault phrase (overlap {ratio:.0%}):\n"
            f"    “{shown}”\n"
            f"  — reword your claim to match the source text exactly"
            + (f" (the vault writes {matched!r} here)" if matched and matched not in needle else "")
            + "\n  — then re-run this verify."
        )

    def _verify_hint_against(self, vault: VaultManager, claim: str,
                             needle: str, recall: int) -> tuple[float, str] | None:
        """Best (overlap_ratio, normalized_phrase) in a single VaultManager."""
        fts = vault._fts_query(claim)
        top: list[str] = []
        if fts:
            with vault._db() as (_conn, cursor):
                cursor.execute(
                    "SELECT text FROM chunks_fts WHERE chunks_fts MATCH ? "
                    "ORDER BY rank LIMIT ?",
                    (fts, recall),
                )
                top = [row[0] for row in cursor.fetchall()]
                # Also pull the oldest rows: a fresh spin-off vault may have
                # better prose in low-ranked chunks than the top FTS hit.
                cursor.execute(
                    "SELECT text FROM chunks_fts ORDER BY rowid LIMIT ?",
                    (recall,),
                )
                top.extend(row[0] for row in cursor.fetchall())
        best: tuple[float, str] | None = None  # (ratio, phrase)
        for raw in top:
            if not raw or raw.isspace():
                continue
            ratio, _start, size = _nearest_phrase_probe(raw, needle)
            if size == 0:
                continue
            phrase = normalize_claim(raw)
            best = (ratio, phrase) if best is None or ratio > best[0] else best
        return best

    def audit_artifact(self, path: str) -> dict[str, Any]:
        """Audit a synthesis artifact's `[V#N]` citation chain against the vault.

        For every `[V#N]` tag in the artifact, three links of the evidence chain
        are checked (this is the execution-provenance gate the stress test
        exposed as a gap — claim-level verify alone never proves the tag maps
        to a listed, ingested source):

          1. VERBATIM  — the sentence containing the tag (or its longest inline
                         double-quoted passage) must `verify` as VERIFIED against
                         the vault. A bare `[V]` (no #N) is verified but not
                         mapping-checked. When several `[V#N]` tags share a
                         line, each is attributed to the double-quoted passage
                         ending nearest before it; a quote wrapped across two
                         physical lines is joined into one claim before
                         extraction. Non-verified claims carry a `hint` field
                         (the nearest stored vault phrase) so the author knows
                         how to reword them instead of guessing.
          2. MAPPED    — N must appear in the artifact's "Source Links /
                         Citations" block as `[#N] <url>`.
          3. INGESTED  — that cited URL must have chunks in the vault.

        Returns a dict with: `claims` (list of {line, n, text, verdict, quote}),
        `unmapped` (referenced Ns with no Source Link), `not_ingested` (Ns whose
        URL has no vault chunks), verdict counts, and `accuracy`
        (verified / total, excluding neither — a partial counts against).
        """
        with open(path, encoding="utf-8") as fh:
            lines = fh.read().splitlines()

        source_block = self._extract_source_links(lines)

        tag_re = re.compile(r"\[(V|E|H)(?:#(\d+))?\]")
        claims: list[dict[str, Any]] = []
        verdicts: dict[str, str] = {}
        emitted: set[tuple[str, str | None]] = set()
        in_links = False
        # Authoring-smell warnings: a `[V#N]` on a line explicitly marked as
        # analysis (`[H]`/`[E]` appearing BEFORE the cite tag) is a strong
        # signal the tag was attached to paraphrase, not a verbatim quote —
        # the recurring live failure mode. Informational only (never changes
        # the exit code); it just calls out the likely fix so a 100% audit
        # needs no hunting.
        warnings: list[dict[str, Any]] = []
        for unit_text, ln_no in self._logical_lines(lines):
            if not unit_text.strip():
                continue
            if re.match(r"^#{1,6}\s", unit_text) and (
                    "source link" in unit_text.lower() or "citation" in unit_text.lower()):
                in_links = True
                continue
            if in_links:
                continue
            # Skip backticked spans (e.g. a prose mention of the literal token
            # `` `[V]` ``) so narrative references are not audited as claims.
            scan = re.sub(r"`[^`]*`", " ", unit_text)
            # Clean heading/list markers but keep cite tags so each tag's
            # position maps to the same offsets as the quoted spans.
            cleaned0 = re.sub(r"^#+\s*", "", scan.strip())
            cleaned0 = re.sub(r"^\s*(?:[-*+]\s+|\d+[.)]\s+)", "", cleaned0)
            cleaned0 = re.sub(r"\s+", " ", cleaned0).strip()
            # Authoring smell: does an [H]/[E] marker precede a [V#N] on this
            # line? If so, the [V#N] is likely riding on analysis prose. Run on
            # the RAW unit_text (backticks intact) because authors mark
            # `` `[H]` ``/`` `[E]` `` in backticks; stripping them first would
            # hide the very marker we are looking for.
            raw_first_vn = next(
                (m for m in tag_re.finditer(unit_text)
                 if m.group(1) == "V" and m.group(2) is not None),
                None)
            if raw_first_vn is not None and any(
                    m.group(1) in ("H", "E")
                    for m in tag_re.finditer(unit_text[:raw_first_vn.start()])):
                warnings.append({
                    "line": ln_no,
                    "type": "v_on_analysis_line",
                    "tag": raw_first_vn.group(0),
                })
            for m in tag_re.finditer(cleaned0):
                if m.group(1) != "V":
                    continue
                text_cand, is_quote = self._claim_text_for_tag(cleaned0, m.start())
                key = normalize_claim(text_cand)
                if not key:
                    continue
                # One claim row per (claim, source-tag) pair: a line carrying
                # two [V#N] tags is two chain links, but repeated tags of the
                # same N on one line must not double-count the claim.
                n = m.group(2)
                if (key, n) in emitted:
                    continue
                emitted.add((key, n))
                verdict = verdicts.get(key) or self.verify_claim(text_cand)
                verdicts[key] = verdict
                hint = None
                if verdict != "verified":
                    hint = self.verify_hint(text_cand, recall=5)
                claims.append({
                    "line": ln_no, "n": n,
                    "text": text_cand, "quote": is_quote,
                    "verdict": verdict, "hint": hint,
                })

        used_n = sorted({c["n"] for c in claims if c["n"]})
        unmapped = [n for n in used_n if n not in source_block]
        not_ingested: list[tuple[str, str]] = []
        if not unmapped:
            for n in used_n:
                url = source_block[n]
                # INGESTED is read across every selected vault: with
                # `--vault a,b` a source living in a companion vault counts as
                # ingested (the profile-spans-two-vaults case).
                if not any(v.get_chunks_for_url(url) for v in self.vaults):
                    not_ingested.append((n, url))

        counts = {"verified": 0, "partial": 0, "unverified": 0}
        for c in claims:
            counts[c["verdict"]] = counts.get(c["verdict"], 0) + 1
        total = len(claims)
        return {
            "path": path, "claims": claims, "counts": counts,
            "total": total, "unmapped": unmapped, "used_n": used_n,
            "not_ingested": not_ingested, "warnings": warnings,
            "accuracy": (counts["verified"] / total) if total else 0.0,
        }

    def _extract_source_links(self, lines: list[str]) -> dict[str, str]:
        """Map `[#N] url` entries in a markdown "Source Links / Citations" block."""
        source_block: dict[str, str] = {}
        in_links = False
        for ln in lines:
            if re.match(r"^#{1,6}\s", ln):
                in_links = ("source link" in ln.lower() or "citation" in ln.lower())
                continue
            if in_links:
                m = re.match(r"\[\s*#\s*(\d+)\s*\](?::)?\s*(\S+)", ln)
                if m:
                    source_block[m.group(1)] = m.group(2).rstrip(",;")
        return source_block

    @staticmethod
    def _unclosed_quotes(text: str) -> bool:
        """True when `text` ends inside an open double-quoted passage.

        Straight `"` toggles a quote (open on the first, close on the next);
        curly `“`/`”` open/close explicitly. Used to join a quote that is
        wrapped across physical lines into one logical audit unit.
        """
        depth = 0
        for ch in text:
            if ch == '“':
                depth += 1
            elif ch == '”':
                depth = max(0, depth - 1)
            elif ch == '"':
                if depth > 0:
                    depth -= 1
                else:
                    depth += 1
        return depth > 0

    @staticmethod
    def _logical_lines(lines: list[str]) -> list[tuple[str, int]]:
        """Group consecutive non-blank lines into logical units, continuing a
        unit while a double-quoted passage is still open across the line break.

        Returns `(joined_text, start_line)` pairs (1-based) so a quote that a
        markdown editor wraps over two physical lines is audited as one claim
        instead of two fragment lines. Headings, list bullets, and balanced
        single-line prose each stay their own unit.
        """
        out: list[tuple[str, int]] = []
        buf: list[str] = []
        start = 1
        for i, ln in enumerate(lines, 1):
            if not ln.strip():
                if buf:
                    out.append((" ".join(buf), start))
                    buf = []
                continue
            if not buf:
                start = i
            buf.append(ln.strip())
            if HoardCore._unclosed_quotes(" ".join(buf)):
                continue
            out.append((" ".join(buf), start))
            buf = []
        if buf:
            out.append((" ".join(buf), start))
        return out

    @staticmethod
    def _quotes_in(text: str) -> list[tuple[int, str]]:
        """Inline double-quoted spans as (char_offset, quote_text) pairs."""
        return [(m.start(1), m.group(1))
                for m in re.finditer(r'[“"]([^”"]+)[”"]', text)]

    @staticmethod
    def _clean_claim_line(line: str) -> str:
        """Strip markdown/cite-tag noise from an artifact line for claim text."""
        cleaned = re.sub(r"^#+\s*", "", line.strip())
        cleaned = re.sub(r"`[^`]*`", " ", cleaned)
        cleaned = re.sub(r"^\s*(?:[-*+]\s+|\d+[.)]\s+)", "", cleaned)
        cleaned = re.sub(r"\[(?:V|E|H)(?:#\d+)?\]|\[CONTROL\]|\[INCOMPLETE[^\]]*\]",
                         "", cleaned)
        return re.sub(r"\s+", " ", cleaned).strip()

    @staticmethod
    def _claim_text_from_line(line: str) -> tuple[str, bool]:
        """Extract the auditable claim from an artifact line.

        Strips list/emphasis/cite-tag noise, then prefers the longest inline
        double-quoted passage (the verbatim burden an author signals by quoting)
        when it is long enough to be distinctive; otherwise the whole cleaned
        line is the claim. Returns (text, was_quoted).
        """
        cleaned = HoardCore._clean_claim_line(line)
        best = max([q for _, q in HoardCore._quotes_in(cleaned)],
                   key=len, default="")
        if best and len(normalize_claim(best)) >= 24:
            return best.strip(" \t.,;:!?’"), True
        return cleaned, False

    @staticmethod
    def _claim_text_for_tag(cleaned0: str, tag_pos: int) -> tuple[str, bool]:
        """Attribute the claim to the inline quote ending nearest before the
        tag (its own citation target) when that quote is distinctive; falls
        back to the whole line's longest quote / whole cleaned line.

        `cleaned0` is the line cleaned of heading/list markers with cite tags
        still present so `tag_pos` maps to the same offsets as `_quotes_in`.
        """
        spans = [(s + len(q), q) for s, q in HoardCore._quotes_in(cleaned0)
                 if s + len(q) <= tag_pos]
        nearest = max(spans, default=(0, ""))[1]
        if nearest and len(normalize_claim(nearest)) >= 24:
            return nearest.strip(" \t.,;:!?’"), True
        return HoardCore._claim_text_from_line(cleaned0)

    async def _process_document(self, url: str, strategy: str, force_refresh: bool) -> tuple[list[Chunk], dict[str, Any]]:
        """
        Core processing pipeline for a single URL.
        Returns (chunks, meta_overrides).
        """
        # Ad-redirect / tracking beacons are never content; don't fetch them.
        if is_ad_tracking_url(url):
            logger.warning(f"Skipping index of {url} (ad/tracking URL).")
            return [Chunk(
                text="",
                metadata={"source": url, "junk": True,
                          "junk_reason": "ad_tracking_url"}
            )], {"junk": True, "junk_reason": "ad_tracking_url"}

        # Check cache
        if not force_refresh and self.vault.document_exists(url, self.config.get('cache.ttl_seconds', 86400)):
            logger.info(f"Cache HIT for {url} (in vault). Skipping network.")
            # Return empty chunks, but indicate cache hit
            return [], {"cached": True, "url": url}

        # Fetch
        try:
            text, binary, content_type, status = await self.fetcher.fetch(url, strategy)
        except RuntimeError as e:
            error_msg = str(e)
            if error_msg == "CF_COOKIE_EXPIRED":
                return [Chunk(
                    text="CF_COOKIE_EXPIRED: Your Cloudflare cookie has expired. Update 'auth.cookie_string' in hoardcore.toml.",
                    metadata={"source": url, "error": True, "code": "CF_COOKIE_EXPIRED"}
                )], {"error": "CF_COOKIE_EXPIRED"}
            elif error_msg == "FETCH_FAILED":
                return [Chunk(
                    text=f"Network error: All strategies failed for {url}.",
                    metadata={"source": url, "error": True, "status": "network_blocked"}
                )], {"error": "FETCH_FAILED"}
            raise

        # Save binary if applicable
        binary_path = None
        if binary and self.save_binary and content_type not in ['text/html', 'text/plain']:
            binary_path = self.vault.save_binary(url, content_type, binary)

        # A body delivered alongside a 4xx/5xx status is an error page, not
        # content — refusing it keeps soft-404 bodies out of the vault.
        if status is not None and status >= 400:
            logger.warning(f"Skipping index of {url} (http_error_status={status}).")
            return [Chunk(
                text="",
                metadata={"source": url, "junk": True,
                          "junk_reason": f"http_error_status={status}"}
            )], {"junk": True, "junk_reason": f"http_error_status={status}"}

        # Parse document
        markdown = ""
        parser_meta: dict[str, Any] = {"content_type": content_type}

        if 'text/html' in content_type or 'text/plain' in content_type:
            if text is None and binary:
                text = binary.decode('utf-8', errors='ignore')
            if text:
                markdown, html_meta = await self.parser.clean_html(text, url)
                parser_meta.update(html_meta)
                parser_meta['source_type'] = 'html'
        elif binary:
            markdown, bin_meta = await self.parser.parse_binary(content_type, binary)
            parser_meta.update(bin_meta)
            parser_meta['source_type'] = 'binary'
            parser_meta['binary_path'] = binary_path

        if not markdown:
            markdown = "[No extractable content found]"

        # Quality assessment
        original_len = len(text) if text else (len(binary) if binary else 1)
        quality_score = len(markdown) / max(original_len, 1)
        parser_meta['quality_score'] = min(quality_score, 1.0)
        parser_meta['quality_label'] = 'high' if quality_score > 0.3 else ('medium' if quality_score > 0.1 else 'low')

        # Junk detection: refuse to persist boilerplate / empty extraction.
        # These otherwise pollute the FTS vault with garbage that drowns real results.
        junk_reason = self._detect_junk(markdown, text, parser_meta, quality_score)
        if junk_reason:
            logger.warning(f"Skipping index of {url} (junk: {junk_reason}).")
            return [Chunk(
                text=markdown,
                metadata={
                    "source": url,
                    "header_path": "",
                    "junk": True,
                    "junk_reason": junk_reason,
                    "quality_score": parser_meta['quality_score'],
                    "parser": parser_meta.get('parser', 'unknown')
                }
            )], {**parser_meta, "junk": True, "junk_reason": junk_reason}

        # Chunk
        chunks = await self.chunker.chunk(markdown, url, parser_meta)

        # Save extracted text to disk
        self.vault.save_extracted_text(url, markdown, chunks, parser_meta)

        # Route through the parallel pipeline when enabled (large batches only).
        self.vault.ingest_chunks_parallel(url, chunks, parser_meta)

        return chunks, parser_meta

    @staticmethod
    def _detect_junk(markdown: str, raw_text: str | None, parser_meta: dict[str, Any], quality_score: float) -> str | None:
        """Return a reason string if extraction is boilerplate/empty, else None."""
        stripped = markdown.strip()

        # Explicit empty-marker produced by the parser pipeline.
        if not stripped or "[No extractable content found]" in stripped:
            return "empty_extraction"

        # Generic block/redirect/captcha/consent pages masquerade as real content.
        boilerplate = [
            "Please click", "if you are not redirected",
            "are you having trouble accessing", "click here",
            "enable javascript", "robot check", "verify you are human",
            "cloudflare", "not a robot", "access denied",
            "the page you are looking for", "we couldn't find the page",
            "bad gateway", "404 not found", "the requested url was not found",
            # login / consent walls (extraction sees only the sign-in chrome)
            "agree & join", "new to linkedin", "join now", "create new account",
            "forgot your password", "log in to see", "sign in",
            "your web browser is not fully supported",
            "instagram from meta", "meta pay",
            "email o telepono", "mag-sign up", "mag-log in",
            "patakaran sa privacy", "log ng aktibidad",
        ]
        lower = stripped.lower()
        matched = [b for b in boilerplate if b in lower]
        # Very short extracted body is almost always a mis-hit.
        if quality_score < 0.02 and len(stripped) < 60:
            return "near_empty_extraction"
        if matched and len(stripped) < 600:
            return f"boilerplate:{matched[0]}"

        # A low-quality flat list (e.g. a site's language-picker while logged
        # out) has no prose at all: many short lines, no sentence-like line.
        if quality_score < 0.15 and 60 <= len(stripped) < 4000:
            lines = [ln for ln in stripped.splitlines() if ln.strip()]
            if (len(lines) >= 12
                    and all(len(ln.split()) <= 3 for ln in lines)):
                return "chrome_shell:flat_list"

        # Some failures leave a low-score body that is still real (PDF ratio is naturally low);
        # rely on structural signals above, not raw length ratio alone.
        return None

    async def _scrape_single(self, url: str, strategy: str, force_refresh: bool) -> list[Chunk]:
        """Scrape a single URL."""
        chunks, meta = await self._process_document(url, strategy, force_refresh)
        if meta.get('cached'):
            # Cache hit: the pipeline fetches nothing, so serve the vaulted
            # chunks back to the caller instead of an empty result.
            return self.vault.get_chunks_for_url(url)
        return chunks

    async def _crawl_domain(self, url: str, strategy: str, force_refresh: bool) -> list[Chunk]:
        """Crawl an entire domain using sitemap."""
        all_chunks = []
        discovered_urls = await self.crawler.discover_urls(url)

        if not discovered_urls:
            logger.warning(f"No URLs discovered for {url}. Falling back to single scrape.")
            return await self._scrape_single(url, strategy, force_refresh)

        logger.info(f"Discovered {len(discovered_urls)} URLs for crawling.")

        # Use semaphore to limit parallel workers
        max_workers = max(1, self.config.get('crawler.parallel_workers', 5))
        semaphore = asyncio.Semaphore(max_workers)

        async def _crawl_one(single_url: str) -> list[Chunk]:
            async with semaphore:
                try:
                    chunks, meta = await self._process_document(single_url, strategy, force_refresh)
                    if meta.get('cached'):
                        # Cache hit: the pipeline fetched nothing, so serve the
                        # vaulted chunks back (mirrors _scrape_single) instead
                        # of silently reporting zero content for the URL.
                        cached = self.vault.get_chunks_for_url(single_url)
                        all_chunks.extend(cached)
                        return cached
                    if chunks and not chunks[0].metadata.get('error'):
                        all_chunks.extend(chunks)
                    return chunks
                except Exception as e:
                    logger.error(f"Failed to crawl {single_url}: {e}")
                    return []

        tasks = [_crawl_one(u) for u in discovered_urls]
        await asyncio.gather(*tasks, return_exceptions=True)

        return all_chunks

    async def fetch(
        self,
        url: str,
        action: str = "scrape",
        strategy: str | None = None,
        query: str | None = None,
        force_refresh: bool = False,
        urls: list[str] | None = None,
        max_results: int = 0,
        mode: str | None = None,
    ) -> list[dict[str, Any]]:
        """
        Public entry point.

        Args:
            url: Target URL or domain.
            action: "scrape", "crawl", "search", "ingest", or "discover".
            strategy: "fast", "balanced", or "aggressive".
            query: Required for "search" and "discover" actions.
            force_refresh: Ignore cache and re-fetch.
            urls: An explicit list of URLs to process in parallel, honored for
                "scrape", "crawl", and "ingest" (a placeholder `_` positional is
                never sent to the fetcher).
            max_results: For "discover", how many top results to ingest; for
                "search", the max chunks returned (falls back to config).
            mode: For "search", "fast" (FTS-only) or "hybrid" (vector+RRF). Note
                that with embeddings.fts_fast_path=true (default), "hybrid" can
                still return the FTS fast path when FTS5 alone fills the result
                set (hits tagged retrieval='fts_fast').

        Returns:
            List of dicts with "text" and "metadata" for the LLM.
        """
        if strategy is None:
            strategy = str(self.config.get('network.default_strategy', 'aggressive'))

        # Route actions
        if action == "search":
            if not query:
                return [{
                    "text": "Error: 'query' parameter required for action='search'.",
                    "metadata": {"error": True}
                }]
            limit = (max_results if max_results > 0
                     else self.config.get('indexer.search_limit', 20))
            domain = urlparse(url).netloc or None
            hybrid: bool | None = None
            if mode == 'fast':
                hybrid = False
            elif mode == 'hybrid':
                hybrid = True
            chunks = self._search_across_vaults(query, limit, domain=domain,
                                                hybrid=hybrid)
            self.bus.emit("search.completed", query=query,
                          n_results=len(chunks), domain=domain)
            return [c.to_dict() for c in chunks]

        elif action == "ingest":
            # Explicit URL list (comma/whitespace separated). Closes the
            # discovery gap by letting callers feed curatined URLs directly.
            if not urls:
                return [{
                    "text": "Error: 'urls' parameter required for action='ingest'.",
                    "metadata": {"error": True}
                }]
            return await self._ingest_many(urls, strategy, force_refresh)

        elif action == "discover":
            # Live web discovery: query -> ranked URLs -> parallel ingest.
            # Closes the discovery gap WITHOUT an API key by going through the
            # same resilient fetch chain (incl. FlareSolverr) as the crawler.
            if not query:
                return [{
                    "text": "Error: 'query' parameter required for action='discover'.",
                    "metadata": {"error": True}
                }]
            return await self._discover_and_ingest(query, max_results, strategy, force_refresh)

        elif action == "crawl":
            if urls:
                return await self._ingest_many(urls, strategy, force_refresh)
            chunks = await self._crawl_domain(url, strategy, force_refresh)
            return [c.to_dict() for c in chunks if not c.metadata.get('error', False)]

        else:  # "scrape" (default)
            if urls:
                return await self._ingest_many(urls, strategy, force_refresh)
            chunks = await self._scrape_single(url, strategy, force_refresh)
            return [c.to_dict() for c in chunks]

    async def _ingest_many(self, urls: list[str], strategy: str, force_refresh: bool) -> list[dict[str, Any]]:
        """Process an explicit list of URLs with a bounded-worker pool."""
        max_workers = max(1, self.config.get('crawler.parallel_workers', 5))
        semaphore = asyncio.Semaphore(max_workers)
        results: list[dict[str, Any]] = []

        async def _ingest_one(target: str) -> None:
            async with semaphore:
                try:
                    chunks, meta = await self._process_document(target, strategy, force_refresh)
                    if meta.get('cached'):
                        # Cache hit: the pipeline fetched nothing, so serve the
                        # vaulted chunks back (mirrors _scrape_single) instead
                        # of silently reporting zero content for the URL.
                        results.extend(
                            c.to_dict() for c in self.vault.get_chunks_for_url(target)
                        )
                        return
                    if meta.get('error'):
                        if chunks:
                            results.append(chunks[-1].to_dict())
                        return
                    if not meta.get('junk'):
                        results.extend(c.to_dict() for c in chunks)
                except Exception as e:
                    logger.error(f"Failed to ingest {target}: {e}")
                    results.append({
                        "text": f"Error ingesting {target}: {e}",
                        "metadata": {"source": target, "error": True}
                    })

        await asyncio.gather(*[_ingest_one(u) for u in urls], return_exceptions=True)
        return results

    async def _discover_and_ingest(self, query: str, max_results: int,
                                   strategy: str, force_refresh: bool) -> list[dict[str, Any]]:
        """Run a live web search, then ingest the top-ranked URLs.

        Uses the free DuckDuckGo HTML provider through the existing
        fetch chain, so FlareSolverr is applied automatically when a search
        result is anti-bot protected. Returns the ingested chunks.
        """
        cfg_max = self.config.get('discovery.max_results', 10)
        cfg_top = self.config.get('discovery.top_rank', 6)
        # `--limit N` now means "ingest the top N results" (per docs); the
        # search pool only needs to be at least that large, so it never
        # shrinks below the configured default.
        ingest_n = max_results if max_results > 0 else cfg_top
        pool = max(cfg_max, ingest_n)

        results = await self.discovery.search(query, max_results=pool, strategy=strategy)
        if not results:
            self.bus.emit("discovery.completed", query=query, urls=[])
            return [{
                "text": f"No URLs discovered for query: {query!r}.",
                "metadata": {"source": "discovery", "error": True, "query": query}
            }]

        logger.info(f"Discovered {len(results)} URLs for query: {query!r}")
        self.bus.emit("discovery.completed", query=query,
                      urls=[r.url for r in results])

        # rank-biased ingest: take the top-N results (--limit, else top_rank)
        targets = [r.url for r in results[:ingest_n]]
        summary = [{
            "text": f"[discovery] top {len(targets)} URLs for {query!r}",
            "metadata": {
                "source": "discovery", "query": query,
                "candidates": [r.__dict__ for r in results]
            }
        }]

        ingested = await self._ingest_many(targets, strategy, force_refresh)
        return summary + ingested

# =============================================================================
# 9. CLI ENTRYPOINT
# =============================================================================

def citation_list(sources: list[str] | dict[str, str]) -> str:
    """Module-level alias for the artifacts **Source Links / Citations** block.

    Accepts either a list of source URLs or a ``{label: url}`` mapping and
    renders the exact ``[#N] <label> — <url>`` lines every artifact should close
    with (the provenance tags use ``[V#N]``/``[E#N]`` -> ``[#N]``). This makes
    ``hoardcore.citation_list(urls)`` available at module level exactly as
    skill.md documents it, without constructing a manager.
    """
    return HoardCore.citation_list(sources)


def write_artifact(filename: str, content: str) -> str:
    """Module-level convenience for encoding a research deliverable into the
    day-sorted artifacts directory (``hoardcore.write_artifact`` per skill.md).

    Constructs a default HoardCore instance so the vault config (artifacts
    dir, day-sorting) is honored; returns the absolute path written.
    """
    return HoardCore().write_artifact(filename, content)


def organize_artifacts_by_day() -> list[str]:
    """Module-level convenience: move any legacy flat artifacts/ files into
    ``artifacts/YYYY-MM-DD/`` subfolders by mtime (``hoardcore.organize_artifacts_by_day``).
    """
    return HoardCore().organize_artifacts_by_day()

_EXAMPLES = (
    "\nExamples:\n"
    "  python hoardcore.py https://example.com --action scrape\n"
    "  python hoardcore.py https://docs.python.org --action crawl --strategy aggressive\n"
    "  python hoardcore.py https://arxiv.org/abs/2110.12345.pdf --action scrape\n"
    "  python hoardcore.py https://example.com --action search --query 'machine learning'\n"
    "  python hoardcore.py _ --action ingest --urls 'u1,u2,u3'\n"
    "  python hoardcore.py _ --action scrape --urls 'u1,u2'\n"
    "  python hoardcore.py _ --action discover --query 'negros renewable energy' --limit 5\n"
    "  python hoardcore.py _ --action research --query 'how does bokashi compost' --discover 5 --recall 6\n"
    "  python hoardcore.py _ --action research --query 'negros economy' --out artifacts/report.md\n"
    "  python hoardcore.py _ --action research --query 'sleep research' --vault sleep\n"
    "  python hoardcore.py _ --action verify --claim 'the Epoch doubling time is 6 months'\n"
    "  python hoardcore.py _ --action verify --claim '500-2000 stars' --hint  # nearest-vault coaching\n"
    "  python hoardcore.py _ --action local --path docs/        # ingest local_inputs/docs/\n"
    "  python hoardcore.py _ --action local --list              # read-only scan of local_inputs/\n"
    "  python hoardcore.py _ --action research --vault career,negros_ai_jobs  # cross-vault read\n"
    "  python hoardcore.py _ --action check   # verify vault integrity\n"
    "  python hoardcore.py _ --action stats   # sources/chunks/vectors summary\n"
)


def _build_parser() -> argparse.ArgumentParser:
    """Argument parser for the hoardcore CLI (D3): typed, --help, no silent
    typo tolerance. `url` is positional so legacy `hoardcore.py <URL>` calls,
    and the `_` placeholder for vault-only actions, keep working verbatim."""
    parser = argparse.ArgumentParser(
        prog="hoardcore",
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=_EXAMPLES,
        # Reject prefix abbreviations so a typo like `--recal` fails loudly
        # instead of silently matching --recall (D3).
        allow_abbrev=False,
    )
    parser.add_argument(
        "url", nargs="?", default="_",
        help="Target URL or domain, or '_' for vault-only actions.",
    )
    parser.add_argument(
        "--action", choices=["scrape", "crawl", "search", "ingest",
                             "discover", "research", "verify", "check", "stats",
                             "audit", "local"],
        default="scrape", help="Action to run (default: scrape).",
    )
    parser.add_argument(
        "--strategy", choices=["fast", "balanced", "aggressive"], default=None,
        help="Fetch strategy (default: network.default_strategy).",
    )
    parser.add_argument(
        "--query", default=None, help="Search/research query."
    )
    parser.add_argument(
        "--claim", default=None,
        help="Claim to verify verbatim against the vault. Exact phrasing "
             "(typographic variants like en/em dashes and curly quotes are "
             "folded, but '%%' vs 'percent' are NOT). In shells, escape '$' as "
             "\\$ (bash expands $13 to empty) or use --claim-file.",
    )
    parser.add_argument(
        "--hint", action="store_true",
        help="With --action verify on a denied claim: print the nearest vault "
             "phrase so the claim can be reworded to the source's own words.",
    )
    parser.add_argument("--discover", type=int, default=None,
                        help="Results to discover+ingest for research.")
    parser.add_argument("--recall", type=int, default=6,
                        help="Chunks to recall for research.")
    parser.add_argument("--out", default=None, dest="out_path",
                        help="Artifact output path for research.")
    parser.add_argument("--force", action="store_true",
                        help="Bypass the cache and re-fetch.")
    parser.add_argument("--path", default=None, dest="local_path",
                        help="With --action local: relative path (file or "
                             "directory) inside storage.local_dir to process; "
                             "defaults to the whole local_dir.")
    parser.add_argument("--list", action="store_true", dest="list_only",
                        help="With --action local: list supported files under "
                             "--path without ingesting them.")
    parser.add_argument("--urls", default=None,
                        help="Comma/space-separated URL list for scrape/crawl/ingest.")
    parser.add_argument("--limit", type=int, default=0,
                        help="Top results to ingest for discover; max chunks for search.")
    parser.add_argument("--vault", default=None, dest="vault_name",
                        help="Per-topic vault name.")
    parser.add_argument("--migrate", action="store_true",
                        help="With --action check: rebuild vault at 16 KB pages.")
    parser.add_argument("--mode", choices=["fast", "hybrid"], default=None,
                        help="Force search mode (FTS-only vs vector+RRF).")
    parser.add_argument(
        "--parallel", action=argparse.BooleanOptionalAction, default=None,
        help="Override threaded ingest for this run (on/off). Engages the "
             "parallel reader->embed->write pipeline for batches of 8+ chunks. "
             "Defaults to config indexer.parallel (true).",
    )
    parser.add_argument("--no-answer-first", action="store_true",
                        help="With --action research: always run live DISCOVER, "
                             "even if the existing vault has a high-confidence answer.")
    parser.add_argument("--log-level", choices=["debug", "info", "warning",
                                                "error"], default=None,
                        help="Override the log verbosity for this run "
                             "(default: info).")
    parser.add_argument("--keep-low", action="store_true",
                        help="With --action research: retain low-confidence hits "
                             "in the grounding context (skip filter_low) — for "
                             "exhaustive/deep hunts that want the full evidence tail.")
    parser.add_argument("--claim-file", default=None,
                        help="With --action verify: read the claim from this file "
                             "instead of --claim. Keeps characters like '$' intact "
                             "(bash would otherwise expand '$13' to empty).")
    parser.add_argument("--claim-list", default=None,
                        help="With --action verify: a file of claims, one per line "
                             "(blank and '#' lines skipped). Verifies each and prints "
                             "an aggregate citation-accuracy percentage. Mutually "
                             "exclusive with --claim/--claim-file.")
    parser.add_argument("--artifact", default=None,
                        help="With --action audit: path to a synthesis artifact (.md). "
                             "Parses every [V#N] tag, verifies each claim verbatim "
                             "against the vault, checks that N maps to a Source Link, "
                             "and that the cited source has chunks in the vault.")
    return parser


async def main(argv: list[str] | None = None) -> None:
    """CLI entry point.

    Wraps the implementation in a safety net: any unexpected exception exits
    cleanly with a short message and code 2 instead of dumping a raw traceback
    (only KeyboardInterrupt and deliberate SystemExit pass through).
    """
    try:
        await _main_impl(argv)
    except KeyboardInterrupt:
        print("\n  ⚠️  interrupted", file=sys.stderr)
        sys.exit(130)
    except SystemExit:
        raise
    except Exception as e:
        logger.exception("Unhandled error during %s", " ".join(argv or []))
        print(f"  ⚠️  unexpected error: {e}", file=sys.stderr)
        print("  (run with --action check to inspect vault integrity; see "
              "hoardcore_data/*/vault.db for details)", file=sys.stderr)
        sys.exit(2)


async def _main_impl(argv: list[str] | None = None) -> None:
    args = _build_parser().parse_args(argv)
    if args.log_level:
        level = getattr(logging, args.log_level.upper(), logging.INFO)
        logging.getLogger().setLevel(level)
        logger.setLevel(level)
    url = args.url
    action = args.action
    strategy = args.strategy
    query = args.query
    claim = args.claim
    force_refresh = args.force
    urls = re.split(r'[,\s]+', args.urls.strip()) if args.urls else None
    urls = [u for u in urls if u] if urls else None
    max_results = args.limit
    discover = args.discover
    recall = args.recall
    out_path = args.out_path
    vault_name = args.vault_name
    migrate_page_size = args.migrate
    mode = args.mode

    # --vault accepts a comma/space list ("career,negros_ai_jobs"); the first
    # name is the primary (write) vault and the rest are read-only companions.
    vault_names = ([n for n in re.split(r"[,\s]+", vault_name) if n]
                   if vault_name else None)
    scraper = HoardCore(vault_names or None)
    # --parallel overrides config indexer.parallel for this run only (in-memory,
    # never written to hoardcore.toml). Read after engine construction so the
    # per-run override is authoritative.
    if args.parallel is not None:
        scraper.config._config.setdefault("indexer", {})["parallel"] = args.parallel
    print(f"\n🚀 HoardCore v{__version__}: Action={action}, URL={url}, Strategy={strategy or 'default'}")
    organized = scraper.organize_artifacts_by_day()
    if organized:
        print(f"   📂 Organized {len(organized)} artifact(s) into day folders")
    vaults = getattr(scraper, 'vaults', None) or [scraper.vault]
    vault_label = ", ".join(getattr(v, 'vault_name', None) or '(default)' for v in vaults)
    print(f"   📁 Vault: {scraper.vault.root_dir}/vault.db"
          + (f" + {len(vaults) - 1} read-only companion(s) [{vault_label}]"
             if len(vaults) > 1 else "")
          + f" | 🏛 Artifacts: {scraper.artifacts_dir}/")
    if len(vaults) > 1:
        print(f"   ⚠️  Cross-vault recall pools {len(vaults)} vault(s) [{vault_label}] — "
              "only combine one coherent context; drop off-topic hits at recall "
              "(each chunk is tagged | vault <name>).")

    if action == "research":
        if not query:
            print("  ⚠️  --query required for --action research", file=sys.stderr)
            sys.exit(2)
        written = await scraper.research(query, out_path=out_path,
                                         discover=discover if discover is not None else 5, recall=recall,
                                         strategy=strategy,
                                         answer_first=not args.no_answer_first,
                                         keep_low=args.keep_low)
        sys.exit(0 if written else 1)

    if action == "verify":
        claim = args.claim
        if args.claim_list and (args.claim or args.claim_file):
            print("  ⚠️  --claim-list cannot be combined with --claim/--claim-file",
                  file=sys.stderr)
            sys.exit(2)
        if args.claim_list:
            try:
                with open(args.claim_list, encoding="utf-8") as cf:
                    claims = [ln.strip() for ln in cf.read().splitlines()
                              if ln.strip() and not ln.strip().startswith("#")]
            except OSError as e:
                print(f"  ⚠️  cannot read --claim-list: {e}", file=sys.stderr)
                sys.exit(2)
            if not claims:
                print("  ⚠️  --claim-list contains no claims", file=sys.stderr)
                sys.exit(2)
            tally: dict[str, int] = {"verified": 0, "partial": 0, "unverified": 0}
            print(f"=== Batch verify: {len(claims)} claim(s) ===")
            for i, c in enumerate(claims, 1):
                result = scraper.verify_claim(c)
                tally[result] = tally.get(result, 0) + 1
                print(f"[{i:>3}] {result.upper():<11} {c}")
            total = len(claims)
            pct = 100.0 * tally["verified"] / total
            print(f"=== citation accuracy: {tally['verified']}/{total} = {pct:.1f}% "
                  f"(verified {tally['verified']}, partial {tally['partial']}, "
                  f"unverified {tally['unverified']}) ===")
            sys.exit(2 if tally["unverified"] else (1 if tally["partial"] else 0))
        if args.claim_file:
            try:
                with open(args.claim_file, encoding="utf-8") as cf:
                    claim = cf.read().strip()
            except OSError as e:
                print(f"  ⚠️  cannot read --claim-file: {e}", file=sys.stderr)
                sys.exit(2)
        if not claim:
            print("  ⚠️  --claim (or --claim-file) required for --action verify", file=sys.stderr)
            sys.exit(2)
        result = scraper.verify_claim(claim)
        print(f"VERIFY: {result.upper()}")
        print(f"claim: {claim}")
        # On denial, coach (--hint): show the nearest vault phrase so the agent
        # can re-express the claim in the source's own words and re-run (the
        # exact-phrasing contract), instead of dead-ending.
        if args.hint and result != "verified":
            hint = scraper.verify_hint(claim, recall=args.recall or 5)
            if hint:
                print(hint)
        # exit codes: 0=verified, 1=partial, 2=unverified (CI-wireable)
        sys.exit(0 if result == "verified" else (1 if result == "partial" else 2))

    if action == "audit":
        if not args.artifact:
            print("  ⚠️  --artifact PATH required for --action audit", file=sys.stderr)
            sys.exit(2)
        if not os.path.exists(args.artifact):
            print(f"  ⚠️  artifact not found: {args.artifact}", file=sys.stderr)
            sys.exit(2)
        audit = scraper.audit_artifact(args.artifact)
        print(f"=== Audit: {audit['path']} ===")
        print(f"Claim chain: {audit['total']} [V] tag(s) analyzed")
        for i, c in enumerate(audit["claims"], 1):
            src = f"[V#{c['n']}] " if c["n"] else "[V] "
            preview = (c["text"] if len(c["text"]) <= 72
                       else c["text"][:69] + "...")
            print(f"[{i:>3}] {c['verdict'].upper():<10} {src}{preview}")
            if c["verdict"] != "verified" and c.get("hint"):
                for hint_line in c["hint"].splitlines():
                    print(f"      {hint_line.strip()}")
        total = audit["total"]
        acc = audit["accuracy"] * 100.0
        c = audit["counts"]
        print(f"=== citation accuracy: {c['verified']}/{total} = {acc:.1f}% "
              f"(verified {c['verified']}, partial {c['partial']}, "
              f"unverified {c['unverified']}) ===")
        if audit["unmapped"]:
            print(f"  ✗ Source-link mapping MISSING for [V#{'],[V#'.join(audit['unmapped'])}]")
        else:
            print("  ✓ every [V#N] maps to a Source Link")
        if audit["not_ingested"]:
            for n, url in audit["not_ingested"]:
                print(f"  ✗ [V#{n}] cites {url} — no chunks in the vault")
        else:
            print("  ✓ every cited source has chunks in the vault")
        if audit.get("warnings"):
            for w in audit["warnings"]:
                print(f"  ⚠ line {w['line']}: {w['tag']} sits on an [H]/[E] "
                      f"analysis line — move the tag to the verbatim quote "
                      f"in the body, or demote the line to [E]")
        bad_map = bool(audit["unmapped"] or audit["not_ingested"])
        sys.exit(2 if (c["unverified"] or bad_map)
                 else (1 if c["partial"] else (0 if total else 0)))

    if action == "local":
        local_dir = scraper.config.get('storage.local_dir', 'local_inputs')
        relpath = args.local_path or "."
        if args.list_only:
            print(f"=== Local scan: {local_dir}/ ({relpath}) ===")
            entries = scraper.scan_local(relpath)
            errors = [e for e in entries if "error" in e]
            if errors:
                for e in errors:
                    print(f"  ⚠️  {e['error']}", file=sys.stderr)
                sys.exit(2)
            if not entries:
                print(f"  (no supported files under {relpath!r})")
            for e in entries:
                print(f"  {e['path']}  ({e['bytes']} B)")
            sys.exit(0)
        print(f"=== Local ingest: {local_dir}/ ({relpath}) ===")
        results = await scraper.local_ingest(relpath, force_refresh=args.force)
        print(f"✅ Done. Ingested {len(results)} chunks from local files.")
        sys.exit(0 if results else 1)

    if action == "check":
        if migrate_page_size:
            migrated = scraper.vault.migrate_page_size()
            print(f"  🔧 Page size: {'migrated to 16 KB' if migrated else 'already at target'}")
        ok = scraper.vault.verify_vault()
        sys.exit(0 if ok else 1)

    if action == "stats":
        for vault in scraper.vaults:
            if len(scraper.vaults) > 1:
                print(f"\n  === {vault.vault_name or '(default)'} ===")
            st = vault.stats()
            print(f"  Vault:      {st['vault'] or '(default)'}")
            print(f"  Sources:    {st['sources']} distinct URLs")
            print(f"  Versions:   {st['doc_versions']} document rows")
            print(f"  Chunks:     {st['chunks']}")
            print(f"  Vectors:    {st['vectors']}")
            print(f"  Embedding:  {st['mode']}, dim {st['dim']}")
            print(f"  Conf mode:  {st['conf_mode']}")
            print(f"  Schema:     v{st['schema_version']} | page {st['page_size']} B")
            mb = st['db_bytes'] / (1024 * 1024)
            print(f"  DB size:    {mb:.1f} MiB")
            try:
                dist = vault.confidence_distribution()
                if any(dist.values()):
                    print(f"  Conf probe: high {dist['high']} | medium {dist['medium']} | low {dist['low']} "
                          "(sampled, set-relative)")
            except Exception as e:
                logger.debug(f"confidence probe skipped: {e}")
        sys.exit(0)

    try:
        result = await scraper.fetch(
            url, action=action, strategy=strategy,
            query=query, force_refresh=force_refresh,
            urls=urls,
            max_results=max_results, mode=mode
        )
    except RuntimeError as e:
        marker = str(e)
        if marker == "SSRF_BLOCKED":
            print("  ⛔ SSRF guard refused the target (private/LAN/loopback/"
                  "non-http URL).", file=sys.stderr)
            print("  Set network.ssrf_protection=false only for trusted internal "
                  "targets.", file=sys.stderr)
        elif marker == "CF_COOKIE_EXPIRED":
            print("  ☁️  Cloudflare challenge not cleared by the current cookie "
                  "string.", file=sys.stderr)
            print("  Update auth.cookie_string in hoardcore.toml (or raise "
                  "network.timeout_seconds) and retry.", file=sys.stderr)
        elif marker == "FETCH_FAILED":
            print("  ⚠️  Fetch chain exhausted (aiohttp →" +
                  (" curl_cffi →" if CURL_AVAILABLE else "") +
                  " FlareSolverr).", file=sys.stderr)
        else:
            print(f"  ⚠️  {e}", file=sys.stderr)
        sys.exit(2)

    print("\n" + "=" * 80)
    print(f"✅ Done. Returned {len(result)} chunks.")

    # Preview
    for i, chunk in enumerate(result[:3]):
        print(f"\n--- CHUNK {i+1} ---")
        print(f"Metadata: {chunk['metadata']}")
        preview = chunk['text'][:300] + "..." if len(chunk['text']) > 300 else chunk['text']
        print(f"Preview: {preview}")

    if len(result) > 3:
        print(f"\n... and {len(result) - 3} more chunks.")

if __name__ == "__main__":
    asyncio.run(main())
else:

    def cli_entry() -> None:
        """Synchronous console-script entry point (setuptools)."""
        asyncio.run(main())
